import csv
from docx import Document
from docx.shared import Cm

def obtener_cometidos(self):
    # Obtener los cometidos del cargo
    if self.cargo.departamento == "Salud":
        self.lista_linas_cometidos = []
        # Leer el archivo CSV y filtrar las filas que coincidan con self.cargo.serv_salud
        with open('cargos_unicos_con_categoria.csv', newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile, delimiter='|')
            for row in reader:
                if row[0].strip().lower() == self.cargo.serv_salud.strip().lower():
                    self.lista_linas_cometidos.append(row[1].strip())
        
        # Crear un documento de Word
        doc = Document()
        
        # Escribir las filas correspondientes en el documento
        self.escribir_parrafos_con_vinetas(self.lista_linas_cometidos, list(range(len(self.lista_linas_cometidos))), doc, margen_izquierdo=Cm(1), margen_derecho=Cm(1))
        
        # Guardar el documento (opcional)
        doc.save('cometidos_salud.docx')
    else:
        self.linea_cometido = self.cargo.text_cometido_sa  # esto en vd quizas no neceita la funcion
    return self.cometidos

def escribir_parrafos_con_vinetas(list_p_cometidos, indices_filas, doc, margen_izquierdo=Cm(1), margen_derecho=Cm(1)):
    """
    Escribe párrafos con viñetas a partir de una lista de textos en un documento de Word.
    
    :param list_p_cometidos: Lista de textos que corresponden a los párrafos.
    :param indices_filas: Lista de índices de las filas que se escribirán.
    :param doc: Documento de Word en el cual se está escribiendo.
    :param margen_izquierdo: Margen izquierdo en centímetros.
    :param margen_derecho: Margen derecho en centímetros.
    """
    # Iterar sobre los índices especificados y agregar párrafos con viñetas
    for i in indices_filas:
        texto = list_p_cometidos[i]
        p = doc.add_paragraph()
        p.add_run(texto)
        p.style = 'List Bullet'
        
        # Configurar los márgenes del párrafo
        p.paragraph_format.left_indent = margen_izquierdo
        p.paragraph_format.right_indent = margen_derecho