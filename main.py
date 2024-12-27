"""
Autor: JK Fienco
Correo: jkfienco@gmail.com
Fecha de creación: 27 de diciembre de 2024
Descripción: Automatización para la creación de documentos (docx) de decretos y contratos a honorarios para la Municipalidad de Vitacura.
Versión: 1.12
Licencia: GNU General Public License v3.0
"""
import sys
import os
import re
import ctypes
import locale
import csv
import pandas as pd
import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox
import ttkbootstrap as ttkb
from ttkbootstrap.widgets import DateEntry
from ttkbootstrap.constants import *
#from tkcalendar import DateEntry
from datetime import datetime
from dateutil.relativedelta import relativedelta

from docx import Document 
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Establecer la configuración regional a español, así escribiremos los meses en español
locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')

# Obtener el path con os getcwd: C:\Users\jfz
path_os_getcwd = os.getcwd()
print("El path de user actual (os_getcwd) es:", path_os_getcwd)

# Obtenemos la ruta del directorio del archivo de script actual, en la aplicación debeía ser: C:\Program Files (x86)\AutoPersonalHonorario
path_app = os.path.dirname(os.path.abspath(__file__))
print("El path del script (path_app) es:", path_app)

path_instalador = r"C:\Program Files (x86)\AutoPersonalHonorario"
print("El path del instalador es:", path_instalador)

# Path a la carpeta en donde creamos el codigo
path_jk = r"C:\Users\jfz\OneDrive - Municipalidad de Vitacura\Documentos\Auto Personas\cx_Freeze"
print("El path usado por JK es:", path_jk)

if path_app.lower() == path_jk.lower():
    path_a_utilizar = path_app # hacemos minuscula para comparar, ya que a veces "C" es minuscula y a veces es mayuscula
else:
    path_a_utilizar = path_instalador

# Utilizaremos el path app

class Persona:
    def __init__(self):

        self.reset()

    def reset(self):
        self.nombre= ""
        self.rut = ""
        self.dig_ver = ""
        self.validez_rut = tk.BooleanVar()
        self.genero = ""
        self.domicilio = ""
        self.nro_domicilio = ""
        self.comuna = ""
        self.aclaracion_domicilio = ""
        self.domicilio_part_o_lab = ""
        self.mail = ""
        self.validez_mail = tk.BooleanVar()
        self.nacionalidad = ""
        self.estado_civil = ""
        self.profesion = ""
        self.beneficios_seleccionados = {}
        self.nombres_beneficios_seleccionados = []
        self.maternidad_seleccionada = ""
        self.sala_cuna = False

class Cargo:
    def __init__(self):

        self.reset()
    
    def reset(self):
        self.tipo_decretoycontrato = ""
        self.numero_y_text_cometido_sa = ""
        self.numero_cometido_sa = ""
        self.text_cometido_sa = ""
        self.text_cometido_salud = ""
        self.programa = ""
        self.departamento = ""
        self.direccion = ""
        self.serv_salud = ""
        self.tipo_contrato = ""
        self.cert_presup = ""
        self.nro_memo_dir = ""
        self.fecha_memo_dir = ""
        self.nro_memo_pers = ""
        self.fecha_memo_pers = ""
        self.renta_bruta = ""
        self.fecha_inicio = ""
        self.fecha_termino = ""
        self.fecha_memo_recep =""
        self.fecha_instrumento_contrato = ""
        self.subtitulo_cuenta = ""
        self.item_cuenta = ""
        self.asig_cuenta = ""
        self.subasig_cuenta = ""
        self.subsubasig_cuenta = ""
        self.fechas_dias_entries = []
        self.nro_cuenta = ""
        self.cuenta_con_puntos = ""
        self.nro_ccosto = ""

        self.text_cometido_educacion = ""

        self.tipo_renta = ""
        self.renta_bruta = ""
        self.renta_liquida = ""
        #self.tope_max_anual = ""
        self.renta_uf_clp = ""
        self.mensual_final = ""
        self.periodo_contratacion = ""
        self.medio_validador = ""

        self.nro_fechas = ""
        self.fecha_dia_contratacion = ""
        self.fechas_dias_contratacion = []

        self.especialidad_salud = ""

        self.nro_decreto_5a1 = "" # creación de programa
        self.fecha_decreto_5a1 = ""
        self.nro_decreto_5a2 = ""
        self.fecha_decreto_5a2 = ""

        self.nro_memo_dem = ""
        self.fecha_memo_dem = ""

        # Para regularización
        self.reg_vigente_o_vencido = ""
        self.motivo_regularizacion = ""

        # Para modificación
        self.tipo_solicitud = "" 
        self.fecha_solicitud = ""
        self.tipo_modificacion = ""
        self.tipo_aprobacion = ""
        self.nro_decreto_siaper = ""
        self.fecha_decreto_siaper = ""

        self.alcaldia = ""
        self.secre_muni_ejercicio_o_subrogancia = ""

        self.visadora_1 = ""
        self.visadora_2 = ""
        self.redactora = ""

class Aplicacion:
    """
    Clase principal para la aplicación de Tkinter.

    Atributos:
    root (tk.Tk): La ventana principal de la aplicación.
    cargo (objeto): Objeto que contiene la información del cargo.
    persona (objeto): Objeto que contiene la información de la persona.
    """
    def __init__(self, root):# Inicializa la clase Aplicacion
        self.root = root
        self.inicializar_interfaz()

    def inicializar_interfaz(self):
        # Configurar la ventana principal
        self.root.title("Automatización de Decretos y Contratos")

        # Inicializa c_decreto con un DataFrame vacío o con los datos necesarios
        self.c_decreto = pd.DataFrame(columns=["texto"])

        # Obtener el ancho y alto de la pantalla del usuario
        self.screen_width = self.root.winfo_screenwidth()
        self.screen_height = self.root.winfo_screenheight()
        print(f"Ancho de la pantalla: {self.screen_width}, Alto de la pantalla: {self.screen_height}")

        # Ajustar la geometría de la ventana principal a x/100 % del tamaño de la pantalla
        self.new_width = int(self.screen_width*1.3)
        self.new_height = int(self.screen_height*1.2)
        print(f"Ancho NUEVO de la ventana: {self.new_width}, Alto NUEVO de la ventana: {self.new_height}")
        #self.root.geometry(f"{self.new_width}x{self.new_height}")

        # Centrar la ventana en la pantalla
        self.x_offset = -int((self.screen_width - self.new_width)/2)
        self.y_offset = -int((self.screen_height - self.new_height)/2)
        print(f"Offset en x: {self.x_offset}, Offset en y: {self.y_offset}")

        # Establecer la geometría de la ventana con tamaño y posición
        self.root.geometry(f"{self.new_width+self.x_offset}x{self.new_height+self.y_offset}")
        # Establecer la ventana en modo de pantalla completa
        #self.root.attributes('-fullscreen', True)

        # Mejora calidad de imagen de la interfaz a crear
        try:
            ctypes.windll.shcore.SetProcessDpiAwareness(1) # 0 se ve mal, 1 y 2 funcionan bien
        except AttributeError:
            pass  # La función no está disponible en esta versión de Windows
        
        # Icono de vitacura .ico en la barra superior, para el ejecutable
        icon_path = os.path.join(sys._MEIPASS, 'logo-vitacura_icono.ico') if hasattr(sys, '_MEIPASS') else 'logo-vitacura_icono.ico'
        if os.path.exists(icon_path):
            self.root.iconbitmap(icon_path)
        else:
            print(f"Icono no encontrado en la ruta: {icon_path}")

        # Crear una instancia de la clase Cargo
        self.cargo = Cargo()

        # Crear una instancia de la clase Persona (empleada)
        self.persona = Persona()
        self.persona.validez_rut.trace_add("write", self.verificar_rut) # con el modo "write", lo que significa que la función verificar_rut se llamará cada vez que validez_rut sea modificada
        self.persona.validez_mail.trace_add("write", self.validar_correo)

        self.subtitulo_var = tk.StringVar()
        self.item_var = tk.StringVar()
        self.asig_var = tk.StringVar()
        self.subasig_var = tk.StringVar()
        self.subsubasig_var = tk.StringVar()

        # Resto del código de inicialización
        self.crear_interfaz()

    def reiniciar_interfaz(self):
        # Destruir todos los widgets actuales
        for widget in self.root.winfo_children():
            widget.destroy()

        self.inicializar_interfaz()

    def reiniciar_variables_interfaz(self):
        # Resetea los widgets de la interfaz
        self.reiniciar_interfaz()
        # Resetear los atributos de Persona y Cargo
        self.persona.reset()
        self.cargo.reset()

    def crear_interfaz(self):
        # Crear un estilo personalizado para el Notebook
        style = ttkb.Style()
        style.configure('TNotebook', foreground="black", font=("", 11, "bold"))  # Cambia el tamaño de la fuente aquí
        style.configure('TNotebook.Tab', padding=[5, 5], font=("", 11, ""))  # Cambia el tamaño de la fuente para las pestañas
        style.configure('TLabel', font=("", 11))  # Cambia el tamaño de la fuente para etiquetas
        # Configurar el estilo del DateEntry
        # style.configure('DateEntry', fieldbackground='darkblue', foreground='white', borderwidth=2)
        #style.configure('TCombobox', font=("", 11))  # Cambia el tamaño de la fuente para combobox
        #style.configure('TComboboxPopdownFrame', font=("", 11))  # Cambia el tamaño de la fuente para la lista desplegable del combobox
        #style.configure('TComboboxPopdownListbox.TListbox', font=("", 11))  # Cambia el tamaño de la fuente para los elementos de la lista desplegable
        style.configure('TButton', font=("", 11))
        style.map("TNotebook.Tab", background=[("selected", "#b0fff8")])  # Color de fondo cuando está seleccionada, otro c5d707. b0fff8 es el aclarado al 75% de 00c4b4

        # Establecer un ancho fijo para las pestañas
        style.configure('TNotebook.Tab', width=15)  # Ajusta el valor según sea necesario

        # Crear un widget Notebook con varios parámetros
        self.notebook = ttkb.Notebook(self.root, style="TNotebook", takefocus=True, padding=10, height=self.new_width+self.x_offset, width=self.new_height+self.y_offset)
        self.notebook.pack(expand=True, fill='both')

        # Crear el primer frame (pestaña)
        self.frame1 = ttk.Frame(self.notebook)
        self.notebook.add(self.frame1, text='Regularización')

        # Crear el segundo frame (pestaña)
        self.frame2 = ttk.Frame(self.notebook)
        self.notebook.add(self.frame2, text='En fecha')

        # Crear el tercer frame (pestaña)
        self.frame3 = ttk.Frame(self.notebook)
        self.notebook.add(self.frame3, text='Modificación')

        # Vincular el evento <<NotebookTabChanged>> al método on_tab_changed
        self.notebook.bind("<<NotebookTabChanged>>", self.on_tab_changed)

    def on_tab_changed(self, event):
        # Obtener el índice de la pestaña activa
        global current_tab_index
        current_tab_index = self.notebook.index(self.notebook.select())
        print(f"Índice de la pestaña activa: {current_tab_index}")

        # Verificar cuál pestaña está activa
        if current_tab_index == 0:
            self.cargo.tipo_decretoycontrato = "Regularización"
            print("Pestaña 'Regularización' activa")
            self.pestana_regularizacion()

        elif current_tab_index == 1:
            self.cargo.tipo_decretoycontrato = "En fecha"
            print("Pestaña 'En fecha' activa")
            self.pestana_en_fecha()

        elif current_tab_index == 2:
            self.cargo.tipo_decretoycontrato = "Modificación"
            print("Pestaña 'Modificación' activa")
            self.pestana_modificacion()

    def agregar_titulo(self, frame_name, titulo, indice_row, indice_col):
        label_titulo = ttk.Label(frame_name, text=titulo, font=("", 14))
        label_titulo.grid(row=indice_row, column=indice_col, pady=5, padx=5, sticky="w")

    def crear_lista_from_csv(self, nombre_csv, nro_columna, separador="|"):
        # Crear una lista a partir de un archivo csv
        df = pd.read_csv(nombre_csv, sep=separador, header='infer')
        lista = df.iloc[:, nro_columna].tolist()
        return lista

    def pestana_regularizacion(self): # Creación de frame 1 (Regularización)
        # Título sección
        global i_cargo_row # pendiente, colocar afuera de las funciones, al comenzar el codigo
        global i_cargo_col
        i_cargo_row = 0
        i_cargo_col = 0
        self.agregar_titulo(self.frame1, "Información del cargo", i_cargo_row, i_cargo_col)

        # Dirección solicitante, presente en todos los tipo de decreto
        self.det_direccion(self.frame1)
        self.tipo_contrato(self.frame1)

        # Título sección
        global i_personas_row  # pendiente, colocar afuera de las funciones, al comenzar el codigo
        i_personas_row = 0
        global i_personas_col
        i_personas_col = 3

        self.agregar_titulo(self.frame1, "Información de la persona", i_personas_row, i_personas_col)
        self.cuenta_renta_bruta_y_medio_validador(self.frame1, i_cargo_row+11) # columna 0 y 1

        self.info_persona(self.frame1) # desde columna 3
        self.beneficios(self.frame1, i_personas_row+12)# desde columna 3

        self.visacion_y_redaccion(self.frame1, i_personas_row+14)

        # Botón para guardar datos, elegir directorio y crear documento word
        guardar_button = ttk.Button(self.contenedor_visa_redacta, text="Crear documentos", command=self.guardar_en_path, width=18) 
        guardar_button.grid(row=0, column=6, pady=5, padx=35, sticky="e")

        # Botón para resetear las variables de la interfaz
        reset_button = ttk.Button(self.contenedor_visa_redacta, text="Resetear", command=self.reiniciar_variables_interfaz, width=10)
        reset_button.grid(row=0, column=7, pady=5, padx=35, sticky="e")

    def pestana_en_fecha(self): # Creación de frame 2 (En fecha)
        # Título sección
        self.agregar_titulo(self.frame2, "Información del cargo", i_cargo_row, i_cargo_col)

        # Dirección solicitante, presente en todos los tipo de decreto
        self.det_direccion(self.frame2)
        self.tipo_contrato(self.frame2)

        # Título sección
        self.agregar_titulo(self.frame2, "Información de la persona", i_personas_row, i_personas_col)
        self.cuenta_renta_bruta_y_medio_validador(self.frame2, i_cargo_row+11) # columna 0 y 1

        self.info_persona(self.frame2) # desde columna 3
        self.beneficios(self.frame2, i_personas_row+12) # desde columna 3
        
        #self.maternidad(self.frame2, i_personas_row+13)

        self.visacion_y_redaccion(self.frame2, i_personas_row+14)

        # Botón para guardar datos
        guardar_button = ttk.Button(self.contenedor_visa_redacta, text="Crear documento", command=self.guardar_en_path, width=15)
        guardar_button.grid(row=0, column=6, pady=5, padx=35, sticky="e")

        # Botón para resetear las variables de la interfaz
        reset_button = ttk.Button(self.contenedor_visa_redacta, text="Resetear", command=self.reiniciar_variables_interfaz, width=10)
        reset_button.grid(row=0, column=7, pady=5, padx=35, sticky="e")

    def pestana_modificacion(self): # Creación de frame 3 (Modificación)
        # Título sección
        self.agregar_titulo(self.frame3, "Información del cargo", i_cargo_row, i_cargo_col)

        self.det_direccion(self.frame3)# Dirección solicitante, presente en todos los tipo de decreto
        self.tipo_solicitud(self.frame3)
        self.tipo_modificacion(self.frame3)
        self.tipo_aprobacion(self.frame3)
        self.detalles_del_contrato(self.frame3)

        # Título sección
        self.agregar_titulo(self.frame3, "Información de la persona", i_personas_row, i_personas_col)
        self.cuenta_renta_bruta_y_medio_validador(self.frame3, i_cargo_row+11) # columna 0 y 1

        self.info_persona(self.frame3) # desde columna 3
        self.beneficios(self.frame3, i_personas_row+12) # desde columna 3

        self.visacion_y_redaccion(self.frame3, i_personas_row+14)

        # Botón para guardar datos
        guardar_button = ttk.Button(self.contenedor_visa_redacta, text="Crear documento", command=self.guardar_en_path, width=15) 
        guardar_button.grid(row=0, column=6, pady=5, padx=35, sticky="e")

        # Botón para resetear las variables de la interfaz
        reset_button = ttk.Button(self.contenedor_visa_redacta, text="Resetear", command=self.reiniciar_variables_interfaz, width=10)
        reset_button.grid(row=0, column=7, pady=5, padx=35, sticky="e")

    def visacion_y_redaccion(self, frame_name, fila_visa_redacta):

        self.entrada_visa_1 = tk.StringVar()
        self.entrada_visa_1.set("RGP")

        self.entrada_visa_2 = tk.StringVar()
        self.entrada_visa_2.set("PFH")

        self.entrada_redacta = tk.StringVar()
        self.entrada_redacta.set("")
        
        self.label_distrib_visa = ttk.Label(master=frame_name, text="Distribución visadores:", style = "TLabel")
        self.label_distrib_visa.grid(row=fila_visa_redacta, column=3, padx=5, pady=5, sticky="w")

        self.contenedor_visa_redacta = ttk.Frame(frame_name)
        self.contenedor_visa_redacta.grid(row=fila_visa_redacta, column=4, pady=5, sticky="w")

        self.label_visa_1 = ttk.Label(master=self.contenedor_visa_redacta, text="Visa", style = "TLabel")
        self.label_visa_1.grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.entry_visa_1 = ttk.Entry(self.contenedor_visa_redacta, font=("", 11), textvariable=self.entrada_visa_1, width=4) # idea: hacer lista desglosable con combobox
        self.entry_visa_1.grid(row=0, column=1, padx=5, pady=5, sticky="w")

        self.label_visa_2 = ttk.Label(master=self.contenedor_visa_redacta, text="Visa:", style = "TLabel")
        self.label_visa_2.grid(row=0, column=2, padx=5, pady=5, sticky="w")
        self.entry_visa_2 = ttk.Entry(self.contenedor_visa_redacta, font=("", 11), textvariable=self.entrada_visa_2, width=4) # idea: hacer lista desglosable con combobox
        self.entry_visa_2.grid(row=0, column=3, padx=5, pady=5, sticky="w")

        self.label_redacta = ttk.Label(master=self.contenedor_visa_redacta, text="Redacta:", style = "TLabel")
        self.label_redacta.grid(row=0, column=4, padx=5, pady=5, sticky="w")
        self.entry_redacta = ttk.Entry(self.contenedor_visa_redacta, font=("", 11), textvariable=self.entrada_redacta, width=4) # idea: hacer lista desglosable con combobox
        self.entry_redacta.grid(row=0, column=5, padx=5, pady=5, sticky="w")

    def det_direccion(self, frame_name):

        # Dirección solicitante
        file_path_direcciones = os.path.join(path_a_utilizar, "clausulas_csv", "nombres_direcciones.csv")
        list_direcciones = self.lista_cuentas = self.crear_lista_from_csv(file_path_direcciones, nro_columna=0, separador="|")

        label_direcciones = ttk.Label(master = frame_name, text="Dirección solicitante: ", style = "TLabel", width=30) # background="#00c4b4"
        label_direcciones.grid(row=1 ,column=0, padx=5, pady=5)

        self.direccionSolicitante = tk.StringVar()
        self.direccionSolicitante.set("Seleccionar")

        menu = ttkb.Combobox(master=frame_name, font=("", 11), width = 65, textvariable = self.direccionSolicitante, state="readonly", values = list_direcciones) # background="#00c4b4"
        menu.grid(row=1, column=1, padx=5, pady=5, sticky="w") # columnspan=3 significa que se extiende por 3 columnas

        menu.bind("<<ComboboxSelected>>", lambda event: self.departamento_salud_educacion(frame_name))

    def departamento_salud_educacion(self, frame_name):
        # Guardamos variables para condicion
        self.cargo.direccion = self.direccionSolicitante.get()

        self.depto_var = tk.StringVar() # podria ser int binario tambien com IntVarblbl() 
        self.depto_var.set("Seleccionar")

        self.serv_salud_var = tk.StringVar() # Definimos esta varible aquí para no tener problemas al guardar en caso de ser muniicipal
        self.serv_salud_var.set("Seleccionar")
    
        if self.cargo.direccion == "Dirección de Salud y Educación, y Demás Incorporados en la Gestión Municipal":
    
            lista_deptos = ["Salud", "Educación"]
        
            self.label_depto = ttk.Label(master=frame_name, text="Departamento:", style="TLabel")
            self.label_depto.grid(row=2, column=i_cargo_col, padx=5, pady=5, sticky="w") #DEPAAAA

            self.contenedor_depto_salud_ed = ttk.Frame(frame_name)
            self.contenedor_depto_salud_ed.grid(row=2, column=i_cargo_col+1, pady=5, sticky="w")
    
            self.menu_depto = ttkb.Combobox(master=self.contenedor_depto_salud_ed, font=("", 11), textvariable=self.depto_var, state="readonly", values=lista_deptos, width=15) #, background="#00c4b4"
            self.menu_depto.grid(row=0, column=0, padx=5, pady=5, sticky="w") # columnspan=3 significa que se extiende por 3 columnas

            self.menu_depto.bind("<<ComboboxSelected>>", lambda event: self.servicios_salud(frame_name))
    
        else:
            self.depto_var.set("Municipal")
            self.actualizar_tipo_contratos()
        
            # Removemos widgets de pedartamento al modificar la dirección
            if hasattr(self, 'menu_depto'):
                print("Removiendo widgets de departamento")
                self.label_depto.grid_remove()
                self.menu_depto.grid_remove()
            
        # Remueven widgets que hayan quedado si antes se eligió programa
        if hasattr(self, 'nombre_programa'):
            # Removiendo widgets de detalle programa
            self.nombre_programa.grid_remove()
            self.entrada_programa.grid_remove()
            self.label_nro_decreto_5a1.grid_remove()
            self.entrada_nro_decreto_5a1.grid_remove()
            self.date_5a1.grid_remove()
            self.cal_5a1.grid_remove()
        if hasattr(self, 'nro_memoDEM_5b'):
            self.nro_memoDEM_5b.grid_remove()
            self.entrada_nro_memoDEM_5b.grid_remove()
            self.date_5b.grid_remove()
            self.cal_5b.grid_remove()
    
    def servicios_salud(self, frame_name):
        # Guardamos variables para condicion
        self.cargo.departamento = self.depto_var.get()
    
        if self.cargo.departamento == "Salud":
    
            lista_serv_salud = ["SAPU", "CESFAM", "COSAM", "Departamento de Salud Municipal"]
        
            self.label_ss = ttk.Label(master=self.contenedor_depto_salud_ed, text="Servicio de salud:", style="TLabel")
            self.label_ss.grid(row=0, column=1, padx=5, pady=5, sticky="w")
    
            self.menu_ss = ttkb.Combobox(master=self.contenedor_depto_salud_ed, font=("", 11), textvariable=self.serv_salud_var, state="readonly", values=lista_serv_salud, width=30) # background="#00c4b4"
            self.menu_ss.grid(row=0, column=2, padx=5, pady=5, sticky="w") # columnspan=3 significa que se extiende por 3 columnas
        else:
            pass

        self.actualizar_tipo_contratos()

    def tipo_contrato(self, frame_name):
        self.label_TC = ttk.Label(master = frame_name, text="Tipo de contrato:", style="TLabel", width=30)
        self.label_TC.grid(row=i_cargo_row+3, column=0, padx=5, pady=5, sticky="w")

        self.contenedor_contrato_reg = ttk.Frame(frame_name)
        self.contenedor_contrato_reg.grid(row=i_cargo_row+3, column=1, pady=5, sticky="w")

        self.entrada_TC = tk.StringVar()
        self.entrada_TC.set("Seleccionar")

        self.menu_TC = ttkb.Combobox(self.contenedor_contrato_reg, font=("", 11), textvariable = self.entrada_TC, state="readonly", values = ["Seleccionar"], width=20)
        self.menu_TC.grid(row=0, column=1, padx=5, pady=5, sticky="w")

        if current_tab_index == 0 or current_tab_index == 1: # Reg o en fecha ## MOVER AL LADO DERECHO DE TIPO DE CONTRATO
            # Fecha instrumento contrato
            self.fecha_instrumento_contrato_var = tk.StringVar()
            self.fecha_instrumento_contrato_var.set("")

            self.label_instrumento_contrato = ttk.Label(self.contenedor_contrato_reg, text='Fecha instrumento contrato:', style="TLabel") 
            self.label_instrumento_contrato.grid(row=0, column=2, padx=5, pady=5, sticky="w")

            self.cal_instrumento_contrato = DateEntry(self.contenedor_contrato_reg, width=17, bootstyle="primary", firstweekday=0, dateformat='%d/%m/%Y')
            self.cal_instrumento_contrato.grid(row=0, column=3, padx=5, pady=5, sticky="w")

        self.menu_TC.bind("<<ComboboxSelected>>", lambda event: self.detalles_del_contrato(frame_name))


    def actualizar_tipo_contratos(self):
        self.cargo.departamento = self.depto_var.get()

        if self.cargo.departamento == "Educación":
            print("Departamento Educación")

            list_contratos_ed = ["SEP Educación", "Mejoramiento a la educación", "FAEP"] # Para salud es "Programa"

            self.menu_TC.config(values = list_contratos_ed)
            # Removemos widgets de servicio salud al modificar el depto
            if hasattr(self, 'menu_ss'):
                self.menu_ss.grid_remove()
            if hasattr(self, 'label_ss'):
                self.label_ss.grid_remove()
            
            # Remueven widgets que hayan quedado si antes se eligió programa
            if hasattr(self, 'nombre_programa'):
                # Removiendo widgets de detalle programa
                self.nombre_programa.grid_remove()
                self.entrada_programa.grid_remove()
                self.label_nro_decreto_5a1.grid_remove()
                self.entrada_nro_decreto_5a1.grid_remove()
                self.date_5a1.grid_remove()
                self.cal_5a1.grid_remove()

        elif self.cargo.departamento == "Salud":
            print("Departamento Salud")
            list_contratos_salud = ["Programa", "Suma alzada"]

            self.menu_TC.config(values = list_contratos_salud)

            # Removemos widgets de detalles de contrato si antes se eligió SEP Educación, Mejoramiento a la educación o FAEP
            if hasattr(self, 'nro_memoDEM_5b'):
                self.nro_memoDEM_5b.grid_remove()
                self.entrada_nro_memoDEM_5b.grid_remove()
                self.date_5b.grid_remove()
                self.cal_5b.grid_remove()

        elif self.cargo.departamento == "Municipal":
            list_contratos_muni = ["Programa", "Suma alzada"]
            print("Departamento Municipal")
            self.menu_TC.config(values = list_contratos_muni)
            # Removemos widgets de servicio salud al modificar el depto
            if hasattr(self, 'menu_ss'):
                self.menu_ss.grid_remove()
            if hasattr(self, 'label_ss'):
                self.label_ss.grid_remove()

    def trans_cuenta(self, nro_cuenta_sin_puntos):
        # Cadena original: nro_cuenta_sin_puntos
        # ejemplo: valor = "2152103001000"

        valor = str(nro_cuenta_sin_puntos) # Convertir a cadena en caso de haber interpretado el número de cuenta como numero

        # Rebanar la cadena en partes
        parte1 = valor[:3] # Cuenta de gastos
        parte2 = valor[3:5] # Subtitulo
        parte3 = valor[5:7] # Item
        parte4 = valor[7:10] # Asignación
        parte5 = valor[10:] # Subasignación

        # Dejamos preescritos los valores de la cuenta
        self.subtitulo_var.set(parte2)
        self.item_var.set(parte3)
        self.asig_var.set(parte4)
        self.subasig_var.set(parte5)

        parte6 = valor[13:] # Subsubasignación
        self.subsubasig_var.set(parte6)

        resultado = f"{parte1}.{parte2}.{parte3}.{parte4}.{parte5}.{parte6}" # hace segunda opcino

        # Retorna el resultado
        return resultado # ejemplo: salida = 215.21.03.001.000.000


    def detalles_del_contrato(self, frame_name):
        self.cargo.tipo_contrato = self.entrada_TC.get()

        if  self.cargo.tipo_contrato == "Programa":
            # Removemos widgets de detalles de contrato si antes se eligió SEP Educación, Mejoramiento a la educación o FAEP
            if hasattr(self, 'nro_memoDEM_5b'):
                self.nro_memoDEM_5b.grid_remove()
                self.entrada_nro_memoDEM_5b.grid_remove()
                self.date_5b.grid_remove()
                self.cal_5b.grid_remove()

            # Nombre del programa
            self.entrada_prog_var = tk.StringVar()
            self.entrada_prog_var.set("Seleccionar")

            self.nro_decreto_5a1_var = tk.StringVar()
            self.nro_decreto_5a1_var.set("")

            self.fecha_5a1_var = tk.StringVar()
            self.fecha_5a1_var.set("")

            self.nro_decreto_5a2_var = tk.StringVar()
            self.nro_decreto_5a2_var.set("")

            self.fecha_5a2_var = tk.StringVar()
            self.fecha_5a2_var.set("")

            self.nombre_programa = ttk.Label(master = frame_name, text='Nombre del programa:', style="TLabel") # cuando tengamos la BBDD con los programas, se podrá seleccionar
            self.nombre_programa.grid(row=i_cargo_row+4, column=0, padx=5, pady=5, sticky="w")

            # Creamos lista de cuentas y lista de programas
            file_path_programas = os.path.join(path_a_utilizar, "clausulas_csv", "cuentas_programas.csv")
            self.lista_cuentas = self.crear_lista_from_csv(file_path_programas, nro_columna=0, separador="|")
            self.lista_programas = self.crear_lista_from_csv(file_path_programas, nro_columna=1, separador="|")

            self.entrada_programa = ttkb.Combobox(frame_name, font=("", 11), textvariable=self.entrada_prog_var, state="readonly", values=self.lista_programas, width=65)
            self.entrada_programa.grid(row=i_cargo_row+4, column=1, padx=5, pady=5, sticky="w")

            self.entrada_programa.bind("<<ComboboxSelected>>", lambda event: self.obtener_cuenta)

            # Crear un contenedor para nro y fecha decreto creacion de programas
            self.contenedor_decreto_creacion_programa = ttk.Frame(frame_name)
            self.contenedor_decreto_creacion_programa.grid(row=i_cargo_row+5, column=1, pady=5, sticky="w")

            ## Nro decreto creación (5a.1)
            self.label_nro_decreto_5a1 = ttk.Label(frame_name, text='Número decreto creación programa:', style="TLabel") # cuando tengamos la BBDD con los programas, se podrá seleccionar
            self.label_nro_decreto_5a1.grid(row=i_cargo_row+5, column=0, padx=5, pady=5, sticky="w")

            self.entrada_nro_decreto_5a1 = ttk.Entry(self.contenedor_decreto_creacion_programa, font=("", 11), textvariable=self.nro_decreto_5a1_var, width=16)
            self.entrada_nro_decreto_5a1.grid(row=0, column=0, padx=5, pady=5, sticky="w")
            
            ## Fecha decreto creación (5a.1)
            self.date_5a1 = ttk.Label(self.contenedor_decreto_creacion_programa, text='Fecha creación programa:', style="TLabel", width=30)
            self.date_5a1.grid(row=0, column=1, padx=5, pady=5, sticky="w")

            self.cal_5a1 = DateEntry(self.contenedor_decreto_creacion_programa, width=17, dateformat='%d/%m/%Y', bootstyle="primary", firstweekday=0)
            self.cal_5a1.grid(row=0, column=2, padx=5, pady=5, sticky="w")

            ## Nro y fecha memo solicitud direccion solicitante (en fecha y regularización) 
            self.contenedor_memo_dir_sol = ttk.Frame(frame_name)
            self.contenedor_memo_dir_sol.grid(row=i_cargo_row+6, column=1, pady=5, sticky="w")

            self.nro_memo_dir_sol_var = tk.StringVar()
            self.nro_memo_dir_sol_var.set("")

            self.label_nro_memo_dir_sol = ttk.Label(frame_name, text='N° memo dirección solicitante:', style="TLabel")
            self.label_nro_memo_dir_sol.grid(row=i_cargo_row+6, column=0, padx=5, pady=5, sticky="w")

            self.entrada_nro_memo_dir_sol = ttk.Entry(self.contenedor_memo_dir_sol, font=("", 11), textvariable=self.nro_memo_dir_sol_var, width=16)
            self.entrada_nro_memo_dir_sol.grid(row=0, column=0, padx=5, pady=5, sticky="w")

            self.fecha_memo_sol_var = tk.StringVar()
            self.fecha_memo_sol_var.set("")

            self.label_cal_memo_sol = ttk.Label(self.contenedor_memo_dir_sol, text="Fecha memo solicitud:", style="TLabel", width=30)
            self.label_cal_memo_sol.grid(row=0, column=1, padx=5, pady=5, sticky="w")

            self.cal_memo_sol = DateEntry(self.contenedor_memo_dir_sol, width=17, dateformat='%d/%m/%Y', bootstyle="primary", firstweekday=0)
            self.cal_memo_sol.grid(row=0, column=2, padx=5, pady=5, sticky="w")

            # FALTA: Agregar nro y fecha de modificación cuando los programas sufren alguna modificacion. Estas modificaciones pieden ser más de 1.

        elif self.cargo.tipo_contrato == "SEP Educación" or self.cargo.tipo_contrato == "Mejoramiento a la educación" or self.cargo.tipo_contrato == "FAEP":
            
            # Remueven widgets que hayan quedado si antes se eligió programa
            if hasattr(self, 'nombre_programa'):
                # Removiendo widgets de detalle programa
                self.nombre_programa.grid_remove()
                self.entrada_programa.grid_remove()
                self.label_nro_decreto_5a1.grid_remove()
                self.entrada_nro_decreto_5a1.grid_remove()
                self.date_5a1.grid_remove()
                self.cal_5a1.grid_remove()
                self.label_nro_memo_dir_sol.grid_remove()
                self.contenedor_memo_dir_sol.grid_remove()

            self.contenerdor_memo_DEM = ttk.Frame(frame_name)
            self.contenerdor_memo_DEM.grid(row=i_cargo_row+4, column=1, pady=5, sticky="w")

            self.nro_memoDEM_var = tk.StringVar()
            self.nro_memoDEM_var.set("")

            self.fecha_DEM_var = tk.StringVar()
            self.fecha_DEM_var.set("")

            ## Número de memo DEM (5b)
            self.nro_memoDEM_5b = ttk.Label(master = frame_name, text='Número memo DEM:', style = "TLabel", width=30) # cuando tengamos la BBDD con los programas, se podrá seleccionar
            self.nro_memoDEM_5b.grid(row=i_cargo_row+4, column=0, padx=5, pady=5, sticky="w")

            self.entrada_nro_memoDEM_5b = ttk.Entry(self.contenerdor_memo_DEM, width=16, font=("", 11), textvariable=self.nro_memoDEM_var)
            self.entrada_nro_memoDEM_5b.grid(row=0, column=0, padx=5, pady=5, sticky="w")
            
            ## Fecha memo DEM (5b)
            self.date_5b = ttk.Label(self.contenerdor_memo_DEM, text="Fecha memo DEM:")
            self.date_5b.grid(row=0, column=1, padx=5, pady=5, sticky="w")  

            self.cal_5b = DateEntry(self.contenerdor_memo_DEM, width=12, dateformat='%d/%m/%Y', bootstyle="primary", firstweekday=0)
            self.cal_5b.grid(row=0, column=2, padx=5, pady=5, sticky="w") 

            ## Cambiamos de lugar en caso de existir los widgets: Nro y fecha memo dirección de personas (en fecha y regularización)
            if hasattr(self, 'nro_memo_dir_sol_var'):
                self.label_nro_memo_dir_pers.grid(row=i_cargo_row+5, column=0, padx=5, pady=5, sticky="w")
                self.contenedor_memo_dir_pers.grid(row=i_cargo_row+5, column=1, padx=5, pady=5, sticky="w")

            ## Cambiamos de lugar en caso de existir los widgets: Certificado factibilidad presupuestaria
            if hasattr(self, 'cert_presup_var'):
                self.label_cert_presup.grid(row=i_cargo_row+6, column=0, padx=5, pady=5, sticky="w")
                self.contenedor_cert_presup_recep.grid(row=i_cargo_row+6, column=1, padx=5, pady=5, sticky="w")
        
        elif self.cargo.tipo_contrato == "Suma alzada":
            # Remueven widgets que hayan quedado si antes se eligió programa
            if hasattr(self, 'nombre_programa'):
                # Removiendo widgets de detalle programa
                self.nombre_programa.grid_remove()
                self.entrada_programa.grid_remove()
                self.label_nro_decreto_5a1.grid_remove()
                self.entrada_nro_decreto_5a1.grid_remove()
                self.date_5a1.grid_remove()
                self.cal_5a1.grid_remove()

            # Removemos widgets de detalles de contrato si antes se eligió SEP Educación, Mejoramiento a la educación o FAEP
            if hasattr(self, 'nro_memoDEM_5b'):
                self.nro_memoDEM_5b.grid_remove()
                self.entrada_nro_memoDEM_5b.grid_remove()
                self.date_5b.grid_remove()
                self.cal_5b.grid_remove()
            
            ## Nro y fecha memo solicitud direccion solicitante (en fecha y regularización) 
            self.contenedor_memo_dir_sol = ttk.Frame(frame_name)
            self.contenedor_memo_dir_sol.grid(row=i_cargo_row+6, column=1, pady=5, sticky="w")

            self.nro_memo_dir_sol_var = tk.StringVar()
            self.nro_memo_dir_sol_var.set("")

            self.label_nro_memo_dir_sol = ttk.Label(frame_name, text='N° memo dirección solicitante:', style="TLabel")
            self.label_nro_memo_dir_sol.grid(row=i_cargo_row+6, column=0, padx=5, pady=5, sticky="w")

            self.entrada_nro_memo_dir_sol = ttk.Entry(self.contenedor_memo_dir_sol, font=("", 11), textvariable=self.nro_memo_dir_sol_var, width=16)
            self.entrada_nro_memo_dir_sol.grid(row=0, column=0, padx=5, pady=5, sticky="w")

            self.fecha_memo_sol_var = tk.StringVar()
            self.fecha_memo_sol_var.set("")

            self.label_cal_memo_sol = ttk.Label(self.contenedor_memo_dir_sol, text="Fecha memo solicitud:", style="TLabel", width=30)
            self.label_cal_memo_sol.grid(row=0, column=1, padx=5, pady=5, sticky="w")

            self.cal_memo_sol = DateEntry(self.contenedor_memo_dir_sol, width=17, dateformat='%d/%m/%Y', bootstyle="primary", firstweekday=0)
            self.cal_memo_sol.grid(row=0, column=2, padx=5, pady=5, sticky="w")

        ## Nro memo direccion de personas

        self.contenedor_memo_dir_pers = ttk.Frame(frame_name)
        self.contenedor_memo_dir_pers.grid(row=i_cargo_row+7, column=1, pady=5, sticky="w")

        self.nro_memo_dir_pers_var = tk.StringVar()
        self.nro_memo_dir_pers_var.set("")

        self.label_nro_memo_dir_pers = ttk.Label(frame_name, text='N° memo dirección de personas:', style="TLabel") 
        self.label_nro_memo_dir_pers.grid(row=i_cargo_row+7, column=0, padx=5, pady=5, sticky="w")

        self.entrada_nro_memo_dir_pers = ttk.Entry(self.contenedor_memo_dir_pers, font=("", 11), textvariable=self.nro_memo_dir_pers_var, width=16)
        self.entrada_nro_memo_dir_pers.grid(row=0, column=0, padx=5, pady=5, sticky="w")

        self.fecha_memo_pers_var = tk.StringVar()
        self.fecha_memo_pers_var.set("")

        self.label_cal_memo_pers = ttk.Label(self.contenedor_memo_dir_pers, text="Fecha memo dirección de personas:", style="TLabel", width=30)
        self.label_cal_memo_pers.grid(row=0, column=1, padx=5, pady=5, sticky="w")

        self.cal_memo_pers = DateEntry(self.contenedor_memo_dir_pers, width=17, dateformat='%d/%m/%Y', bootstyle="primary", firstweekday=0)
        self.cal_memo_pers.grid(row=0, column=2, padx=5, pady=5, sticky="w")

        ## Certificado factibilidad presupuestaria
        self.cert_presup_var = tk.StringVar()
        self.cert_presup_var.set("")
        
        self.label_cert_presup = ttk.Label(frame_name, text='N° cert. factibilidad presupuestaria:', style="TLabel") 
        self.label_cert_presup.grid(row=i_cargo_row+8, column=0, padx=5, pady=5, sticky="w")

        self.contenedor_cert_presup_recep = ttk.Frame(frame_name)
        self.contenedor_cert_presup_recep.grid(row=i_cargo_row+8, column=1, pady=5, sticky="w")

        self.entrada_cert_presup = ttk.Entry(self.contenedor_cert_presup_recep, font=("", 11), textvariable=self.cert_presup_var, width=16)
        self.entrada_cert_presup.grid(row=0, column=0, padx=5, pady=5, sticky="w")

        # Fecha recepción memo
        self.fecha_recepcion_var = tk.StringVar()
        self.fecha_recepcion_var.set("")

        self.label_fecha_recep = ttk.Label(self.contenedor_cert_presup_recep, text="Fecha recepción memo:", style="TLabel", width=30)
        self.label_fecha_recep.grid(row=0, column=1, padx=5, pady=5, sticky="w")

        self.cal_fecha_recep = DateEntry(self.contenedor_cert_presup_recep, width=17, bootstyle="primary", firstweekday=0, dateformat='%d/%m/%Y')
        self.cal_fecha_recep.grid(row=0, column=2, padx=5, pady=5, sticky="w")

        # Especialidad salud

        if self.cargo.departamento == "Salud":

            file_path_especialidades = os.path.join(path_a_utilizar, "clausulas_csv", "cargos_unicos.csv")
            lista_especialidades_salud = self.crear_lista_from_csv(file_path_especialidades, nro_columna=0, separador="|")
            lista_especialidades_salud.sort() # ordenamos la lista por orden alfabetico
            
            self.especialidad_salud_var = tk.StringVar()
            self.especialidad_salud_var.set("Seleccionar")

            self.label_esp_salud = ttk.Label(frame_name, text="Cargo:", style="TLabel", width=30)
            self.label_esp_salud.grid(row=i_cargo_row+9, column=0, padx=5, pady=5, sticky="w")

            self.menu_esp_salud = ttkb.Combobox(frame_name, font=("", 11), textvariable=self.especialidad_salud_var, state="readonly", values=lista_especialidades_salud, width=65)
            self.menu_esp_salud.grid(row=i_cargo_row+9, column=1, padx=5, pady=5, sticky="w")
        else:
            if hasattr(self, 'menu_esp_salud'):
                self.menu_esp_salud.grid_remove()
                self.label_esp_salud.grid_remove()
            pass
        
        # Opcion solo para modificación
        if current_tab_index == 2: 
            ## Nro decreto siaper y fecha, para modificación
            self.contenedor_siaper = ttk.Frame(frame_name)
            self.contenedor_siaper.grid(row=i_cargo_row+10, column=1, pady=5, sticky="w")

            self.nro_siaper_var = tk.StringVar()
            self.nro_siaper_var.set("")

            self.label_nro_siaper = ttk.Label(frame_name, text='N° decreto SIAPER a modificar:', style="TLabel") 
            self.label_nro_siaper.grid(row=i_cargo_row+10, column=0, padx=5, pady=5, sticky="w")

            self.entrada_nro_siaper = ttk.Entry(self.contenedor_siaper, font=("", 11), textvariable=self.nro_siaper_var, width=16)
            self.entrada_nro_siaper.grid(row=0, column=0, padx=5, pady=5, sticky="w")

            self.fecha_siaper_var = tk.StringVar()
            self.fecha_siaper_var.set("")

            self.label_cal_siaper = ttk.Label(self.contenedor_siaper, text="Fecha decreto SIAPER:", style="TLabel", width=30)
            self.label_cal_siaper.grid(row=0, column=1, padx=5, pady=5, sticky="w")

            self.cal_siaper = DateEntry(self.contenedor_siaper, width=17, bootstyle="primary", firstweekday=0, dateformat='%d/%m/%Y')
            self.cal_siaper.grid(row=0, column=2, padx=5, pady=5, sticky="w")

        # Cometido
        if self.cargo.departamento == "Salud":

            # Cometidos Salud
            #self.cargo.text_cometido_salud = self.cometido_text_var.get()
            self.path_cometidos_salud = os.path.join(path_a_utilizar, "clausulas_csv", "cargos_cometidos.csv")
            self.lista_text_cometido_salud = self.crear_lista_from_csv(nombre_csv=self.path_cometidos_salud, nro_columna=3, separador="|") 

            if hasattr(self, 'contenedor_cometido'):
                self.contenedor_cometido.grid_remove()
                self.label_cometido_nro.grid_remove()

        elif self.cargo.departamento == "Municipal":

            if hasattr(self, 'contenedor_cometido'):
                self.contenedor_cometido.grid_remove()
                self.label_cometido_nro.grid_remove()

            path_cometidos_suma_alzada = os.path.join(path_a_utilizar, "clausulas_csv", "cometidos_suma_alzada.csv")
            self.lista_numero_cometido_sa = self.crear_lista_from_csv(nombre_csv=path_cometidos_suma_alzada, nro_columna=0, separador="|") 
            self.lista_text_cometido_sa = self.crear_lista_from_csv(nombre_csv=path_cometidos_suma_alzada, nro_columna=1, separador="|") 
            self.lista_nro_y_text_cometido_sa = [f"{num}. {text}" for num, text in zip(self.lista_numero_cometido_sa, self.lista_text_cometido_sa)]

            self.cometido_nro_y_text_var = tk.StringVar()
            self.cometido_nro_y_text_var.set("Seleccionar")

            self.label_cometido_nro = ttk.Label(master = frame_name, text="Cometido:", style = "TLabel")
            self.label_cometido_nro.grid(row=i_cargo_row+10, column=0, padx=5, pady=5, sticky="w", columnspan=3)

            self.contenedor_cometido = ttk.Frame(frame_name)
            self.contenedor_cometido.grid(row=i_cargo_row+10, column=1, pady=5, sticky="w")

            #self.sublabel_cometido_nro = ttk.Label(self.contenedor_cometido, text="N°:", style="TLabel")
            #self.sublabel_cometido_nro.grid(row=0, column=0, padx=5, pady=5, sticky="w")

            self.cometido_text_entry = ttkb.Combobox(self.contenedor_cometido, font=("", 11), textvariable=self.cometido_nro_y_text_var, state="readonly", values= self.lista_nro_y_text_cometido_sa, width=65)
            self.cometido_text_entry.grid(row=0, column=0, padx=5, pady=5, sticky="w")

        elif self.cargo.departamento == "Educación":
            if hasattr(self, 'contenedor_cometido'):
                self.contenedor_cometido.grid_remove()
                self.label_cometido_nro.grid_remove()
                # aqui eliminar n memo direaccion de personas y nro cett fact presup repretidos.

            self.label_cometido_nro = ttk.Label(master = frame_name, text="Cometido:", style = "TLabel")
            self.label_cometido_nro.grid(row=i_cargo_row+10, column=0, padx=5, pady=5, sticky="w", columnspan=3)
            
            self.contenedor_cometido = ttk.Frame(frame_name)
            self.contenedor_cometido.grid(row=i_cargo_row+10, column=1, pady=5, sticky="w")
                
            self.cometido_educacion_var = tk.StringVar()
            self.cometido_educacion_var.set("")

            self.cometido_entry_ed = ttk.Entry(self.contenedor_cometido, font=("", 11), textvariable=self.cometido_educacion_var, width=65) 
            self.cometido_entry_ed.grid(row=0, column=0, padx=5, pady=5, rowspan=2, sticky="w")

    def consulta_cuenta(self, frame_name, fila):
        # Cuenta: Subtitulo, item, asignación, aubasignacino, subsubasignacion.
        if self.cargo.tipo_contrato != "Programa":
            self.subtitulo_var.set("21")
            self.item_var.set("")
            self.asig_var.set("")
            self.subasig_var.set("")
            self.subsubasig_var.set("")
        else:
            pass # Se obtiene la cuenta desde el combobox, al consultar anteriormente por programa

        self.label_cuenta = ttk.Label(master=frame_name, text="Cuenta (215.XX.XX.XXX.XXX.XXX):", style = "TLabel")
        self.label_cuenta.grid(row=fila, column=0, padx=5, pady=5, sticky="w")

        self.contenedor_visa_redacta = ttk.Frame(frame_name)
        self.contenedor_visa_redacta.grid(row=fila, column=1, pady=5, padx=5, sticky="w")

        self.label_subt = ttk.Label(master=self.contenedor_visa_redacta, text="Subt.:", style = "TLabel")
        self.label_subt.grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.nro_subt = ttk.Entry(self.contenedor_visa_redacta, font=("", 11), textvariable=self.subtitulo_var, width=2) # idea: hacer lista desglosable con combobox
        self.nro_subt.grid(row=0, column=1, padx=5, pady=5, sticky="w")

        self.label_item = ttk.Label(master=self.contenedor_visa_redacta, text="Ítem:", style = "TLabel")
        self.label_item.grid(row=0, column=2, padx=5, pady=5, sticky="w")
        self.nro_item = ttk.Entry(self.contenedor_visa_redacta, font=("", 11), textvariable=self.item_var, width=2) # idea: hacer lista desglosable con combobox
        self.nro_item.grid(row=0, column=3, padx=5, pady=5, sticky="w")

        self.label_asig = ttk.Label(master=self.contenedor_visa_redacta, text="Asig.:", style = "TLabel")
        self.label_asig.grid(row=0, column=4, padx=5, pady=5, sticky="w")
        self.nro_asig = ttk.Entry(self.contenedor_visa_redacta, font=("", 11), textvariable=self.asig_var, width=3) # idea: hacer lista desglosable con combobox
        self.nro_asig.grid(row=0, column=5, padx=5, pady=5, sticky="w")

        self.label_subasig = ttk.Label(master=self.contenedor_visa_redacta, text="Subasig.:", style = "TLabel")
        self.label_subasig.grid(row=0, column=6, padx=5, pady=5, sticky="w")
        self.nro_subasig = ttk.Entry(self.contenedor_visa_redacta, font=("", 11), textvariable=self.subasig_var, width=3) # idea: hacer lista desglosable con combobox
        self.nro_subasig.grid(row=0, column=7, padx=5, pady=5, sticky="w")

        self.label_subsubasig = ttk.Label(master=self.contenedor_visa_redacta, text="Subsubasig.:", style = "TLabel")
        self.label_subsubasig.grid(row=0, column=8, padx=5, pady=5, sticky="w")
        self.nro_subsubasig = ttk.Entry(self.contenedor_visa_redacta, font=("", 11), textvariable=self.subsubasig_var, width=3) # idea: hacer lista desglosable con combobox
        self.nro_subsubasig.grid(row=0, column=9, padx=5, pady=5, sticky="w")

    def obtener_cuenta(self, event): 
        # Obtener el índice del elemento seleccionado en el Combobox
        indice = self.entrada_programa.current()
        # Verificar si se ha seleccionado un elemento
        if indice != -1:
            # Obtener el valor correspondiente en lista_cuentas
            cuenta_correspondiente = self.lista_cuentas[indice]
            print(f"Cuenta correspondiente: {cuenta_correspondiente}")
            self.cargo.cuenta_con_puntos = self.trans_cuenta(cuenta_correspondiente)
            return self.cargo.cuenta_con_puntos
        else:
            print("No se ha seleccionado ningún programa")

    def cuenta_renta_bruta_y_medio_validador(self, frame_name, indice_row):
        # Cuenta
        # Tipo renta: mensual o total
        # Renta bruta monto
        # Tope maximo anual
        # Contratación por día/mes 
        # Medio validador: Informe/Certificado

        self.consulta_cuenta(frame_name, indice_row)
        
        self.tipo_renta_var = tk.StringVar()
        self.tipo_renta_var.set("Seleccionar")

        self.renta_bruta_var = tk.StringVar()
        self.renta_bruta_var.set("")

        self.renta_uf_clp_var = tk.StringVar()
        self.renta_uf_clp_var.set("CLP")

        self.ccosto_var = tk.StringVar()
        self.ccosto_var.set("")

        #self.tope_max_anual_var = tk.StringVar()
        #self.tope_max_anual_var = tk.StringVar()

        self.contratacion_diames = tk.StringVar()
        self.contratacion_diames.set("Seleccionar")

        self.medio_validador_var = tk.StringVar()
        self.medio_validador_var.set("Seleccionar")

        self.alcalia_ejercicio_subrogancia_var = tk.StringVar()
        self.alcalia_ejercicio_subrogancia_var.set("Seleccionar")

        self.secre_muni_ejercicio_subrogancia_var = tk.StringVar()
        self.secre_muni_ejercicio_subrogancia_var.set("Seleccionar")

        # Tipo renta: mensual o total
        self.label_tipo_renta = ttk.Label(master = frame_name, text="Tipo renta:", style="TLabel")
        self.label_tipo_renta.grid(row=indice_row+1, column=0, padx=5, pady=5, sticky="w")

        self.contenedor_tipo_renta = ttk.Frame(frame_name)
        self.contenedor_tipo_renta.grid(row=indice_row+1, column=1, pady=5, sticky="w")

        self.menu_tipo_renta = ttkb.Combobox(self.contenedor_tipo_renta, font=("", 11), textvariable=self.tipo_renta_var, state="readonly", values=["Mensual", "Total"], width=12)
        self.menu_tipo_renta.grid(row=0, column=0, padx=5, pady=5, sticky="w")

        # Renta bruta
        self.label_renta_bruta = ttk.Label(master = self.contenedor_tipo_renta, text="Renta bruta:", style="TLabel", width=13)
        self.label_renta_bruta.grid(row=0, column=1, padx=5, pady=5, sticky="w")

        self.entrada_renta_bruta = ttk.Entry(self.contenedor_tipo_renta, font=("", 11), textvariable=self.renta_bruta_var, width=15)
        self.entrada_renta_bruta.grid(row=0, column=2, padx=5, pady=5, sticky="w")

        self.label_uf_clp = ttk.Label(self.contenedor_tipo_renta, text="UF/CLP:", style="TLabel")
        self.label_uf_clp.grid(row=0, column=3, padx=5, pady=5, sticky="w")

        self.entrada_uf_clp = ttkb.Combobox(self.contenedor_tipo_renta, font=("", 11), textvariable=self.renta_uf_clp_var, state="readonly", values=["UF", "CLP"], width=5)
        self.entrada_uf_clp.grid(row=0, column=4, padx=5, pady=5, sticky="w")

        # Alcalde en ejercicio 

        self.contenedor_ejecicio_subrogancia = ttk.Frame(frame_name)
        self.contenedor_ejecicio_subrogancia.grid(row=indice_row+2, column=0, columnspan=2, pady=5, sticky="w")

        self.label_alcalde_en_ejercicio = ttk.Label(self.contenedor_ejecicio_subrogancia, text="Alcaldía en ejercicio:", style="TLabel", width=30)
        self.label_alcalde_en_ejercicio.grid(row=0, column=0, padx=5, pady=5, sticky="w")

        self.menu_alcalde_en_ejercicio = ttkb.Combobox(self.contenedor_ejecicio_subrogancia, font=("", 11), textvariable=self.alcalia_ejercicio_subrogancia_var, state="readonly", values=["Camila Merino Catalán", "Rodrigo Zalaquett (S)"], width=20)
        self.menu_alcalde_en_ejercicio.grid(row=0, column=1, padx=5, pady=5, sticky="w")

        # Secretarie municipal en ejercicio 

        self.label_secre_en_ejercicio = ttk.Label(self.contenedor_ejecicio_subrogancia, text="Secretario municipal:", style="TLabel")
        self.label_secre_en_ejercicio.grid(row=0, column=2, padx=5, pady=5, sticky="w")

        self.menu_secre_en_ejercicio = ttkb.Combobox(self.contenedor_ejecicio_subrogancia, font=("", 11), textvariable=self.secre_muni_ejercicio_subrogancia_var, state="readonly", values=["En ejercicio", "Subrogancia"], width=12)
        self.menu_secre_en_ejercicio.grid(row=0, column=3, padx=5, pady=5, sticky="w")

        # Tipo de regularización

        self.contenedor_regularizacion = ttk.Frame(frame_name)
        self.contenedor_regularizacion.grid(row=indice_row+3, column=0, columnspan=2, pady=5, sticky="w")

        if current_tab_index == 0: # Regularización

            self.vigente_vencido_var = tk.StringVar()
            self.vigente_vencido_var.set("Seleccionar")

            self.list_vigente_vencido = ["Vigente", "Vencido"]
            self.label_vigente_vencido = ttk.Label(self.contenedor_regularizacion, text="Contrato:", style="TLabel", width=30)
            self.label_vigente_vencido.grid(row=0, column=0, padx=5, pady=5, sticky="w")

            self_menu_vigente_vencido = ttkb.Combobox(self.contenedor_regularizacion, font=("", 11), textvariable = self.vigente_vencido_var, state="readonly", values = self.list_vigente_vencido, width=12)
            self_menu_vigente_vencido.grid(row=0, column=1, padx=5, pady=5, sticky="w") #  agregar función para hacer aparecer combobox de motivo de regularización solo para contrato vencido
        
            self_menu_vigente_vencido.bind("<<ComboboxSelected>>", lambda event: self.crear_widget_vigente_vencido(self.contenedor_regularizacion))
        
        else:
            pass
    

        # Medio validador
        self.label_medio_validador = ttk.Label(master = frame_name, text="Medio validador:", style="TLabel")
        self.label_medio_validador.grid(row=indice_row+4, column=0, padx=5, pady=5, sticky="w")

        self.contenedor_contratacion_diames_ccosto = ttk.Frame(frame_name)
        self.contenedor_contratacion_diames_ccosto.grid(row=indice_row+4, column=1, pady=5, sticky="w")

        self.menu_medio_validador = ttkb.Combobox(self.contenedor_contratacion_diames_ccosto, font=("", 11), textvariable=self.medio_validador_var, state="readonly", values=["Informe", "Certificado"], width=12)
        self.menu_medio_validador.grid(row=0, column=0, padx=5, pady=5, sticky="w")

        # Centro de costo 
        self.label_ccosto = ttk.Label(master = self.contenedor_contratacion_diames_ccosto, text="Centro Costo:", style="TLabel")
        self.label_ccosto.grid(row=0, column=1, padx=5, pady=5, sticky="w")

        self.menu_ccosto = ttk.Entry(self.contenedor_contratacion_diames_ccosto, font=("", 11), textvariable=self.ccosto_var, width=8)
        self.menu_ccosto.grid(row=0, column=2, padx=5, pady=5, sticky="w")

        # Periodo contratación 
        self.label_contratacion_diames = ttk.Label(master = self.contenedor_contratacion_diames_ccosto, text="Contratación por:", style="TLabel")
        self.label_contratacion_diames.grid(row=0, column=3, padx=5, pady=5, sticky="w")

        self.menu_contratacion_diames = ttkb.Combobox(self.contenedor_contratacion_diames_ccosto, font=("", 11), textvariable=self.contratacion_diames, state="readonly", values=["Por un día", "Por días", "Por periodo"], width=10)
        self.menu_contratacion_diames.grid(row=0, column=4, padx=5, pady=5, sticky="w")

        self.menu_contratacion_diames.bind("<<ComboboxSelected>>", lambda event: self.fechas_contratacion(frame_name, indice_row+5))

    def crear_widget_vigente_vencido(self, frame):

        self.motivo_reg_var = tk.StringVar()
        self.motivo_reg_var.set("Seleccionar")

        self.cargo.reg_vigente_o_vencido = self.vigente_vencido_var.get()

        if self.cargo.reg_vigente_o_vencido == "Vencido":
            self.list_motivos_reg = ["Urgencia", "Recepción memo tardio", "Firma tardia de prestador"]

            self.label_motivo_reg = ttk.Label(frame, text="Motivo regularización:", style="TLabel", width=20)
            self.label_motivo_reg.grid(row=0, column=2, padx=5, pady=5, sticky="w")

            self.menu_motivo_reg = ttkb.Combobox(frame, font=("", 11), textvariable = self.motivo_reg_var, state="readonly", values = self.list_motivos_reg, width=25)
            self.menu_motivo_reg.grid(row=0, column=3, padx=5, pady=5, sticky="w")

        elif self.cargo.reg_vigente_o_vencido == "Vigente":
            if hasattr(self, 'label_motivo_reg'):
                self.label_motivo_reg.grid_remove()
                self.menu_motivo_reg.grid_remove()
            pass

    def fechas_contratacion(self, frame_name, indice_row):
        self.cargo.periodo_contratacion = self.contratacion_diames.get()

        if self.cargo.periodo_contratacion == "Por periodo":
            if hasattr(self, 'nro_dias_entry'):
                self.nro_dias_entry.grid_remove()
                self.boton_ingresar_fechas.grid_remove()
                self.widgets_frame.grid_remove()

            # Fecha inicio labores
            self.contenedor_fechas_inicio_recepcion = ttk.Frame(frame_name)
            self.contenedor_fechas_inicio_recepcion.grid(row=indice_row, column=1, pady=5, sticky="w")

            self.fecha_inicio_var = tk.StringVar()
            self.fecha_inicio_var.set("")

            self.label_fecha_inicio = ttk.Label(frame_name, text="Fecha inicio labores:", style="TLabel", width=30)
            self.label_fecha_inicio.grid(row=indice_row, column=0, padx=5, pady=5, sticky="w")

            self.cal_fecha_inicio = DateEntry(self.contenedor_fechas_inicio_recepcion, width=17, bootstyle="primary", firstweekday=0, dateformat='%d/%m/%Y')
            self.cal_fecha_inicio.grid(row=0, column=0, padx=5, pady=5, sticky="w")

            # Fecha termino labores
            self.fecha_termino_var = tk.StringVar()
            self.fecha_termino_var.set("")

            self.label_fecha_termino = ttk.Label(self.contenedor_fechas_inicio_recepcion, text="Fecha término labores:", style="TLabel", width=30)
            self.label_fecha_termino.grid(row=0, column=1, padx=5, pady=5, sticky="w")

            self.cal_fecha_termino = DateEntry(self.contenedor_fechas_inicio_recepcion, width=17, bootstyle="primary", firstweekday=0, dateformat='%d/%m/%Y')
            self.cal_fecha_termino.grid(row=0, column=2, padx=5, pady=5, sticky="w")

        elif self.cargo.periodo_contratacion == "Por un día" or self.cargo.periodo_contratacion == "Por días": # hacer todo en un frame para luego hacer destroy()
            if hasattr(self, 'fecha_inicio_var'):
                self.fecha_inicio_var.set("")
                self.fecha_termino_var.set("")
                # self.cal_fecha_inicio.grid_remove() #estan dentro de contenedor_fechas_inicio_recepcion
                # self.cal_fecha_termino.grid_remove()
                self.contenedor_fechas_inicio_recepcion.grid_remove()
            if hasattr(self, 'boton_ingresar_fechas'):
                # self.nro_dias_entry.grid_remove()
                # self.boton_ingresar_fechas.grid_remove()
                # self.widgets_frame.grid_remove()
                self.contenedor_fechas_contr.grid_remove()

            self.label_nro_dias_contr = ttk.Label(frame_name, text="N° días contrato:", style="TLabel", width=30)
            self.label_nro_dias_contr.grid(row=indice_row, column=0, padx=5, pady=5, sticky="w")

            self.contenedor_fechas_contr = ttk.Frame(frame_name)
            self.contenedor_fechas_contr.grid(row=indice_row, column=1, columnspan=4, pady=5, sticky="w")
            self.widgets_frame = tk.Frame(self.contenedor_fechas_contr)

            self.nro_fechas_var = tk.IntVar()
            if self.cargo.periodo_contratacion == "Por un día":
                self.nro_fechas_var.set(1)
                if hasattr(self, 'nro_dias_entry'):
                    self.nro_dias_entry.grid_remove()
                self.nro_dias_entry = ttk.Entry(self.contenedor_fechas_contr, font=("", 11), textvariable=self.nro_fechas_var, width=5, state="readonly")
                text_boton_un_dia = "Ingresar fecha"
                self.widgets_frame.grid(row=0, column=3, padx=5, pady=5)
            else:
                if hasattr(self, 'nro_dias_entry'):
                    self.nro_dias_entry.grid_remove()
                self.nro_dias_entry = ttk.Entry(self.contenedor_fechas_contr, font=("", 11), textvariable=self.nro_fechas_var, width=5, state="normal")
                text_boton_un_dia = "Ingresar fechas"
                self.widgets_frame.grid(row=0, column=3, padx=5, pady=5)
            # self.spinbox = tk.Spinbox(root, from_=0, to=15, textvariable=self.nro_fechas_var, state="state")
            self.nro_dias_entry.grid(row=0, column=0, padx=5, pady=5, sticky="w")

            self.boton_ingresar_fechas = ttk.Button(self.contenedor_fechas_contr, text=text_boton_un_dia, command=self.crear_cal_widgets)
            self.boton_ingresar_fechas.grid(row=0, column=1, padx=5, pady=5)

    def crear_cal_widgets(self):
        print("entró a función crear cal widgets")
            
        try:
            self.cargo.nro_fechas = self.nro_fechas_var.get()
        except ValueError:
            print("Por favor, ingrese un número válido.")
            return

        for widget in self.widgets_frame.winfo_children(): #  Antes de crear nuevos widgets, destruimos los widgets anteriores en frame_name para evitar superposiciones.
            widget.destroy()

        i_col_fechas = 0
        i_row_fechas = 0
        #fechas_dias_entries = []
        
        for i in range(self.cargo.nro_fechas): # IDEA PENDIENTE 3 DIC: CREAR UNA VARIABLE DISTINTA PARA CADA WIDGET CREADO, LUEGO GUARDAR SU ULTIMO VALOR Y LUEGO RECIEN AGREGARLO A UNA LISTA AL FINAL, PARA ORDENAR DE MENOR A MAYOR CON EL SORT.
            fila_i = i-i_row_fechas
            col_i = i_col_fechas 
            label_fechas_dias = ttk.Label(self.widgets_frame, text=f"Fecha {i + 1}:")
            label_fechas_dias.grid(row=fila_i+1, column=col_i+1, padx=5, pady=5)
                
            fechas_dias_entry = DateEntry(self.widgets_frame, dateformat='%d/%m/%Y', bootstyle="primary", firstweekday=0, width=10)
            fechas_dias_entry.grid(row=fila_i+1, column=col_i+2, padx=5, pady=5)
            print("fechas_dias_entry: ", fechas_dias_entry)
            self.cargo.fechas_dias_entries.append(fechas_dias_entry)

            # Bind la seleccion de fecha a la funcion de guardar la fecha en el listado de fechas
            ## fechas_dias_entry.bind("<<DateEntrySelected>>", lambda event, de=fechas_dias_entry: save_date(event, de))
            if i != 0 and ((i + 1) % 3) == 0: # si es multiplo de 4, entra. si son mas de 5 fechas, sumar 1 a la columna, para crear nueva columna
                i_row_fechas += 3
                i_col_fechas += 2
        

    def obtener_fechas(self):
        print("entró a función obtener_fechas")

        print("fechas_dias_entries: ", self.cargo.fechas_dias_entries)
        print("fechas_dias_contratacion pre clean: ", self.cargo.fechas_dias_contratacion)
        self.cargo.fechas_dias_contratacion = [] # lo limpiamos en caso de que usuario se haya equivocado en alguna cosa y requiera volver a apretar el boton guardar, asi no se suman todas las fechas

        for entry in self.cargo.fechas_dias_entries:
            #date = entry.get_date()
            #formatted_date = date.strftime("%d de %B de %Y")
            #self.cargo.fechas_dias_contratacion.append(formatted_date)
            formatted_date = self.obtener_fecha_from_ttkb(entry,variable_a_guardar=None)
            print("formatted_date fn obtener_fechas: ", formatted_date)
            self.cargo.fechas_dias_contratacion.append(formatted_date)
        
        fechas_str = self.cargo.fechas_dias_contratacion
        print("fechas_str: ", fechas_str) # Para verificar los resultados
        fechas_datetime = [datetime.strptime(fecha, "%d de %B de %Y") for fecha in fechas_str] # lista de objetos datetime
        fechas_datetime.sort() # ordenar fechas de menor a mayor, inplace
        print("fechas_datetime: ", fechas_datetime) # Para verificar los resultados
        
        # Convertir los objetos datetime de vuelta a cadenas de fechas si es necesario
        self.cargo.fechas_dias_contratacion = [fecha.strftime("%d de %B de %Y") for fecha in fechas_datetime]
        
        # Verificar los resultados 
        print("verificando fechas guardadas ordenadas dias contratacion: " ,self.cargo.fechas_dias_contratacion)  # Para verificar los resultados


    def especialidad_salud(self, frame_name, indice_row): # traspasar widgets aqui?
        pass

    def tipo_solicitud(self, frame_name): # Solo para frame 3: "Modificación"
    
        list_solitudes = ["Memo", "Correo electrónico"] 

        label = ttk.Label(master = frame_name, text="Solicitado vía:", style="TLabel", width=30) 
        label.grid(row=2, column=0, padx=5, pady=5, sticky="w")

        self.contenedor_solicitud_via_mod = ttk.Frame(frame_name)
        self.contenedor_solicitud_via_mod.grid(row=2, column=1, pady=5, sticky="w")

        self.entrada_SSV = tk.StringVar() # podria ser int binario tambien com IntVarblbl() 
        self.entrada_SSV.set("Seleccionar")

        self.fecha_sol_mod_var = tk.StringVar() 
        self.fecha_sol_mod_var.set("")
        
        menu = ttkb.Combobox(self.contenedor_solicitud_via_mod, font=("", 11), textvariable = self.entrada_SSV, state="readonly", values = list_solitudes, width=13)
        menu.grid(row=0, column=0, padx=5, pady=5, sticky="w") 

        ## Fecha solicitud
        self.label_fecha_sol_mod = ttk.Label(self.contenedor_solicitud_via_mod, text='Fecha solicitud:', style="TLabel", width=30)
        self.label_fecha_sol_mod.grid(row=0, column=1, padx=5, pady=5, sticky="w")

        self.cal_fecha_sol_mod = DateEntry(self.contenedor_solicitud_via_mod, width=17, bootstyle="primary", firstweekday=0, dateformat='%d/%m/%Y')
        self.cal_fecha_sol_mod.grid(row=0, column=2, padx=5, pady=5, sticky="w")

    def tipo_modificacion(self, frame_name):

        tipos_modificacion = ["Beneficio", "Renta", "Plazo", "Cometido"]

        label = ttk.Label(master = frame_name, text = "Tipo de Modificación:", style="TLabel", width=30)
        label.grid(row=3, column=0, padx=5, pady=5,  sticky="w")

        self.tipo_mod_stringvar = tk.StringVar()
        self.tipo_mod_stringvar.set("Seleccionar") 

        menu = ttkb.Combobox(frame_name, font=("", 11), textvariable = self.tipo_mod_stringvar, values = tipos_modificacion, state="readonly", style="TCombobox") 
        menu.grid(row=3, column=1, padx=5, pady=5, sticky="w")

    def tipo_aprobacion(self, frame_name):
    
        lista = ["Memo", "Correo electrónico"] 

        label = ttk.Label(master = frame_name, text="Aprobación a través de:", style="TLabel")
        label.grid(row=4, column=0, padx=5, pady=5, sticky="w")

        self.entrada_AV = tk.StringVar() # podria ser int binario tambien com IntVarblbl() 
        self.entrada_AV.set("Seleccionar")

        menu = ttkb.Combobox(master=frame_name, font=("", 11), textvariable = self.entrada_AV, state="readonly", values = lista, width=30)#, background="#00c4b4")
        menu.grid(row=4, column=1, padx=5, pady=5, sticky="w") # columnspan=3 significa que se extiende por 3 columnas

    def info_persona(self, frame_name):
        # Nombre completo, rut, genero, domicilio, comuna, mail, nacionalidad, estado civil, profesion, cargo, cometido

        # Nombre completo
        self.entrada_p_nombre = tk.StringVar()
        self.entrada_p_nombre.set("")

        self.label_nombre = ttk.Label(master = frame_name, text="Nombre completo:", style = "TLabel", width=20)
        self.label_nombre.grid(row=i_personas_row+1, column=i_personas_col, padx=5, pady=5, sticky="w")

        self.menu_nombre = ttk.Entry(frame_name, font=("", 11), textvariable=self.entrada_p_nombre, width=50)
        self.menu_nombre.grid(row=i_personas_row+1, column=i_personas_col+1, padx=5, pady=5, sticky="w")

        #Rut
        self.entrada_p_rut = tk.StringVar()
        self.entrada_p_rut.set("")

        self.entrada_p_rut_dig_ver = tk.StringVar()
        self.entrada_p_rut_dig_ver.set("")

        self.label_rut = ttk.Label(master = frame_name, text="Rut (sin puntos):", style = "TLabel", width=20)
        self.label_rut.grid(row=i_personas_row+3, column=i_personas_col, padx=5, pady=5, sticky="w")

        # Crear un contenedor para los widgets para el rut
        self.contenedor = ttk.Frame(frame_name)
        self.contenedor.grid(row=i_personas_row+3, column=i_personas_col+1, pady=5, padx=5, sticky="w")

        self.menu_rut = ttk.Entry(self.contenedor, font=("", 11), textvariable=self.entrada_p_rut, width=16)
        self.menu_rut.grid(row=0, column=0, pady=5, sticky="w")

        self.label_guion = ttk.Label(self.contenedor, text="-", font=("", 11))
        self.label_guion.grid(row=0, column=1, padx=5, pady=5, sticky="w") 
        
        self.menu_rut_dig_ver = ttk.Entry(self.contenedor, font=("", 11), textvariable=self.entrada_p_rut_dig_ver, width=3)
        self.menu_rut_dig_ver.grid(row=0, column=2, padx=5, pady=5, sticky="w")

        # Género != sexo
        self.contenedor_genero_y_maternidad =ttk.Frame(frame_name)
        self.contenedor_genero_y_maternidad.grid(row=i_personas_row+4, column=i_personas_col+1, pady=5, columnspan=1, rowspan=1 , sticky="w")

        self.entrada_p_genero = tk.StringVar()
        self.entrada_p_genero.set("Seleccionar")

        self.label_genero = ttk.Label(master = frame_name, text="Género:", style = "TLabel", width=20)
        self.label_genero.grid(row=i_personas_row+4, column=i_personas_col, padx=5, pady=5, sticky="w")

        self.menu_genero = ttkb.Combobox(self.contenedor_genero_y_maternidad, font=("", 11), textvariable=self.entrada_p_genero, values=["Femenino", "Masculino", "Otro"], state="readonly", width=13) # hacer aparecer casilla de "cual?"
        self.menu_genero.grid(row=0, column=0, padx=5, pady=5, sticky="w")

        # Clausula Protección a la maternidad: Solamente para salud y educación
        self.maternidad_var = tk.BooleanVar(value=False) # Variable para almacenar la opción seleccionada

        self.label_maternidad = ttk.Label(self.contenedor_genero_y_maternidad, text="Protección maternal:", style = "TLabel")
        self.label_maternidad.grid(row=0, column=2, pady=5, padx=5, sticky="w")

        self.checkbutton_maternidad = ttk.Checkbutton(self.contenedor_genero_y_maternidad, text="(Marque si corresponde)", variable=self.maternidad_var, onvalue=True, offvalue=False)
        self.checkbutton_maternidad.grid(row=0, column=3, pady=5, sticky="w")

        # Clausula sala cuna
        self.sala_cuna_var = tk.BooleanVar(value=False) # Variable para almacenar la opción seleccionada

        self.label_sala_cuna = ttk.Label(self.contenedor_genero_y_maternidad, text="Sala cuna/Jardín infantil:", style = "TLabel")
        self.label_sala_cuna.grid(row=0, column=4, pady=5, padx=5, sticky="w")

        self.checkbutton_sala_cuna = ttk.Checkbutton(self.contenedor_genero_y_maternidad, text=" ", variable=self.sala_cuna_var, onvalue=True, offvalue=False)
        self.checkbutton_sala_cuna.grid(row=0, column=5, pady=5, sticky="w")

        # Domicilio
        self.entrada_p_domicilio = tk.StringVar()
        self.entrada_p_domicilio.set("")

        self.entrada_p_nro_domicilio = tk.StringVar()
        self.entrada_p_nro_domicilio.set("")

        self.label_domicilio = ttk.Label(master = frame_name, text="Domicilio:", style = "TLabel", width=20)
        self.label_domicilio.grid(row=i_personas_row+5, column=i_personas_col, padx=5, pady=5, sticky="w")

        # Crear un contenedor para los widgets de domicilio
        self.contenedor_domicilio = ttk.Frame(frame_name)
        self.contenedor_domicilio.grid(row=i_personas_row+5, column=i_personas_col+1, pady=5, padx=5, sticky="w")

        self.entry_domicilio = ttk.Entry(self.contenedor_domicilio, font=("", 11), textvariable=self.entrada_p_domicilio, width=50)
        self.entry_domicilio.grid(row=0, column=0, pady=5, sticky="w")
        self.label_nro = ttk.Label(self.contenedor_domicilio, text="#", font=("", 11))
        self.label_nro.grid(row=0, column=1, padx=5, pady=5, sticky="w") 
        self.entry_nro_domicilio = ttk.Entry(self.contenedor_domicilio, font=("", 11), textvariable=self.entrada_p_nro_domicilio, width=10)
        self.entry_nro_domicilio.grid(row=0, column=2, padx=5, pady=5, sticky="w")

        # Comuna
        self.entrada_p_comuna = tk.StringVar()
        self.entrada_p_comuna.set("Seleccionar")

        file_path_comunas = os.path.join(path_a_utilizar, "clausulas_csv", "comunas_chile.csv")
        list_comunas = self.crear_lista_from_csv(file_path_comunas, nro_columna=1, separador="|")
        list_comunas.sort() # ordenamos la lista por orden alfabetico

        self.label_comuna = ttk.Label(master = frame_name, text="Comuna:", style = "TLabel", width=20)
        self.label_comuna.grid(row=i_personas_row+6, column=i_personas_col, padx=5, pady=5, sticky="w")

        # Crear un contenedor para los widgets
        self.contenedor_comuna_y_aclaracion = ttk.Frame(frame_name)
        self.contenedor_comuna_y_aclaracion.grid(row=i_personas_row+6, column=i_personas_col+1, pady=5, sticky="w")

        self.menu_comuna = ttkb.Combobox(self.contenedor_comuna_y_aclaracion, font=("", 11), width = 25, textvariable = self.entrada_p_comuna, state="normal", values = list_comunas, height=10)
        self.menu_comuna.grid(row=0, column=0, padx=5, pady=5, sticky="w")

        # Buscador
        self.menu_comuna.bind('<KeyRelease>', lambda event: self._filter_combobox(event, self.menu_comuna, list_comunas))

        # Aclaración de domicilio (e.g. depto, nro casa, block, etc)
        self.entrada_p_aclaracion_domicilio= tk.StringVar()
        self.entrada_p_aclaracion_domicilio.set("")

        self.label_aclaracion_domicilio = ttk.Label(master = self.contenedor_comuna_y_aclaracion, text="Aclaración domicilio:", style = "TLabel")
        self.label_aclaracion_domicilio.grid(row=0, column=1, padx=5, pady=5, sticky="w")

        self.entry_aclaracion_domicilio = ttk.Entry(self.contenedor_comuna_y_aclaracion, font=("",11), textvariable = self.entrada_p_aclaracion_domicilio, width=29)
        self.entry_aclaracion_domicilio.grid(row=0, column=2, padx=5, pady=5, sticky="w")


        ########
        # Domicilio notificación: particular o laboral
        self.entrada_p_domicilio_part_o_lab = tk.StringVar()
        self.entrada_p_domicilio_part_o_lab.set("Seleccionar")

        self.label_domicilio_part_o_lab = ttk.Label(master = frame_name, text="Domicilio notificación:", style = "TLabel", width=20)
        self.label_domicilio_part_o_lab.grid(row=i_personas_row+7, column=i_personas_col, padx=5, pady=5, sticky="w")

        self.menu_domicilio_part_o_lab = ttkb.Combobox(frame_name, font=("", 11), textvariable=self.entrada_p_domicilio_part_o_lab, values=["Particular", "Laboral"], state="readonly", width=16)
        self.menu_domicilio_part_o_lab.grid(row=i_personas_row+7, column=i_personas_col+1, padx=5, pady=5, sticky="w")

        ######

        # Mail
        self.entrada_p_mail = tk.StringVar()
        self.entrada_p_mail.set("")

        self.label_mail = ttk.Label(master = frame_name, text="Mail:", style = "TLabel", width=20)
        self.label_mail.grid(row=i_personas_row+8, column=i_personas_col, padx=5, pady=5, sticky="w")

        self.menu_mail = ttk.Entry(frame_name, font=("", 11), textvariable=self.entrada_p_mail, width=50)
        self.menu_mail.grid(row=i_personas_row+8, column=i_personas_col+1, padx=5, pady=5, sticky="w")

        # Nacionalidad

        # Obtener la lista de nacionalidades desde el CSV
        file_path_nacionalidades = os.path.join(path_a_utilizar, "clausulas_csv", "gentilicios_nacionalidades.csv")
        self.lista_nacionalidades = self.crear_lista_from_csv(file_path_nacionalidades, nro_columna=2, separador=",")
        # Capitalizar y ordenar en orden alfabetico cada elemento de la lista
        self.lista_nacionalidades = sorted([nacionalidad.capitalize() for nacionalidad in self.lista_nacionalidades])

        self.label_nacionalidad = ttk.Label(frame_name, text="Nacionalidad:", style="TLabel", width=20)
        self.label_nacionalidad.grid(row=i_personas_row+9, column=i_personas_col, padx=5, pady=5, sticky="w")

        self.entrada_p_nacionalidad = tk.StringVar()
        self.entrada_p_nacionalidad.set("Chilena")
        self.widget_nacionalidad = ttkb.Combobox(frame_name, textvariable=self.entrada_p_nacionalidad, values=self.lista_nacionalidades, state="normal", width=25, font=("", 11), height=10) # postcommand=lambda: variable.set(variable.get()) para que se actualice el valor del combobox, ttkb permite hacer busqueda en combobox
        self.widget_nacionalidad.grid(row=i_personas_row+9, column=i_personas_col+1, padx=5, pady=5, sticky="w")

        # Buscador
        self.widget_nacionalidad.bind('<KeyRelease>', lambda event: self._filter_combobox(event, self.widget_nacionalidad, self.lista_nacionalidades))

        # Estado civil
        lista_estados_civiles = ["Soltero/a", "Casado/a", "Divorciado/a", "Viudo/a", "Conviviente civil", "Separado/a", "Separado/a judicialmente"] # Categorias dadas por Dirección de Personas

        self.entrada_p_estado_civil = tk.StringVar()
        self.entrada_p_estado_civil.set("Seleccionar")

        self.label_estado_civil = ttk.Label(master = frame_name, text="Estado civil:", style = "TLabel", width=20)
        self.label_estado_civil.grid(row=i_personas_row+10, column=i_personas_col, padx=5, pady=5, sticky="w")

        self.menu_estado_civil = ttkb.Combobox(frame_name, font=("", 11), textvariable=self.entrada_p_estado_civil, width=25, values=lista_estados_civiles, state="readonly")
        self.menu_estado_civil.grid(row=i_personas_row+10, column=i_personas_col+1, padx=5, pady=5, sticky="w")

        # Profesión
        self.entrada_p_profesion = tk.StringVar()
        self.entrada_p_profesion.set("")

        self.label_profesion = ttk.Label(master = frame_name, text="Profesión:", style = "TLabel", width=20)
        self.label_profesion.grid(row=i_personas_row+11, column=i_personas_col, padx=5, pady=5, sticky="w")

        self.menu_profesion = ttk.Entry(frame_name, font=("", 11), textvariable=self.entrada_p_profesion, width=50)
        self.menu_profesion.grid(row=i_personas_row+11, column=i_personas_col+1, padx=5, pady=5, sticky="w")
        
    def beneficios(self, frame_name, fila):
        # Beneficios
        self.contenedor_beneficios =ttk.Frame(frame_name)
        self.contenedor_beneficios.grid(row=fila, column=i_personas_col+1, pady=5,columnspan=2, rowspan=2 , sticky="w")

        self.lista_beneficios = []
        self.opciones_beneficios = ["Feriados legales","Días administrativos", "Capacitación", "Uniforme clínico", 
                                    "Aguinaldo Fiestas Patrias", "Aguinaldo Navidad", "Tarjeta de vestuario", "Vestuario prestado", "Laptop", "Celular"]

        self.label_beneficios = ttk.Label(master = frame_name, text="Beneficios:", style = "TLabel", width=20)
        self.label_beneficios.grid(row=fila, column=i_personas_col, pady=5, padx=5, sticky="w")

        j=0 # Para controlar la fila
        for i, opcion in enumerate(self.opciones_beneficios): # i para controlar la columna
            var = tk.BooleanVar()
            if i <= 4: 
                i=i+1
            elif i > 4:
                i=i-4
                j=1
            checkbutton = ttk.Checkbutton(self.contenedor_beneficios, text=opcion, variable=var)
            checkbutton.grid(row=j, column=i, pady=5, padx=5, sticky="w")
            self.lista_beneficios.append(var)

    def obtener_fecha_from_ttkb(self, calendario, variable_a_guardar=None):
        print(f"Obteniendo fecha desde calendario: {calendario}")
        print("Entró a función obtener_fecha_from_ttkb")
        # Obtener la fecha seleccionada como cadena
        fecha_str = calendario.entry.get()
        print(f"Fecha seleccionada: {fecha_str}")
        # Convertir la cadena a un objeto datetime
        fecha_obj = datetime.strptime(fecha_str, "%d/%m/%Y")
        print(f"Fecha en objeto datetime: {fecha_obj}")
        # Formatear la fecha como se desee
        fecha_formateada = fecha_obj.strftime("%d de %B de %Y")
        print(f"Fecha formateada: {fecha_formateada}")
        if variable_a_guardar:
            variable_a_guardar.set(fecha_formateada)
            print(f"Fecha guardada en variable_a_guardar: {variable_a_guardar.get()}")
        return fecha_formateada
    

    def guardar(self):
        # Guardar el valor del widget textvariable en class.atributo
        self.cargo.direccion = self.direccionSolicitante.get()
        self.cargo.departamento = self.depto_var.get()
        self.cargo.serv_salud = self.serv_salud_var.get()
        self.cargo.tipo_contrato = self.entrada_TC.get()
        self.cargo.cert_presup = self.cert_presup_var.get()

        self.cargo.nro_memo_pers = self.nro_memo_dir_pers_var.get()

        self.cargo.fecha_memo_pers =  self.obtener_fecha_from_ttkb(self.cal_memo_pers, self.fecha_memo_pers_var)
        #self.fecha_memo_pers_var.set(self.cal_memo_pers.strftime("%d de %B de %Y")) # %B para mes en texto, en vez de %m
        #self.cargo.fecha_memo_pers = self.fecha_memo_pers_var.get() 

        self.cargo.fecha_memo_recep = self.obtener_fecha_from_ttkb(self.cal_fecha_recep, self.fecha_recepcion_var)
        #self.fecha_recepcion_var.set(self.cal_fecha_recep.strftime("%d de %B de %Y")) # %B para mes en texto, en vez de %m
        #self.cargo.fecha_memo_recep = self.fecha_recepcion_var.get()

        self.cargo.subtitulo_cuenta = self.subtitulo_var.get()
        self.cargo.item_cuenta = self.item_var.get()
        self.cargo.asig_cuenta = self.asig_var.get()
        self.cargo.subasig_cuenta = self.subasig_var.get()
        self.cargo.subsubasig_cuenta =  self.subsubasig_var.get()

        
        if self.cargo.departamento == "Municipal":
            # Cometido Suma Alzada
            self.cargo.numero_y_text_cometido_sa = self.cometido_nro_y_text_var.get()
            print(f"Cometido: {self.cargo.numero_y_text_cometido_sa}")
            partes = self.cargo.numero_y_text_cometido_sa.split(". ", 1) # hace split solamente 1 vez, es decir divide en 2 partes, la primera vez que correponde luego del numero

            if len(partes) == 2:
                self.cargo.numero_cometido_sa = partes[0]
                self.cargo.text_cometido_sa = partes[1]
            else:
                self.cargo.numero_cometido_sa = ""
                self.cargo.text_cometido_sa = ""
            
        elif self.cargo.departamento == "Salud":
            # Cometidos Salud: podria ser programa o suma alzada, pero siempre tiene especialidad/cargo
            self.cargo.especialidad_salud = self.especialidad_salud_var.get()
            print(f"Especialidad en salud en Cargo: {self.cargo.especialidad_salud}")
        else: # educacion, no tenemos los cometidos de los progrmas, y muchos contratos son de talleres
            self.cargo.text_cometido_educacion = self.cometido_educacion_var.get()
            print(f"Cometido educación en Cargo: {self.cargo.text_cometido_educacion}")

        self.cargo.tipo_renta = self.tipo_renta_var.get()
        self.cargo.renta_bruta = self.renta_bruta_var.get()

        self.cargo.periodo_contratacion = self.contratacion_diames.get()
        self.cargo.medio_validador = self.medio_validador_var.get()

        self.cargo.renta_uf_clp = self.renta_uf_clp_var.get()
        self.cargo.nro_ccosto = self.ccosto_var.get()

        if self.cargo.periodo_contratacion == "Por periodo":
            self.cargo.fecha_inicio = self.obtener_fecha_from_ttkb(self.cal_fecha_inicio, self.fecha_inicio_var)
            self.cargo.fecha_termino = self.obtener_fecha_from_ttkb(self.cal_fecha_termino, self.fecha_termino_var)
        elif self.cargo.periodo_contratacion == "Por un día" or self.cargo.periodo_contratacion == "Por días":
            self.obtener_fechas()

        if current_tab_index == 2: # modificacion
            self.cargo.tipo_solicitud = self.entrada_SSV.get()
            self.cargo.tipo_modificacion = self.tipo_mod_stringvar.get()
            self.cargo.tipo_aprobacion = self.entrada_AV.get()

            self.cargo.fecha_solicitud = self.obtener_fecha_from_ttkb(self.cal_fecha_sol_mod, self.fecha_sol_mod_var)
            self.cargo.nro_decreto_siaper = self.nro_siaper_var.get()
            self.cargo.fecha_decreto_siaper = self.obtener_fecha_from_ttkb(self.cal_siaper, self.fecha_siaper_var)

            print(f"Tipo de solicitud guardada en Cargo: {self.cargo.tipo_solicitud}")
            print(f"Tipo de modificación guardada en Cargo: {self.cargo.tipo_modificacion}")
            print(f"Tipo de aprobación guardada en Cargo: {self.cargo.tipo_aprobacion}")

        if current_tab_index == "Regularización":
            self.cargo.reg_vigente_o_vencido = self.vigente_vencido_var.get()
            self.motivo_regularizacion = self.motivo_reg_var.get()
            self.cargo.fecha_instrumento_contrato = self.obtener_fecha_from_ttkb(self.cal_instrumento_contrato, self.fecha_instrumento_contrato_var)

        else: # En fecha
            self.cargo.fecha_instrumento_contrato = self.obtener_fecha_from_ttkb(self.cal_instrumento_contrato, self.fecha_instrumento_contrato_var)

        print(f"Dirección guardada en Cargo: {self.cargo.direccion}")
        print(f"Tipo de solicitud guardada en Cargo: {self.cargo.tipo_solicitud}")
        print(f"Tipo de modificación guardada en Cargo: {self.cargo.tipo_modificacion}")
        print(f"Tipo de aprobación guardada en Cargo: {self.cargo.tipo_aprobacion}")
        print(f"Departamento guardado en Cargo: {self.cargo.departamento}")
        print(f"Servicio de salud guardado en Cargo: {self.cargo.serv_salud}")
        print(f"Tipo de contrato guardado en Cargo: {self.cargo.tipo_contrato}")
        print(f"N° certificado factibilidad presupuestaria guardado en Cargo: {self.cargo.cert_presup}")
        print(f"N° memo dirección solicitante guardado en Cargo: {self.cargo.nro_memo_dir}")
        print(f"N° memo dirección de personas guardado en Cargo: {self.cargo.nro_memo_pers}")
        print(f"Fecha memo dirección solicitante guardado en Cargo: {self.cargo.fecha_memo_dir}")
        print(f"Fecha memo dirección de personas guardado en Cargo: {self.cargo.fecha_memo_pers}")
        print(f"N° memo DEM guardado en Cargo: {self.cargo.nro_memo_dem}")
        print(f"Fecha memo DEM guardado en Cargo: {self.cargo.fecha_memo_dem}")

        if self.cargo.tipo_contrato == "Programa":
            self.cargo.programa = self.entrada_prog_var.get()
            print(f"Nombre del programa en Cargo: {self.cargo.programa}")
            self.cargo.nro_decreto_5a1 = self.nro_decreto_5a1_var.get() # decreto creacion
            print(f"Nro decreto creación en Cargo: {self.cargo.nro_decreto_5a1}")
            self.cargo.fecha_decreto_5a1 = self.obtener_fecha_from_ttkb(self.cal_5a1, self.fecha_5a1_var)
            print(f"Fecha decreto creación en Cargo: {self.cargo.fecha_decreto_5a1}")
            self.cargo.nro_memo_dir = self.nro_memo_dir_sol_var.get()
            self.cargo.fecha_memo_dir = self.obtener_fecha_from_ttkb(self.cal_memo_sol, self.fecha_memo_sol_var) 

        # Educacion: Memo DEM
        elif self.cargo.tipo_contrato == "SEP Educación" or self.cargo.tipo_contrato == "Mejoramiento a la educación" or self.cargo.tipo_contrato == "FAEP":
            self.cargo.nro_memo_dem = self.nro_memoDEM_var.get()
            self.cargo.fecha_memo_dem = self.obtener_fecha_from_ttkb(self.cal_5b, self.fecha_DEM_var)
        else: # Suma alzada
            self.cargo.nro_memo_dir = self.nro_memo_dir_sol_var.get()
            self.cargo.fecha_memo_dir = self.obtener_fecha_from_ttkb(self.cal_memo_sol, self.fecha_memo_sol_var)


        # Motivo de regularización
        if current_tab_index == 0:
            self.cargo.motivo_regularizacion = self.motivo_reg_var.get()
            print(f"Motivo de regularización en Cargo: {self.cargo.motivo_regularizacion}")
        else:
            pass
                
        # Clausula de obtencion fecha inicio y fin para dia y dias
        if self.cargo.periodo_contratacion == "Por días":
            print("fechas para la funcion obtener fechas: ", self.cargo.fechas_dias_contratacion)
            self.fecha_menor, self.fecha_mayor = self.obtener_fecha_menor_y_mayor(self.cargo.fechas_dias_contratacion)
            self.cargo.fecha_inicio = self.fecha_menor
            self.cargo.fecha_termino = self.fecha_mayor
        elif self.cargo.periodo_contratacion == "Por un día":
            print("fechas para la funcion obtener fechas: ", self.cargo.fechas_dias_contratacion)
            self.cargo.fecha_inicio = self.cargo.fechas_dias_contratacion[0] # fecha de inicio/termino es el único día
            self.cargo.fecha_termino = self.cargo.fechas_dias_contratacion[0]

        # Alcaldia y direccion: en ejercicio o subrogancia
        self.cargo.alcaldia = self.alcalia_ejercicio_subrogancia_var.get()
        self.cargo.secre_muni_ejercicio_o_subrogancia = self.secre_muni_ejercicio_subrogancia_var.get()
        print(f"Alcaldía: {self.cargo.alcaldia}")
        print(f"Dirección solicitante: {self.cargo.secre_muni_ejercicio_o_subrogancia}")

        self.cargo.visadora_1 = self.entrada_visa_1.get()
        self.cargo.visadora_2 = self.entrada_visa_2.get()
        self.cargo.redactora = self.entrada_redacta.get()

        self.guardar_persona()
    
    def guardar_persona(self):
        self.persona.nombre = self.entrada_p_nombre.get().upper() # Hace mayusculas en las iniciales title()
        # el rut lo guardamo haciendo la erificacion antes de guardar estas variables
        self.persona.genero = self.entrada_p_genero.get()
        self.persona.domicilio = self.entrada_p_domicilio.get().title()
        self.persona.nro_domicilio = self.entrada_p_nro_domicilio.get()
        self.persona.comuna = self.entrada_p_comuna.get().title()
        self.persona.aclaracion_domicilio = self.entrada_p_aclaracion_domicilio.get().title() # primera letra mayuscula
        self.persona.domicilio_part_o_lab = self.entrada_p_domicilio_part_o_lab.get().lower() # minuscula
        self.persona.mail = self.entrada_p_mail.get()
        self.persona.nacionalidad = self.entrada_p_nacionalidad.get().lower()
        print(f"Nacionalidad guardada: {self.persona.nacionalidad}")

        self.persona.estado_civil = self.entrada_p_estado_civil.get().lower() ## aqui: cambia o/a segun genero

        if self.persona.genero == "Femenino":
            self.persona.estado_civil = self.persona.estado_civil.replace("o/a", "a") # Cambia por a en estado civil
        elif self.persona.genero == "Masculino":
            self.persona.estado_civil = self.persona.estado_civil.replace("o/a", "o") # Cambia por 0 en estado civil
        elif self.persona.genero == "Otro":
            self.persona.estado_civil = self.persona.estado_civil.replace("o/a", "e") # Cambia por e en estado civil

        self.persona.profesion = self.entrada_p_profesion.get().title() # Hace mayusculas en las iniciales title()

        self.art_nombre, self.art_prestador, self.art_servidor = self.join_art_nombre(self.persona.genero)

        # Recorre la lista de beneficios y obtén el valor de cada uno
        self.persona.beneficios_seleccionados = [var.get() for var in self.lista_beneficios]
        # Recorre la lista de maternidad
        self.persona.maternidad_seleccionada = self.maternidad_var.get()
        # Sala cuna
        self.persona.sala_cuna = self.sala_cuna_var.get()

        print(f"Persona guardada: {self.persona.nombre}")
        print(f"Rut guardado: {self.persona.rut}-{self.persona.dig_ver}")
        print(f"Género guardado: {self.persona.genero}")
        print(f"Domicilio guardado: {self.persona.domicilio}")
        print(f"Comuna guardada: {self.persona.comuna}")
        print(f"Mail guardado: {self.persona.mail}")
        print(f"Nacionalidad guardada: {self.persona.nacionalidad}")
        print(f"Estado civil guardado: {self.persona.estado_civil}")
        print(f"Profesión guardada: {self.persona.profesion}")
        print(f"Beneficios seleccionados: {self.persona.beneficios_seleccionados}")
        print(f"Protección maternal seleccionada: {self.persona.maternidad_seleccionada}")

    def clausulas_beneficios(self, df_csv, nro_fila):
        # Recorre la lista de beneficios y obtén el valor de cada uno
        self.persona.beneficios_seleccionados = [var.get() for var in self.lista_beneficios]
        print("self.persona.beneficios_seleccionados: ", self.persona.beneficios_seleccionados)
        for i_count in range(len(self.persona.beneficios_seleccionados)):
            if self.persona.beneficios_seleccionados[i_count] == True:
                select = self.opciones_beneficios[i_count]
                print(f"Beneficio seleccionado: {select}")
                self.persona.nombres_beneficios_seleccionados.append(select) # agrego beneficios seleccionados a lista de beneficios (nombres, no booleano)
                # Escritura de clausulas
                df_csv["texto"] = df_csv["texto"].str.replace("[BENEFICIO]", select) # crear clausula por cada beneficio seleccionado
                df_csv.at[nro_fila, "texto"] = df_csv.at[nro_fila, "texto"].replace("[#]", str(self.count_parrafo))
                content_b = df_csv["texto"][nro_fila] # clausula modificacion (por cada beneficio)
                self.frases_lineas(documento=document, text=content_b, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False)
                print(f"Clausula de beneficio {select}: listo")
                self.count_parrafo += 1
            else:
                pass
        self.count_parrafo -= 1

    def obtencion_var_beneficios(self):
        # Recorre la lista de beneficios y obtén el valor de cada uno
        self.persona.beneficios_seleccionados = [var.get() for var in self.lista_beneficios]
        print("self.persona.beneficios_seleccionados: ", self.persona.beneficios_seleccionados)
        for i_count in range(len(self.persona.beneficios_seleccionados)):
            if self.persona.beneficios_seleccionados[i_count] == True:
                select = self.opciones_beneficios[i_count]
                print(f"Beneficio seleccionado: {select}")
                self.persona.nombres_beneficios_seleccionados.append(select) # agrego beneficios seleccionados a lista de beneficios (nombres, no booleano)
            else:
                pass

    def crear_word_decretos(self):
        ###################################
        ########## CREACIÓN WORD ##########
        ###################################
        print("Creando documento Word de decreto...")
        global document
        # Abrir el documento existente
        document = Document()

        # Acceder a la sección del documento
        seccion = document.sections[0]

        # Configurar el tamaño de la página a formato oficio (8.5 x 13 pulgadas)
        seccion.page_height = Cm(33.02)  # 13 pulgadas en cm
        seccion.page_width = Cm(21.59)   # 8.5 pulgadas en cm

        # Modificar los márgenes de la sección
        seccion.top_margin = Cm(2.75)    # Margen superior de 2.5 cm
        seccion.bottom_margin = Cm(1.25) # Margen inferior de 2.5 cm
        seccion.left_margin = Cm(3)   # Margen izquierdo de 2.5 cm
        seccion.right_margin = Cm(2.54)  # Margen derecho de 2.5 cm

        # Agregar una imagen
        image_path_logo_vert = os.path.join(path_a_utilizar, 'logos-vitacura_sineslogan_vert.png')
        # image_path_logo_vert = os.path.join(os.path.dirname(__file__), 'logos-vitacura_sineslogan_vert.png')

        parrafo_imagen = document.add_paragraph()
        logo = parrafo_imagen.add_run()
        logo.add_picture(image_path_logo_vert, width=Inches(1.3))

        # Alinear y margenes de imagen logo
        parrafo_imagen.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_logo = parrafo_imagen.paragraph_format
        format_logo.right_indent = Cm(10)  # Margen derecho de 11 cm

        # Parrafo bajo la imagen
        # PENDIENTE: witget para <<XXX>>/<<YYYY>> (/año)
        parrafo_superior_izquierdo = document.add_paragraph()

        # Agregar negrita al parrafo

        self.year_today = str(datetime.now().year)

        run = parrafo_superior_izquierdo.add_run(f"MUNICIPALIDAD DE VITACURA\nDIRECCIÓN DE PERSONAS\nDEPARTAMENTO DE PERSONAL\n<<XXX>>/{self.year_today}")
        run.font.size = Pt(10)   # Establecer el tamaño de la fuente
        run.bold = True

        # Margenes del parrafo
        parrafo_superior_izquierdo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parrafo_format_sub_logo = parrafo_superior_izquierdo.paragraph_format
        parrafo_format_sub_logo.right_indent = Cm(10)  # Margen derecho de 10 cm

        # Modificar el interlineado del párrafo
        parrafo_format_sub_logo.line_spacing = Pt(10)  # Establecer el interlineado a 1.15 puntos
        
        ######################################
        # vistos y considerando, titulos
        ######################################

        if current_tab_index == 0: # Regularizacion
            parrafo_superior_derecho_1 = document.add_paragraph(f"REGULARIZASE LA CONTRATACIÓN A HONORARIOS DE {self.persona.nombre.upper()}")
        elif current_tab_index == 1: # En fecha
            parrafo_superior_derecho_1 = document.add_paragraph(f"APRUÉBASE LA CONTRATACIÓN DE PRESTACIÓN DE SERVICIOS A HONORARIOS DE {self.persona.nombre.upper()}")
        elif current_tab_index == 2: # Modificación
            fecha_siaper = datetime.strptime(self.cargo.fecha_decreto_siaper, "%d de %B de %Y")  # Ajusta el formato según sea necesario
            parrafo_superior_derecho_1 = document.add_paragraph(f"MODIFICA CONTRATO A HONORARIOS DE {self.persona.nombre.upper()}, APROBADO POR DECRETO SIAPER N°{self.cargo.nro_decreto_siaper} DE {fecha_siaper.year}")
        else:
            pass

        parrafo_superior_derecho_1.alignment = WD_ALIGN_PARAGRAPH.LEFT.JUSTIFY
        run_s_d_1 = parrafo_superior_derecho_1.runs[0]
        run_s_d_1.bold = True          # Establecer el texto en negrita
        run_s_d_1.font.size = Pt(12)   # Establecer el tamaño de la fuente

        parrafo_superior_derecho_2 = document.add_paragraph("DECRETO SIAPER Nº_______")
        parrafo_superior_derecho_2.alignment = WD_ALIGN_PARAGRAPH.LEFT.JUSTIFY
        run_s_d_2 = parrafo_superior_derecho_2.runs[0]
        run_s_d_2.bold = True          # Establecer el texto en negrita
        run_s_d_2.font.size = Pt(12)   # Establecer el tamaño de la fuente

        parrafo_superior_derecho_3 = document.add_paragraph("VITACURA,\n")
        parrafo_superior_derecho_3.alignment = WD_ALIGN_PARAGRAPH.LEFT.JUSTIFY
        run_s_d_3 = parrafo_superior_derecho_3.runs[0]
        run_s_d_3.bold = True          # Establecer el texto en negrita
        run_s_d_3.font.size = Pt(12)   # Establecer el tamaño de la fuente

        # Ajustar los márgenes del párrafo
        parrafo_format_1 = parrafo_superior_derecho_1.paragraph_format
        parrafo_format_1.left_indent = Cm(8)  # Margen izquierdo de 9 cm
        parrafo_format_1.line_spacing = Pt(11)  # Establecer el interlineado a 1 puntos


        parrafo_format_2 = parrafo_superior_derecho_2.paragraph_format
        parrafo_format_2.left_indent = Cm(8)  # Margen izquierdo de 9 cm
        #parrafo_format_2.line_spacing = Pt(10)  # Establecer el interlineado a 1 puntos, omito por ser una linea

        parrafo_format_3 = parrafo_superior_derecho_3.paragraph_format
        parrafo_format_3.left_indent = Cm(8)  # Margen izquierdo de 9 cm
        #parrafo_format_2.line_spacing = Pt(10)  # Establecer el interlineado a 1 puntos, omito por ser una linea

        #parrafo_format = parrafo.paragraph_format
        #parrafo_format.right_indent = Cm(1)  # Margen derecho de 1 cm

        # Ajustar los espacios despues y antes del párrafo
        #parrafo_format_1.space_before = Pt(12)  # Espacio antes del párrafo
        parrafo_format_1.space_after = Pt(7)  # Espacio después del párrafo
        parrafo_format_2.space_after = Pt(7)  # Espacio después del párrafo
        parrafo_format_3.space_after = Pt(7)  # Espacio después del párrafo

        # Titulo de vistos y considerando
        parrafo_vyc = document.add_paragraph()

        # Agregar negrita al parrafo
        run = parrafo_vyc.add_run("VISTOS Y CONSIDERANDO:")
        run.bold = True
        run.font.size = Pt(12)   # Establecer el tamaño de la fuente

        parrafo_vyc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        ######################################
        # vistos y considerando, parrafo 1 y 2
        ######################################
        self.word_decretos()

        ###################################

    def crear_word_contratos(self):
        ############################################
        ########## CREACIÓN WORD CONTRATO ##########
        ############################################
        print("Creando documento Word de contrato...")
        global document_contrato
        # Abrir el documento existente
        document_contrato = Document()

        # Acceder a la sección del documento
        seccion_contratos = document_contrato.sections[0]

        # Configurar el tamaño de la página a formato oficio (8.5 x 13 pulgadas)
        seccion_contratos.page_height = Cm(33.02)  # 13 pulgadas en cm
        seccion_contratos.page_width = Cm(21.59)   # 8.5 pulgadas en cm

        # Modificar los márgenes de la sección
        seccion_contratos.top_margin = Cm(2.75)    # Margen superior de 2.5 cm
        seccion_contratos.bottom_margin = Cm(1.25) # Margen inferior de 2.5 cm
        seccion_contratos.left_margin = Cm(3)   # Margen izquierdo de 2.5 cm
        seccion_contratos.right_margin = Cm(2.54)  # Margen derecho de 2.5 cm

        # Agregar una imagen
        image_path_logo_vert = os.path.join(path_a_utilizar, 'logos-vitacura_sineslogan_vert.png')
        parrafo_imagen_contrato = document_contrato.add_paragraph()
        logo_contrato = parrafo_imagen_contrato.add_run()
        logo_contrato.add_picture(image_path_logo_vert, width=Inches(1.3))

        # Alinear y margenes de imagen logo
        parrafo_imagen_contrato.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_logo_contrato = parrafo_imagen_contrato.paragraph_format
        format_logo_contrato.right_indent = Cm(10)  # Margen derecho de 11 cm

        # Parrafo bajo la imagen
        parrafo_superior_izquierdo_contrato = document_contrato.add_paragraph()

        # Agregar negrita al parrafo

        self.year_today = str(datetime.now().year)

        run = parrafo_superior_izquierdo_contrato.add_run(f"MUNICIPALIDAD DE VITACURA\nDIRECCIÓN DE PERSONAS\nDEPARTAMENTO DE PERSONAL\n<<XXX>>/{self.year_today}")
        run.font.size = Pt(10)   # Establecer el tamaño de la fuente
        run.bold = True

        # Margenes del parrafo superior izquierdo
        parrafo_superior_izquierdo_contrato.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parrafo_format_sub_logo_contrato = parrafo_superior_izquierdo_contrato.paragraph_format
        parrafo_format_sub_logo_contrato.right_indent = Cm(10)  # Margen derecho de 10 cm

        # Modificar el interlineado del párrafo
        parrafo_format_sub_logo_contrato.line_spacing = Pt(10)  # Establecer el interlineado a 1.15 puntos

        # Iterar sobre todos los párrafos y `runs` del documento
        for parrafo in document_contrato.paragraphs:
            for run in parrafo.runs:
                run.font.name = 'Times New Roman'  # Establecer la fuente usada en los contratos
        
        ######################################
        ############## Titulos ###############
        ######################################

        # Titulo de contrato: Agregar negrita y subrayado al parrafo
        self.frases_lineas(documento=document_contrato, text="\nCONTRATO DE PRESTACIÓN DE SERVICIO A HONORARIOS", posicion=WD_ALIGN_PARAGRAPH.CENTER, fuente=12, bold_bool=True, subrayado=True)

    def fuente_en_negrita(self, parrafo_n, texto, palabra_en_negrita): # PROBAR FUNCIÓN
        # Agregar un párrafo con una palabra en negrita
        parrafo_n = document.add_paragraph()
        # HACE SPLIT CON LA PALABRA EN NEGRITA
        # USA EL VALOR 0 COMO ORACIÓN PREVIA A PALABRA EN NEGRITA
        # USA EL VALOR 1 COMO ORACIÓN POSTERIOR A PALABRA EN NEGRITA
        lista_oraciones = texto.split(palabra_en_negrita)
        texto_previo = lista_oraciones[0]
        texto_posterior = lista_oraciones[1]

        parrafo_n.add_run(texto_previo).bold = False
        run_negrita = parrafo_n.add_run(palabra_en_negrita)
        run_negrita.bold = True
        parrafo_n.add_run(texto_posterior).bold = False

    def word_decretos(self):                                                 
        # Vistos y Considerandos
        # Clausulas decretos
        ######################################
        # vistos y considerando, parrafo 1 y 2
        ######################################
        file_path_vyc_1y2 = os.path.join(path_a_utilizar, "clausulas_csv", "vyc_1+2_siempre.csv") # os.path.join se encarga de construir la ruta del archivo de manera correcta y portátil, independientemente del sistema operativo.
        vyc_1_2 = pd.read_csv(file_path_vyc_1y2, sep="|", encoding="utf-8", header=None, usecols=[0], names=["texto"]) 
        vyc_1_2["texto"] = vyc_1_2["texto"].astype("string")

        # Escribir el docx
        self.count_parrafo = 0
        for i in range(len(vyc_1_2)):
            self.count_parrafo += 1
            vyc_1_2.at[i, "texto"] = vyc_1_2.at[i, "texto"].replace("[#]", f"{self.count_parrafo}.")
            content_1_2 = vyc_1_2["texto"][i]
            self.frases_lineas(documento=document, text=content_1_2, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
        print("Párrafo 1 y 2: listo")
        
        ###################################
        # vistos y considerando, parrafo 3: REGULARIZACION O EN FECHA
        ###################################
        # Separador: "|",  a modificar:  <<serviciosalud>> 
        file_path_vyc_3_csv = os.path.join(path_a_utilizar, "clausulas_csv", "vyc_3.csv")
        vyc_3_csv = pd.read_csv(file_path_vyc_3_csv, sep="|", encoding="utf-8", header=None, usecols=[1], names=["texto"]) 
        vyc_3_csv["texto"] = vyc_3_csv["texto"].astype("string") # hacemos columna string, podria haberlo hecho en el read_csv
        
        self.count_parrafo += 1 # se hace igual a 3

        if current_tab_index == 0 or current_tab_index == 1: # Regularización o en fecha
            # escribir el docx, dependiendo si es municipal, salud o educacion
            if self.cargo.direccion == "Dirección de Salud y Educación, y Demás Incorporados en la Gestión Municipal": 
                if self.cargo.departamento == "Salud":
                    vyc_3_csv["texto"] = vyc_3_csv["texto"].str.replace("<<serviciosalud>>", self.cargo.serv_salud) # str.replace reemplaza incluso dentro de la oracion a diferencia de solamente replace que requiere que toda la celda sea igual
                    vyc_3_csv.at[2, "texto"] = vyc_3_csv.at[2, "texto"].replace("[#]", str(self.count_parrafo))
                    content_3 = vyc_3_csv["texto"][2] # fila para tipo salud
                elif self.cargo.departamento == "Educación":
                    vyc_3_csv.at[1, "texto"] = vyc_3_csv.at[1, "texto"].replace("[#]", str(self.count_parrafo))
                    content_3 = vyc_3_csv["texto"][1] # fila para tipo educacion
            else: # (depto = "Municipal")
                vyc_3_csv.at[0, "texto"] = vyc_3_csv.at[0, "texto"].replace("[#]", str(self.count_parrafo))
                content_3 = vyc_3_csv["texto"][0] # fila para tipo municipales
            self.frases_lineas(documento=document, text=content_3, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
            #parrafo_3 = document.add_paragraph(content_3)
            #parrafo_3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            print("Párrafo 3: listo")

        ###################################
        # vistos y considerando, parrafo 3: MODIFICACIÓN
        ###################################

        elif current_tab_index == 2: # Modificación
            # self.art_nombre, self.art_prestador, self.art_servidor = self.join_art_nombre(self.persona.genero)
            vyc_3_csv["texto"] = vyc_3_csv["texto"].str.replace("[ARTÍCULO + Sr./Sra.]", self.art_nombre)
            vyc_3_csv["texto"] = vyc_3_csv["texto"].str.replace("[NOMBRE PERSONA]", self.persona.nombre)
            vyc_3_csv["texto"] = vyc_3_csv["texto"].str.replace("[NOMBRE DIRECCIÓN]", self.cargo.direccion)
            vyc_3_csv["texto"] = vyc_3_csv["texto"].str.replace("[X MEMO DIR]", self.cargo.nro_memo_dir)
            vyc_3_csv["texto"] = vyc_3_csv["texto"].str.replace("[DÍA de MES de AÑO DIR]", self.cargo.fecha_memo_dir)
            vyc_3_csv["texto"] = vyc_3_csv["texto"].str.replace("[DÍA de MES de AÑO INICIO]", self.cargo.fecha_inicio)

            if self.cargo.tipo_solicitud == "Memo" and self.cargo.tipo_modificacion == "Beneficio": 
                #vyc_3_csv.at[3, "texto"] = vyc_3_csv.at[3, "texto"].replace("[#]", str(self.count_parrafo)) # idea opti: hacer funcion con esta y la siguiente linea
                #content_3 = vyc_3_csv["texto"][3] 
                self.clausulas_beneficios(df_csv=vyc_3_csv, nro_fila=3)
                # rompemos if para no usar add_paragraph nuevamente
            elif self.cargo.tipo_solicitud == "Correo electrónico" and self.cargo.tipo_modificacion == "Beneficio": 
                #vyc_3_csv.at[4, "texto"] = vyc_3_csv.at[4, "texto"].replace("[#]", str(self.count_parrafo))
                #content_3 = vyc_3_csv["texto"][4]
                self.clausulas_beneficios(df_csv=vyc_3_csv, nro_fila=4)
            elif self.cargo.tipo_modificacion == "Renta":
                vyc_3_csv.at[5, "texto"] = vyc_3_csv.at[5, "texto"].replace("[#]", str(self.count_parrafo))
                content_3 = vyc_3_csv["texto"][5]
                self.frases_lineas(documento=document, text=content_3, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
                print("Párrafo 3: listo")
            elif self.cargo.tipo_modificacion == "Plazo":
                vyc_3_csv.at[6, "texto"] = vyc_3_csv.at[6, "texto"].replace("[#]", str(self.count_parrafo))
                content_3 = vyc_3_csv["texto"][6]
                self.frases_lineas(documento=document, text=content_3, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
                print("Párrafo 3: listo")
            elif self.cargo.tipo_modificacion == "Cometido":
                vyc_3_csv.at[7, "texto"] = vyc_3_csv.at[7, "texto"].replace("[#]", str(self.count_parrafo))
                content_3 = vyc_3_csv["texto"][7]
                self.frases_lineas(documento=document, text=content_3, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
                print("Párrafo 3: listo")
        else: 
            pass

        ###################################
        # vistos y considerando, parrafo 4
        ###################################
        file_path_vyc_4_csv = os.path.join(path_a_utilizar, "clausulas_csv", "vyc_4.csv")
        vyc_4_csv = pd.read_csv(file_path_vyc_4_csv, sep="|", encoding="utf-8", header=None, usecols=[1], names=["texto"]) 
        vyc_4_csv["texto"] = vyc_4_csv["texto"].astype("string")

        self.count_parrafo += 1
        if current_tab_index != 2: # Regularización o en fecha
            vyc_4_csv.at[0, "texto"] = vyc_4_csv.at[0, "texto"].replace("[#]", str(self.count_parrafo))
            content_4 = vyc_4_csv["texto"][0]
            self.frases_lineas(documento=document, text=content_4, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
            print("Párrafo 4: listo")

        #self.art_nombre, self.art_prestador, self.art_servidor = self.join_art_nombre(self.persona.genero)

        elif current_tab_index == 2: # Modificación
            #self.art_nombre, self.art_prestador, self.art_servidor = self.join_art_nombre(self.persona.genero)
            vyc_4_csv["texto"] = vyc_4_csv["texto"].str.replace("[ARTÍCULO + Sr./Sra.]", self.art_nombre)
            vyc_4_csv["texto"] = vyc_4_csv["texto"].str.replace("[NOMBRE PERSONA]", self.persona.nombre)
            vyc_4_csv["texto"] = vyc_4_csv["texto"].str.replace("[X MEMO PERS]", self.cargo.nro_memo_pers)
            vyc_4_csv["texto"] = vyc_4_csv["texto"].str.replace("[DÍA de MES de AÑO PERS]", self.cargo.fecha_memo_pers) 

            if self.cargo.tipo_solicitud == "Correo electrónico" and self.cargo.tipo_modificacion == "Beneficio": # Beneficio via Correo
                #vyc_4_csv.at[1, "texto"] = vyc_4_csv.at[1, "texto"].replace("[#]", str(self.count_parrafo))
                #content_4 = vyc_4_csv["texto"][1]
                self.clausulas_beneficios(df_csv=vyc_4_csv, nro_fila=1)
            elif self.cargo.tipo_solicitud == "Memo" and self.cargo.tipo_modificacion == "Beneficio": # Beneficio via Memo
                #vyc_4_csv.at[2, "texto"] = vyc_4_csv.at[2, "texto"].replace("[#]", str(self.count_parrafo)) 
                #content_4 = vyc_4_csv["texto"][2]
                self.clausulas_beneficios(df_csv=vyc_4_csv, nro_fila=2)
            elif self.cargo.tipo_modificacion == "Renta" or self.cargo.tipo_modificacion == "Plazo": # Renta o plazo
                vyc_4_csv.at[3, "texto"] = vyc_4_csv.at[3, "texto"].replace("[#]", str(self.count_parrafo))
                content_4 = vyc_4_csv["texto"][3]
                self.frases_lineas(documento=document, text=content_4, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
                print("Párrafo 4: listo")
            else: 
                pass

        ###################################
        # vistos y considerando, parrafo 5: Reg o en fecha
        ###################################

        # Separador: "|",  a modificar:  <<>>
        file_path_vyc_5_csv = os.path.join(path_a_utilizar, "clausulas_csv", "vyc_5.csv")
        vyc_5_csv = pd.read_csv(file_path_vyc_5_csv, sep="|", encoding="utf-8", header=None, usecols=[1], names=["texto"]) 
        vyc_5_csv["texto"] = vyc_5_csv["texto"].astype("string") # hacemos columna string, podria haberlo hecho en el read_csv

        self.count_parrafo += 1

        if current_tab_index != 2: # Regularización o en fecha
            vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("[NOMBRE DIRECCIÓN]", self.cargo.direccion) # reemplazamos los parametros
            vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("<<nro_decreto_5a1>>", self.cargo.nro_decreto_5a1)
            vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("<<fecha_decreto_5a1_dma>>", self.cargo.fecha_decreto_5a1)
            vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("<<nro_decreto_5a2>>", self.cargo.fecha_decreto_5a2)
            vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("<<fecha_decreto_5a2_dma>>", self.cargo.fecha_decreto_5a2) # reemplazamos los parametros
            vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("[NOMBRE PROGRAMA]", self.cargo.programa) # reemplazamos los parametros
            vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("[nro_memo_dem]", self.cargo.nro_memo_dem) # reemplazamos los parametros
            vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("[fecha_memo_dem]", self.cargo.fecha_memo_dem) # reemplazamos los parametros

            if self.cargo.tipo_contrato == "Programa":
                if self.cargo.direccion == "Dirección de Salud y Educación, y Demás Incorporados en la Gestión Municipal" and self.cargo.departamento == "Salud":
                    vyc_5_csv.at[2, "texto"] = vyc_5_csv.at[2, "texto"].replace("[#]", str(self.count_parrafo))  # Numero clausula
                    content_5 = vyc_5_csv["texto"][2] 
                    print("Programa: salud")
                else: 
                    vyc_5_csv.at[0, "texto"] = vyc_5_csv.at[0, "texto"].replace("[#]", str(self.count_parrafo))  # Numero clausula
                    content_5 = vyc_5_csv["texto"][0] 
                    print("Programa")
            elif self.cargo.tipo_contrato == "Suma alzada":
                vyc_5_csv.at[3, "texto"] = vyc_5_csv.at[3, "texto"].replace("[#]", str(self.count_parrafo))  # Numero clausula
                content_5 = vyc_5_csv["texto"][3] 
                print("Suma alzada")
            else: # tipo_contrato = "FAEP", "SEP Educación", "Mejoramiento a la educación"
                vyc_5_csv.at[1, "texto"] = vyc_5_csv.at[1, "texto"].replace("[#]", str(self.count_parrafo))  # Numero clausula
                content_5 = vyc_5_csv["texto"][1] 
                print("FAEP, SEP Educación, Mejoramiento a la educación")
            self.frases_lineas(documento=document, text=content_5, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
            print("Párrafo 5: listo")

        # self.art_nombre, self.art_prestador, self.art_servidor = self.join_art_nombre(self.persona.genero)

        else: # Modificación 
            #self.art_nombre, self.art_prestador, self.art_servidor = self.join_art_nombre(self.persona.genero)
            vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("[ARTÍCULO + Sr./Sra.]", self.art_nombre)
            vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("[NOMBRE PERSONA]", self.persona.nombre)
            vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("[X FACT PRESUP]", self.cargo.cert_presup) 
            
            if self.cargo.tipo_solicitud == "Correo electrónico" and self.cargo.tipo_modificacion == "Beneficio": # Beneficio via Correo
                #vyc_5_csv.at[4, "texto"] = vyc_5_csv.at[4, "texto"].replace("[#]", str(self.count_parrafo))  # Numero clausula
                #content_5 = vyc_5_csv["texto"][4]
                self.clausulas_beneficios(df_csv=vyc_5_csv, nro_fila=4)
            elif self.cargo.tipo_solicitud == "Memo" and self.cargo.tipo_modificacion == "Beneficio": # Beneficio via Correo
                #vyc_5_csv.at[5, "texto"] = vyc_5_csv.at[5, "texto"].replace("[#]", str(self.count_parrafo))  # Numero clausula
                #content_5 = vyc_5_csv["texto"][5]
                self.clausulas_beneficios(df_csv=vyc_5_csv, nro_fila=5)
            elif self.cargo.tipo_modificacion == "Renta" or self.cargo.tipo_modificacion == "Plazo": # Renta o plazo
                vyc_5_csv.at[6, "texto"] = vyc_5_csv.at[6, "texto"].replace("[#]", str(self.count_parrafo))  # Numero clausula
                content_5 = vyc_5_csv["texto"][6]
                self.frases_lineas(documento=document, text=content_5, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
                print("Párrafo 5: listo")
            elif self.cargo.tipo_modificacion == "Cometido":
                pass

        ###################################
        # vistos y considerando, parrafo 6 (siempre menos SEP)
        ###################################
        file_path_vyc_6_csv = os.path.join(path_a_utilizar, "clausulas_csv", "vyc_6_siempre_menos_SEP.csv")
        vyc_6_csv = pd.read_csv(file_path_vyc_6_csv, sep="|", encoding="utf-8", header=None, usecols=[1], names=["texto"]) 
        vyc_6_csv["texto"] = vyc_6_csv["texto"].astype("string") # hacemos columna string, podria haberlo hecho en el read_csv

        vyc_6_csv["texto"] = vyc_6_csv["texto"].str.replace("<<direccion>>", self.cargo.direccion) 
        vyc_6_csv["texto"] = vyc_6_csv["texto"].str.replace("<<nro_memo_solicitante>>", self.cargo.nro_memo_dir) 
        vyc_6_csv["texto"] = vyc_6_csv["texto"].str.replace("<<fecha_memo_solicitante>>", self.cargo.fecha_memo_dir) 
        vyc_6_csv["texto"] = vyc_6_csv["texto"].str.replace("<<nombre_persona>>", self.persona.nombre)

        self.count_parrafo += 1
        if current_tab_index != 2: # Regularizacion o en fecha
            if self.cargo.tipo_contrato != "SEP Educación":
                vyc_6_csv.at[0, "texto"] = vyc_6_csv.at[0, "texto"].replace("[#]", str(self.count_parrafo))  # Numero clausula
                content_6 = vyc_6_csv["texto"][0]
                self.frases_lineas(documento=document, text=content_6, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
                print("Párrafo 6 (excepto sep): listo")
            else:
                print("Párrafo 6 no aplica por ser sep")
        elif current_tab_index == 2: # Modificación
            #[DÍA de MES de AÑO INICIO] Decreto Siaper N°[X SIAPER] [DÍA de MES de AÑO SIAPER] [ARTÍCULO + Sr./Sra.] [NOMBRE PERSONA]
            #self.art_nombre, self.art_prestador, self.art_servidor = self.join_art_nombre(self.persona.genero)
            vyc_6_csv["texto"] = vyc_6_csv["texto"].str.replace("[ARTÍCULO + Sr./Sra.]", self.art_nombre)
            vyc_6_csv["texto"] = vyc_6_csv["texto"].str.replace("[NOMBRE PERSONA]", self.persona.nombre)
            vyc_6_csv["texto"] = vyc_6_csv["texto"].str.replace("[X SIAPER]", self.cargo.nro_decreto_siaper) 
            vyc_6_csv["texto"] = vyc_6_csv["texto"].str.replace("[DÍA de MES de AÑO SIAPER]", self.cargo.fecha_decreto_siaper) 
            vyc_6_csv["texto"] = vyc_6_csv["texto"].str.replace("[DÍA de MES de AÑO INICIO]", self.cargo.fecha_inicio) # fecha inicio es fecha contrato, es lo mismo
            vyc_6_csv.at[1, "texto"] = vyc_6_csv.at[1, "texto"].replace("[#]", str(self.count_parrafo))  # Numero clausula
            content_6 = vyc_6_csv["texto"][1]
            self.frases_lineas(documento=document, text=content_6, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
            print("Párrafo 6 (excepto sep): listo")

        
        ###################################
        # vistos y considerando, parrafo 7: memo contratacion honorario reg o en fecha
        ###################################
        
        if current_tab_index != 2: # Regularización o en fecha
            # <<nro_memo_pers>> de fecha <<fecha_memo_pers>>
            file_path_vyc_7_csv = os.path.join(path_a_utilizar, "clausulas_csv", "vyc_7_siempre.csv")
            vyc_7_csv = pd.read_csv(file_path_vyc_7_csv, sep="|", encoding="utf-8", header=None, usecols=[1], names=["texto"])
            vyc_7_csv["texto"] = vyc_7_csv["texto"].astype("string")

            vyc_7_csv["texto"] = vyc_7_csv["texto"].str.replace("<<nro_memo_pers>>", self.cargo.nro_memo_pers)
            vyc_7_csv["texto"] = vyc_7_csv["texto"].str.replace("<<fecha_memo_pers>>", self.cargo.fecha_memo_pers)

            self.count_parrafo += 1
            vyc_7_csv.at[0, "texto"] = vyc_7_csv.at[0, "texto"].replace("[#]", str(self.count_parrafo))  # Numero clausula
            content_7 = vyc_7_csv["texto"][0]
            self.frases_lineas(documento=document, text=content_7, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
            print("Párrafo 7: listo")
        else:
            pass
        

        ###################################
        # vistos y considerando, parrafo 8: Factibilidad presup reg o en fecha
        ###################################
        if current_tab_index != 2: # Regularización o en fecha
            # <<nro_cert_presup>>
            file_path_vyc_8_csv = os.path.join(path_a_utilizar, "clausulas_csv", "vyc_8_siempre.csv")
            vyc_8_csv = pd.read_csv(file_path_vyc_8_csv, sep="|", encoding="utf-8", header=None, usecols=[1], names=["texto"])
            vyc_8_csv["texto"] = vyc_8_csv["texto"].astype("string")

            vyc_8_csv["texto"] = vyc_8_csv["texto"].str.replace("<<nro_cert_presup>>", self.cargo.cert_presup)

            self.count_parrafo += 1
            vyc_8_csv.at[0, "texto"] = vyc_8_csv.at[0, "texto"].replace("[#]", str(self.count_parrafo))  # Numero clausula
            content_8 = vyc_8_csv["texto"][0]
            self.frases_lineas(documento=document, text=content_8, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
            print("Párrafo 8: listo")
        else:
            pass

        ###################################
        # vistos y considerando, parrafo 9 con gestante ed y salud
        ###################################

        #self.count_parrafo = 8
        file_path_vyc_9_plus_csv = os.path.join(path_a_utilizar, "clausulas_csv", "vyc_9_en_adelante.csv")
        vyc_9_plus = pd.read_csv(file_path_vyc_9_plus_csv, sep="|", encoding="utf-8", header=None, usecols=[1], names=["texto"])
        vyc_9_plus["texto"] = vyc_9_plus["texto"].astype("string")

        if self.persona.maternidad_seleccionada == True and self.cargo.direccion == "Dirección de Salud y Educación, y Demás Incorporados en la Gestión Municipal":
            self.count_parrafo += 1
            vyc_9_plus.at[0, "texto"] = vyc_9_plus.at[0, "texto"].replace("[#]", str(self.count_parrafo))  # Numero clausula
            content_maternidad = vyc_9_plus["texto"][0] # Clausula maternal 
            self.frases_lineas(documento=document, text=content_maternidad, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
            print("Párrafo 9 maternal: listo")

        ###################################
        # vistos y considerando, parrafo 10
        ###################################

        if current_tab_index != 2: # art 5
            self.count_parrafo += 1
            vyc_9_plus.at[1, "texto"] = vyc_9_plus.at[1, "texto"].replace("[#]", str(self.count_parrafo))  # Convert self.count_parrafo to a string before replacing
            content_9o10 = vyc_9_plus["texto"][1]
            self.frases_lineas(documento=document, text=content_9o10, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
            print("Párrafo 9 o 10: listo")
        else:
            pass

        # para "regularización"
        # self.art_nombre, self.art_prestador, self.art_servidor = self.join_art_nombre(self.persona.genero)

        if current_tab_index == 0: # Regularización
            #self.art_nombre, self.art_prestador, self.art_servidor = self.join_art_nombre(self.persona.genero)
            vyc_9_plus["texto"] = vyc_9_plus["texto"].str.replace("[ARTÍCULO + Sr./Sra.]", self.art_nombre)
            vyc_9_plus["texto"] = vyc_9_plus["texto"].str.replace("[NOMBRE PERSONA]", self.persona.nombre)
            vyc_9_plus["texto"] = vyc_9_plus["texto"].str.replace("[NOMBRE DIRECCIÓN]", self.cargo.direccion)
            vyc_9_plus["texto"] = vyc_9_plus["texto"].str.replace("[X MEMO DIR]", self.cargo.nro_memo_dir)
            vyc_9_plus["texto"] = vyc_9_plus["texto"].str.replace("[DÍA de MES de AÑO RECEPCION]", self.cargo.fecha_memo_recep) 
            vyc_9_plus["texto"] = vyc_9_plus["texto"].str.replace("[EL PRESTADOR o LA PRESTADORA]", self.art_prestador)
            vyc_9_plus["texto"] = vyc_9_plus["texto"].str.replace("[DÍA de MES de AÑO INICIO]", self.cargo.fecha_inicio)
        
            self.count_parrafo += 1
            # 5 opciones de regularización: 4 vencidos y 1 vigente

            if self.cargo.reg_vigente_o_vencido == "Vencido":
                if self.cargo.motivo_regularizacion == "Urgencia":
                    vyc_9_plus.at[2, "texto"] = vyc_9_plus.at[2, "texto"].replace("[#]", str(self.count_parrafo))  # Numero clausula
                    content_10_reg = vyc_9_plus["texto"][2]
                elif self.cargo.motivo_regularizacion == "Recepción memo tardio":
                    vyc_9_plus.at[3, "texto"] = vyc_9_plus.at[3, "texto"].replace("[#]", str(self.count_parrafo))  # Numero clausula
                    content_10_reg = vyc_9_plus["texto"][3]
                elif self.cargo.motivo_regularizacion == "Firma tardia de prestador":
                    vyc_9_plus.at[4, "texto"] = vyc_9_plus.at[4, "texto"].replace("[#]", str(self.count_parrafo))  # Numero clausula
                    content_10_reg = vyc_9_plus["texto"][4]
                else:
                    pass
            elif self.cargo.reg_vigente_o_vencido == "Vigente":
                vyc_9_plus.at[5, "texto"] = vyc_9_plus.at[5, "texto"].replace("[#]", str(self.count_parrafo))
                content_10_reg = vyc_9_plus["texto"][5]

            self.frases_lineas(documento=document, text=content_10_reg, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
        else:
            pass

        ##############################################
        # vistos y considerando, parrafo 11/12 y 12/13
        ##############################################

        self.count_parrafo += 1
        vyc_9_plus.at[6, "texto"] = vyc_9_plus.at[6, "texto"].replace("[#]", str(self.count_parrafo))
        content_11 = vyc_9_plus["texto"][6]
        self.frases_lineas(documento=document, text=content_11, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
        
        # Clausula tranformación digital
        self.count_parrafo += 1
        vyc_9_plus.at[7, "texto"] = vyc_9_plus.at[7, "texto"].replace("[#]", str(self.count_parrafo))
        content_12 = vyc_9_plus["texto"][7]
        self.frases_lineas(documento=document, text=content_12, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))

        self.clausulas_decreto()

    def zfill_if_not_empty(self, value, length): # Función para completar con ceros los numeros de cuenta en caso de ser inexistentes i.e. ""
        return value.zfill(length) if value != "" else value

    def clausulas_decreto(self):

        # Titulo de decreto
        titulo_decreto = document.add_paragraph()

        # Agregar negrita al titulo
        run = titulo_decreto.add_run("DECRETO")
        run.bold = True
        run.font.size = Pt(12)   # Establecer el tamaño de la fuente del tutulo
        titulo_decreto.alignment = WD_ALIGN_PARAGRAPH.CENTER # centramos el titulo

        file_path_decretos = os.path.join(path_a_utilizar, "clausulas_csv", "decretos.csv")
        self.c_decreto = pd.read_csv(file_path_decretos, sep="|", encoding="utf-8", header=None, usecols=[1], names=["texto"])
        self.c_decreto["texto"] = self.c_decreto["texto"].astype("string")

        self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[ARTÍCULO + Sr./Sra.]", self.art_nombre)
        self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[NOMBRE PERSONA]", self.persona.nombre)
        self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[NOMBRE DIRECCIÓN]", self.cargo.direccion) 
        self.rut_con_puntos = int(self.persona.rut.replace(".", ""))
        self.rut_con_puntos = f"{self.rut_con_puntos:,}".replace(",", ".") # quito puntos y pongo puntos en caso de que existan o no
        self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[RUT PERSONA]", f"{self.rut_con_puntos}-{self.persona.dig_ver}")
        self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[PARTICULAR/LABORAL]", self.persona.domicilio_part_o_lab)

        # Subrogancia de secretario municipal
        if self.cargo.secre_muni_ejercicio_o_subrogancia == "Subrogancia":
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[(S)]", " (S)")
        else:
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[(S)]", "")

        # Creamos variable con número de cuenta [NRO CUENTA]
        self.cargo.subtitulo_cuenta = self.zfill_if_not_empty(self.cargo.subtitulo_cuenta, 2) # rellenamos con ceros la cadena hasta tener 2 caracteres, ejemplo: Input:"5", Output:"05"
        self.cargo.item_cuenta = self.zfill_if_not_empty(self.cargo.item_cuenta, 2)
        self.cargo.asig_cuenta = self.zfill_if_not_empty(self.cargo.asig_cuenta, 3) # rellenamos con ceros la cadena hasta tener 3 caracteres, ejemplo: Input:"5", Output:"005"
        self.cargo.subasig_cuenta = self.zfill_if_not_empty(self.cargo.subasig_cuenta, 3)
        self.cargo.subsubasig_cuenta = self.zfill_if_not_empty(self.cargo.subsubasig_cuenta, 3)

        # Creamos cuenta "215.XX.XX.XXX.XXX.XXX"
        # Se usa una lista por comprensión para filtrar las variables que no son cadenas vacías (""). Luego, se unen las variables con un punto "."
        variables = ["215", self.cargo.subtitulo_cuenta, self.cargo.item_cuenta, self.cargo.asig_cuenta, self.cargo.subasig_cuenta, self.cargo.subsubasig_cuenta]
        self.cargo.nro_cuenta = ".".join([var for var in variables if var != ""])

        self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[NRO CUENTA]", self.cargo.nro_cuenta) # mencionar cuenta de diferente forma para el caso de educacion, pero tb para otros programas?

        if self.cargo.nro_ccosto:
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[, Centro de Costos NRO CCOSTO]", f", Centro de Costos {self.cargo.nro_ccosto}") # si no se escribe centro de costo, no se mencionará en la clausula
        else:
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[, Centro de Costos NRO CCOSTO]", f"") 

        self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[CORREO ELECTRONICO]", self.persona.mail)
        self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[NRO COMETIDO]", self.cargo.numero_cometido_sa)
        self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[NOMBRE PROGRAMA]", self.cargo.programa) 
        # self.cargo.tipo_renta, self.cargo.renta_bruta, self.cargo.periodo_contratacion, self.cargo.medio_validador 
        self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[Total/Mensual]", self.cargo.tipo_renta)
        self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[ARTÍCULO + servidor/servidora]", self.art_servidor)

        # Fecha de instrumento contratacion o en fecha
        if self.cargo.tipo_decretoycontrato == "Regularización":
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[(de fecha DÍA de MES de AÑO)/(de igual fecha)]", f"de fecha {self.cargo.fecha_instrumento_contrato}") # 
        elif self.cargo.tipo_decretoycontrato == "En fecha":
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[(de fecha DÍA de MES de AÑO)/(de igual fecha)]", "de igual fecha") 
        
        # Impuesto Renta BBHH
        file_path_imp_renta_bbhh = os.path.join(path_a_utilizar, "clausulas_csv", "imp_bbhh.csv")
        imp_csv = pd.read_csv(file_path_imp_renta_bbhh, sep="|", encoding="utf-8", header=0 , usecols=[0,1], names=["año","impuesto"]) 
        imp_csv["año"] = imp_csv["año"].astype("int") # hacemos columna float
        imp_csv["impuesto"] = imp_csv["impuesto"].astype("float") # hacemos columna float

        # Filtrar el DataFrame por el año actual

        # Suponiendo que self.cargo.fecha_inicio es una cadena en formato "YYYY-MM-DD"
        fecha_inicio =  datetime.strptime(self.cargo.fecha_inicio, "%d de %B de %Y") #self.obtener_fecha_from_ttkb(self.cal_fecha_inicio, self.fecha_inicio_var) # cadena a datetime
        year_inicio = fecha_inicio.year # obtengo año de fecha de inicio
        self.imp = imp_csv.loc[imp_csv["año"] == int(year_inicio), "impuesto"].values[0] # obtengo el impuesto correspondiente al año en que comenzó el contrato
        self.renta_bruta_con_puntos = int(self.cargo.renta_bruta.replace(".", ""))
        self.RB = f"{self.renta_bruta_con_puntos:,}".replace(",", ".") # quito puntos y pongo puntos en caso de que existan o no
        self.cargo.renta_liquida = round(float(self.cargo.renta_bruta.replace(".", "")) * (100-self.imp)/100) # quito puntos en caso de que existan
        self.cargo.renta_liquida = f"{self.cargo.renta_liquida:,}".replace(",", ".") # string y puntos de peso chileno
        if self.cargo.renta_uf_clp == "CLP":
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[RB IMP RL]", f"de ${self.RB}.- con la deducción del {self.imp}% por concepto de Impuesto a la Renta, lo que da un líquido a pagar de ${self.cargo.renta_liquida}.-") 
        elif self.cargo.renta_uf_clp == "UF":
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[RB IMP RL]", f"de UF {self.RB}.- (impuesto incluido)") 

        # Domicilio
        if self.persona.aclaracion_domicilio == "": # no existe aclaración de domicilio
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[DOMICILIO]", f"{self.persona.domicilio} #{self.persona.nro_domicilio}, comuna de {self.persona.comuna}") 
        else: # existe aclaración de domicilio
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[DOMICILIO]", f"{self.persona.domicilio} #{self.persona.nro_domicilio}, {self.persona.aclaracion_domicilio}, comuna de {self.persona.comuna}")

        # Regularización, en fecha o modificación
        if current_tab_index == 0: # Regularización
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[APRUÉBESE/REGULARIZESE]", "REGULARIZESE")
        elif current_tab_index == 1: # En fecha
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[APRUÉBESE/REGULARIZESE]", "APRUÉBESE") 
        else: # Modificación
            pass 
        
        # Clausula de comienzo de funciones
        if self.cargo.periodo_contratacion == "Por periodo":
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[DÍA de MES de AÑO INICIO]", self.cargo.fecha_inicio)
            fecha_inicio_date = datetime.strptime(self.cargo.fecha_inicio, "%d de %B de %Y") # string a objeto datetime
            fecha_termino_date = datetime.strptime(self.cargo.fecha_termino, "%d de %B de %Y") 
            self.dias_periodo = fecha_termino_date - fecha_inicio_date
            self.dias_periodo = self.dias_periodo.days
        elif self.cargo.periodo_contratacion == "Por días":
            print(f"min: {self.fecha_menor}")
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[DÍA de MES de AÑO INICIO]", self.fecha_menor)
            self.dias_periodo = len(self.cargo.fechas_dias_contratacion)
        elif self.cargo.periodo_contratacion == "Por un día":
            print("fechas para la funcion obtener fechas: ", self.cargo.fechas_dias_contratacion)
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[DÍA de MES de AÑO INICIO]", self.cargo.fechas_dias_contratacion[0])
            self.dias_periodo = 1
        
        # Fecha de inicio y temrino para diferentes periodos de contratación, ya sea Informe o Certificación
        print("fecha i:", self.cargo.fecha_inicio)
        print("fecha t:", self.cargo.fecha_termino)
        fecha_inicio_date = datetime.strptime(self.cargo.fecha_inicio, "%d de %B de %Y") # string a objeto datetime
        fecha_termino_date = datetime.strptime(self.cargo.fecha_termino, "%d de %B de %Y")

        if self.dias_periodo >= 60: # periodo mayor a 2 meses: mensual y final
            self.cargo.mensual_final = "Mensual"
        elif self.dias_periodo < 60: # periodo menor a 2 meses: final
            self.cargo.mensual_final = "Final"
        
        # A quien se presenta informe
        if self.cargo.medio_validador == "Informe":
            if self.cargo.direccion == "Administración Municipal":
                self.directivo_a_quien_se_informa = "Administrador Municipal"
                self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[DIRECTIVO A QUIEN SE INFORMA]", self.directivo_a_quien_se_informa)
            else:
                self.directivo_a_quien_se_informa = f"el/la Director(a) de {self.cargo.direccion}"
                self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[DIRECTIVO A QUIEN SE INFORMA]", self.directivo_a_quien_se_informa)
        elif self.cargo.departamento == "Educación":
            self.directivo_a_quien_se_informa = "Jefa de Departamento de Educación Municipal y por la Directora de Servicios de Salud, Educación y demás Incorporados a la Gestión Municipal"
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[DIRECTIVO A QUIEN SE INFORMA]", self.directivo_a_quien_se_informa)
        else:
            pass

        # Comenzamos a escribir
        self.count_parrafo = 0 # reiniciamos contador
        self.count_parrafo += 1
        if self.cargo.tipo_contrato == "Suma alzada": # no hay para salud, aquí tb hay casos de educación
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[Honorarios Suma Alzada/Programa]", "Honorarios Suma Alzada")
            if  self.cargo.departamento == "Salud": 
                self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[COSAM/CESFAM/Departamento de Salud Municipal/SAPU]", "Departamento de Salud Municipal de Vitacura")
                self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[ESPECIALIDAD SALUD]", self.cargo.especialidad_salud)
                self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[, en el NOMBRE PROGRAMA]", "")
                self.c_decreto.at[3, "texto"] = self.c_decreto.at[3, "texto"].replace("[#]", f"{self.count_parrafo}°")
                content_decreto = self.c_decreto["texto"][3]
            else:
                self.c_decreto.at[0, "texto"] = self.c_decreto.at[0, "texto"].replace("[#]", f"{self.count_parrafo}°")
                content_decreto = self.c_decreto["texto"][0]
        elif self.cargo.tipo_contrato == "Programa": # hay programas para municipal y salud
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[Honorarios Suma Alzada/Programa]", f"{self.cargo.programa}") 
            if self.cargo.departamento  == "Municipal":
                self.c_decreto.at[1, "texto"] = self.c_decreto.at[1, "texto"].replace("[#]", f"{self.count_parrafo}°")
                content_decreto = self.c_decreto["texto"][1]
            elif self.cargo.departamento == "Salud": 
                self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[COSAM/CESFAM/Departamento de Salud Municipal/SAPU]", "Departamento de Salud Municipal de Vitacura")
                self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[ESPECIALIDAD SALUD]", self.cargo.especialidad_salud)
                self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[, en el NOMBRE PROGRAMA]", f", en el {self.cargo.programa}")
                self.c_decreto.at[3, "texto"] = self.c_decreto.at[3, "texto"].replace("[#]", f"{self.count_parrafo}°")
                content_decreto = self.c_decreto["texto"][3]
        else: # SEP, FAEP, Mejoramiento a a la educación
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[Honorarios Suma Alzada/Programa]", f"{self.cargo.tipo_contrato}") 
            self.c_decreto.at[2, "texto"] = self.c_decreto.at[2, "texto"].replace("[#]", f"{self.count_parrafo}°")
            content_decreto = self.c_decreto["texto"][2]
        
        self.frases_lineas(documento=document, text=content_decreto, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
        #self.escribir_parrafos_con_vinetas(list_p_cometidos = self.cargo.text_cometido_salud, filas_parrafos_en_vineta=[1,2,3,4,5,6,7,8,9], doc=document, margen_izquierdo=Cm(1.75), margen_derecho=Cm(0), sangria_primera_linea=Cm(0.5)) # text cometido
        self.obtener_y_escribir_cometidos(docu=document)

        # Medio validador: Informe o certificado
        self.count_parrafo += 1
        if self.cargo.medio_validador == "Informe":
            texto_separador = "Esta suma será pagada"
            self.c_decreto.at[4, "texto"] = self.c_decreto.at[4, "texto"].replace("[#]", f"{self.count_parrafo}°") 
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[Final/Mensual]", self.cargo.mensual_final) # Puede ser "Final" o "Mensual"
            if self.cargo.mensual_final == "Final":
                self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[un Informe Final/Informes Mensuales]", "un Informe Final")
            elif self.cargo.mensual_final == "Mensual":
                self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[un Informe Final/Informes Mensuales]", "Informes Mensuales")
            content_medio_validador_decreto = self.c_decreto["texto"][4]

        elif self.cargo.medio_validador == "Certificado":
            texto_separador = "Estas sumas serán pagadas"
            self.c_decreto.at[5, "texto"] = self.c_decreto.at[5, "texto"].replace("[#]", f"{self.count_parrafo}°") 
            if self.cargo.mensual_final == "Final":
                self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[Certificaciones de Cumplimientos Mensuales/certificación de trabajo efectivamente realizado]", "certificación de trabajo efectivamente realizado")
            elif self.cargo.mensual_final == "Mensual":
                self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[Certificaciones de Cumplimientos Mensuales/certificación de trabajo efectivamente realizado]", "Certificaciones de Cumplimientos Mensuales")
            content_medio_validador_decreto = self.c_decreto["texto"][5]
                
        dos_parrafos = content_medio_validador_decreto.split(texto_separador) # 
        content_medio_validador_decreto_1 = dos_parrafos[0] # si entra a educacion, se modifica variable
        content_medio_validador_decreto_2 = f"{texto_separador} {dos_parrafos[1]}"
            
        self.frases_lineas(documento=document, text=content_medio_validador_decreto_1, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
        self.frases_lineas(documento=document, text=content_medio_validador_decreto_2, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1))
        
        # Clausula ley 21.133 BBHH siempre
        self.count_parrafo += 1
        self.c_decreto.at[7, "texto"] = self.c_decreto.at[7, "texto"].replace("[#]", f"{self.count_parrafo}°")
        self.frases_lineas(documento=document, text=self.c_decreto["texto"][7], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))

        # Periodo contratacion
        self.count_parrafo += 1
        if self.cargo.periodo_contratacion == "Por periodo":
            print("Estoy escribiendo fecha e inicio y término, para periodo de contratación por periodo")
            print(f"Fecha inicio: {self.cargo.fecha_inicio}, Fecha término: {self.cargo.fecha_termino}")
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[DÍA de MES de AÑO INICIO]", self.cargo.fecha_inicio)
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[DÍA de MES de AÑO TÉRMINO]", self.cargo.fecha_termino)
            self.c_decreto.at[9, "texto"] = self.c_decreto.at[9, "texto"].replace("[#]", f"{self.count_parrafo}°")
            content_periodo_decreto = self.c_decreto["texto"][9]

        elif self.cargo.periodo_contratacion == "Por días":
            print("Estoy escribiendo fecha e inicio y término, para periodo de contratación por días")
            #self.fecha_menor = self.obtener_fecha_menor_y_mayor(self.cargo.fechas_dias_contratacion)
            #self.cargo.fecha_inicio = self.fecha_menor
            #print(f"min: {self.fecha_menor}")
            #self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[DÍA de MES de AÑO INICIO]", self.fecha_menor)
            print(f"Largo: {len(self.cargo.fechas_dias_contratacion)}, Lista: {self.cargo.fechas_dias_contratacion}")
            self.fechas_unidas = " , ".join(self.cargo.fechas_dias_contratacion[:-1]) + " y " + self.cargo.fechas_dias_contratacion[-1]
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[(el día (DÍA de MES de AÑO)/(los días (DÍA de MES de AÑO) y (DÍA de MES de AÑO))]", f"los días {self.fechas_unidas}") # hacer funcion para varios dias
            self.c_decreto.at[8, "texto"] = self.c_decreto.at[8, "texto"].replace("[#]", f"{self.count_parrafo}°")
            content_periodo_decreto = self.c_decreto["texto"][8]

        elif self.cargo.periodo_contratacion == "Por un día":
            print("Estoy escribiendo fecha e inicio y término, para periodo de contratación por un día")
            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[(el día (DÍA de MES de AÑO)/(los días (DÍA de MES de AÑO) y (DÍA de MES de AÑO))]", f"el día {self.cargo.fechas_dias_contratacion[0]}")
            self.c_decreto.at[8, "texto"] = self.c_decreto.at[8, "texto"].replace("[#]", f"{self.count_parrafo}°") 
            content_periodo_decreto = self.c_decreto["texto"][8]

        self.frases_lineas(documento=document, text=content_periodo_decreto, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))

        # Clausula presentacion informe
        
        if self.cargo.medio_validador == "Informe":

            self.mes_inicio = datetime.strftime(fecha_inicio_date, "%B") # cambiamos el formato del objeto datetime a una cadena de texto, escribiendo solamente el mes
            calculo_mes_penultimo = fecha_termino_date - relativedelta(months=1)
            self.mes_penultimo = datetime.strftime(calculo_mes_penultimo, "%B") 
            self.mes_termino_de_anio = datetime.strftime(fecha_termino_date, "%B de %Y") # datetime a string, escribiendo mes y año

            self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[MES TERMINO de AÑO]", self.mes_termino_de_anio)  

            self.count_parrafo += 1
            if self.cargo.mensual_final == "Mensual":
                self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[MES INICIO]", self.mes_inicio)  
                self.c_decreto["texto"] = self.c_decreto["texto"].str.replace("[MES PENÚLTIMO]", self.mes_penultimo) 
                self.c_decreto.at[10, "texto"] = self.c_decreto.at[10, "texto"].replace("[#]", f"{self.count_parrafo}°") 
                content_clausula_presentacion_informe = self.c_decreto["texto"][10]
            elif self.cargo.mensual_final == "Final": # Tipo informe: Final
                self.c_decreto.at[11, "texto"] = self.c_decreto.at[11, "texto"].replace("[#]", f"{self.count_parrafo}°")
                content_clausula_presentacion_informe = self.c_decreto["texto"][11] # informe por dia (final) o periodo menor a 2 meses 
            self.frases_lineas(documento=document, text=content_clausula_presentacion_informe, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
        else: # Certificado
            pass

        # Pendiente: Opcion tallerista, y parrafos de beneficios

        ## Creamos el word de contratos aquí para después poder ingresar los beneficios en conjunto con la escritura de los decretos

        self.crear_word_contratos()
        self.word_contratos()

        ## Clausula 1-5 para contrato
        self.index_nro_subt_clausulas = 0
        self.index_csv_contrato = 1 # parte siendo 1 ("Primero:""), luego va cambiando dentro de la funcion
        self.clausulas_contrato(nro_o_tipo_clausula="1")
        self.clausulas_contrato(nro_o_tipo_clausula="2")
        self.clausulas_contrato(nro_o_tipo_clausula="3")
        self.clausulas_contrato(nro_o_tipo_clausula="4")
        self.clausulas_contrato(nro_o_tipo_clausula="5")

        if self.persona.maternidad_seleccionada == True:
            self.count_parrafo += 1
            self.index_nro_subt_clausulas += 1
            nro_fila = 12
            self.index_csv_contrato = 11 # clausula para contratos
            self.c_decreto.at[nro_fila, "texto"] = self.c_decreto.at[nro_fila, "texto"].replace("[#]", f"{self.count_parrafo}°")
            self.frases_lineas(documento=document, text=self.c_decreto["texto"][nro_fila], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
            self.clausula_beneficio_contrato = self.csv_contratos["texto"][self.index_csv_contrato] # escribo la clausula de bemeficio maternidad
            self.clausulas_contrato(nro_o_tipo_clausula="6")

        #############################################
        ########### PARRAFOS BENEFICIOS #############
        #############################################

        # Escritura de clausulas de beneficios
        
        # Obtención nombres de beneficios seleccionados
        self.obtencion_var_beneficios()
        print("Nombres beneficios seleccionados: ", self.persona.nombres_beneficios_seleccionados)

        # Reemplazo de texto en clausulas de beneficios y escritura de clausulas
        # Feriados legales y administrativos
        if "Feriados legales" in self.persona.nombres_beneficios_seleccionados and "Días administrativos" not in self.persona.nombres_beneficios_seleccionados: # Solo feriados legales
            self.count_parrafo += 1
            self.index_nro_subt_clausulas += 1
            nro_fila = 16 # Feriados legales para decretos
            self.index_csv_contrato = 15 # clausula para contratos
            for csv in [self.c_decreto, self.csv_contratos]: # reemplazo en ambos csv
                csv["texto"] = csv["texto"].str.replace("[, Feriado Legal y Permisos por Matrimonio, Nacimiento de hijos u otros, que establezca la Ley 18.883 en las mismas condiciones que se otorgan al Personal Municipal]", "Feriado Legal y Permisos por Matrimonio, Nacimiento de hijos u otros, que establezca la Ley 18.883 en las mismas condiciones que se otorgan al Personal Municipal")
                csv["texto"] = csv["texto"].str.replace("[06 días administrativos]", "") 

            self.c_decreto.at[nro_fila, "texto"] = self.c_decreto.at[nro_fila, "texto"].replace("[#]", f"{self.count_parrafo}°")
            self.frases_lineas(documento=document, text=self.c_decreto["texto"][nro_fila], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))

            self.clausula_beneficio_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
            self.clausulas_contrato(nro_o_tipo_clausula="6") # escribe claustula beneficio de contrato
        elif "Feriados legales" not in self.persona.nombres_beneficios_seleccionados and "Días administrativos" in self.persona.nombres_beneficios_seleccionados: # Solo días administrativos
            if self.cargo.departamento == "Municipal":
                self.count_parrafo += 1
                self.index_nro_subt_clausulas += 1
                nro_fila = 16
                self.index_csv_contrato = 15
                for csv in [self.c_decreto, self.csv_contratos]:
                    csv["texto"] = csv["texto"].str.replace("[, Feriado Legal y Permisos por Matrimonio, Nacimiento de hijos u otros, que establezca la Ley 18.883 en las mismas condiciones que se otorgan al Personal Municipal]", "")
                    csv["texto"] = csv["texto"].str.replace("[06 días administrativos]", "06 días administrativos")  
                
                self.c_decreto.at[nro_fila, "texto"] = self.c_decreto.at[nro_fila, "texto"].replace("[#]", f"{self.count_parrafo}°")
                self.frases_lineas(documento=document, text=self.c_decreto["texto"][nro_fila], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))

                self.clausula_beneficio_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
                self.clausulas_contrato(nro_o_tipo_clausula="6") # escribe claustula beneficio de contrato

            elif self.cargo.departamento == "Salud":  # Salud tiene administrativos diferentes
                self.count_parrafo += 1
                self.index_nro_subt_clausulas += 1
                nro_fila = 17
                self.index_csv_contrato = 16
                self.c_decreto.at[nro_fila, "texto"] = self.c_decreto.at[nro_fila, "texto"].replace("[#]", f"{self.count_parrafo}°")
                self.frases_lineas(documento=document, text=self.c_decreto["texto"][nro_fila], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))

                self.clausula_beneficio_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
                self.clausulas_contrato(nro_o_tipo_clausula="6") # escribe claustula beneficio de contrato

            else: # Educacion no tiene administrativos
                pass
        elif "Feriados legales" in self.persona.nombres_beneficios_seleccionados and "Días administrativos" in self.persona.nombres_beneficios_seleccionados: # Ambos
            self.count_parrafo += 1
            self.index_nro_subt_clausulas += 1
            nro_fila = 16
            self.index_csv_contrato = 15
            for csv in [self.c_decreto, self.csv_contratos]:
                csv["texto"] = csv["texto"].str.replace("[, Feriado Legal y Permisos por Matrimonio, Nacimiento de hijos u otros, que establezca la Ley 18.883 en las mismas condiciones que se otorgan al Personal Municipal]", ", Feriado Legal y Permisos por Matrimonio, Nacimiento de hijos u otros, que establezca la Ley 18.883 en las mismas condiciones que se otorgan al Personal Municipal")
                csv["texto"] = csv["texto"].str.replace("[06 días administrativos]", "06 días administrativos")

            self.c_decreto.at[nro_fila, "texto"] = self.c_decreto.at[nro_fila, "texto"].replace("[#]", f"{self.count_parrafo}°")
            self.frases_lineas(documento=document, text=self.c_decreto["texto"][nro_fila], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))

            self.clausula_beneficio_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
            self.clausulas_contrato(nro_o_tipo_clausula="6") # escribe claustula beneficio de contrato
        
        # Capacitaciones
        if "Capacitación" in self.persona.nombres_beneficios_seleccionados:
            print("Estoy escribiendo capacitacion")
            self.count_parrafo += 1
            self.index_nro_subt_clausulas += 1
            nro_fila = 15
            self.index_csv_contrato = 14
            self.c_decreto.at[nro_fila, "texto"] = self.c_decreto.at[nro_fila, "texto"].replace("[#]", f"{self.count_parrafo}°")
            self.frases_lineas(documento=document, text=self.c_decreto["texto"][nro_fila], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))

            self.clausula_beneficio_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
            self.clausulas_contrato(nro_o_tipo_clausula="6") # escribe claustula beneficio de contrato
        
        # Uniforme clinico
        if "Uniforme clínico" in self.persona.nombres_beneficios_seleccionados:
            print("Estoy escribiendo uniforme clinico")
            self.count_parrafo += 1
            self.index_nro_subt_clausulas += 1
            nro_fila = 20
            self.index_csv_contrato = 19
            self.c_decreto.at[nro_fila, "texto"] = self.c_decreto.at[nro_fila, "texto"].replace("[#]", f"{self.count_parrafo}°")
            self.frases_lineas(documento=document, text=self.c_decreto["texto"][nro_fila], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))

            self.clausula_beneficio_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
            self.clausulas_contrato(nro_o_tipo_clausula="6") # escribe claustula beneficio de contrato

        # Aguinaldos
        if ("Aguinaldo Fiestas Patrias" in self.persona.nombres_beneficios_seleccionados) and ("Aguinaldo Navidad" not in self.persona.nombres_beneficios_seleccionados): # Solo fiestas patrias
            print("Estoy escribiendo aguinaldo decreto: Solo fiestas patrias")
            self.count_parrafo += 1
            self.index_nro_subt_clausulas += 1
            nro_fila = 18
            self.index_csv_contrato = 17
            for csv in [self.c_decreto, self.csv_contratos]: # Reemplazo variables tanto para csv de decretos como para csv de contratos
                csv["texto"] = csv["texto"].str.replace("[Fiestas Patrias]", "Fiestas Patrias") 
                csv["texto"] = csv["texto"].str.replace("[Navidad]", "") 
                csv["texto"] = csv["texto"].str.replace("[y_1]", "")

            # Escribo en decreto
            print("Estoy escribiendo aguinaldo decreto")
            self.c_decreto.at[nro_fila, "texto"] = self.c_decreto.at[nro_fila, "texto"].replace("[#]", f"{self.count_parrafo}°")
            self.frases_lineas(documento=document, text=self.c_decreto["texto"][nro_fila], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
            # Escrito en contrato
            print("Estoy escribiendo aguinaldo contrato")
            self.clausula_beneficio_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
            self.clausulas_contrato(nro_o_tipo_clausula="6") # escribe claustula beneficio de contrato
            
        elif ("Aguinaldo Fiestas Patrias" not in self.persona.nombres_beneficios_seleccionados) and ("Aguinaldo Navidad" in self.persona.nombres_beneficios_seleccionados): # Solo navidad
            print("Estoy escribiendo aguinaldo decreto: Solo navidad")
            self.count_parrafo += 1
            self.index_nro_subt_clausulas += 1
            nro_fila = 18
            self.index_csv_contrato = 17
            for csv in [self.c_decreto, self.csv_contratos]:
                csv["texto"] = csv["texto"].str.replace("[Fiestas Patrias]", "") 
                csv["texto"] = csv["texto"].str.replace("[Navidad]", "Navidad") 
                csv["texto"] = csv["texto"].str.replace("[y_1]", "")

            self.c_decreto.at[nro_fila, "texto"] = self.c_decreto.at[nro_fila, "texto"].replace("[#]", f"{self.count_parrafo}°")
            self.frases_lineas(documento=document, text=self.c_decreto["texto"][nro_fila], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))

            self.clausula_beneficio_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
            self.clausulas_contrato(nro_o_tipo_clausula="6") # escribe claustula beneficio de contrato
            
        elif ("Aguinaldo Fiestas Patrias" in self.persona.nombres_beneficios_seleccionados) and ("Aguinaldo Navidad" in self.persona.nombres_beneficios_seleccionados): # Ambos
            print("Estoy escribiendo aguinaldo decreto: Ambos")
            self.count_parrafo += 1
            self.index_nro_subt_clausulas += 1
            nro_fila = 18
            self.index_csv_contrato = 17
            for csv in [self.c_decreto, self.csv_contratos]:
                csv["texto"] = csv["texto"].str.replace("[Fiestas Patrias]", "Fiestas Patrias")
                csv["texto"] = csv["texto"].str.replace("[Navidad]", "Navidad")
                csv["texto"] = csv["texto"].str.replace("[y_1]", " y ")
                
            self.c_decreto.at[nro_fila, "texto"] = self.c_decreto.at[nro_fila, "texto"].replace("[#]", f"{self.count_parrafo}°")
            self.frases_lineas(documento=document, text=self.c_decreto["texto"][nro_fila], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))

            self.clausula_beneficio_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
            self.clausulas_contrato(nro_o_tipo_clausula="6") # escribe claustula beneficio de contrato

        # Tarjeta de vestuario
        if "Tarjeta de vestuario" in self.persona.nombres_beneficios_seleccionados:
            print("Estoy escribiendo tarjeta de vestuario")
            self.count_parrafo += 1
            self.index_nro_subt_clausulas += 1
            nro_fila = 13
            self.index_csv_contrato = 12
            self.c_decreto.at[nro_fila, "texto"] = self.c_decreto.at[nro_fila, "texto"].replace("[#]", f"{self.count_parrafo}°")
            self.frases_lineas(documento=document, text=self.c_decreto["texto"][nro_fila], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))

            self.clausula_beneficio_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
            self.clausulas_contrato(nro_o_tipo_clausula="6") # escribe claustula beneficio de contrato

        # Vestuario prestado
        if "Vestuario prestado" in self.persona.nombres_beneficios_seleccionados:
            print("Estoy escribiendo vestuario prestado")
            self.count_parrafo += 1
            self.index_nro_subt_clausulas += 1
            nro_fila = 14
            self.index_csv_contrato = 13
            self.c_decreto.at[nro_fila, "texto"] = self.c_decreto.at[nro_fila, "texto"].replace("[#]", f"{self.count_parrafo}°")
            self.frases_lineas(documento=document, text=self.c_decreto["texto"][nro_fila], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))

            self.clausula_beneficio_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
            self.clausulas_contrato(nro_o_tipo_clausula="6") # escribe claustula beneficio de contrato

        # Laptop y celular
        if "Laptop" in self.persona.nombres_beneficios_seleccionados and "Celular" not in self.persona.nombres_beneficios_seleccionados: # Solo laptop
            self.count_parrafo += 1
            self.index_nro_subt_clausulas += 1
            nro_fila = 19
            self.index_csv_contrato = 18
            for csv in [self.c_decreto, self.csv_contratos]:
                csv["texto"] = csv["texto"].str.replace("[una laptop]", "una laptop") 
                csv["texto"] = csv["texto"].str.replace("[un celular]", "")
                csv["texto"] = csv["texto"].str.replace("[y_2]", "")
            self.c_decreto.at[nro_fila, "texto"] = self.c_decreto.at[nro_fila, "texto"].replace("[#]", f"{self.count_parrafo}°")
            self.frases_lineas(documento=document, text=self.c_decreto["texto"][nro_fila], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))

            self.clausula_beneficio_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
            self.clausulas_contrato(nro_o_tipo_clausula="6") # escribe claustula beneficio de contrato
            
        elif  "Laptop" not in self.persona.nombres_beneficios_seleccionados and "Celular" in self.persona.nombres_beneficios_seleccionados: # Solo celular
            self.count_parrafo += 1
            self.index_nro_subt_clausulas += 1
            nro_fila = 19
            self.index_csv_contrato = 18
            for csv in [self.c_decreto, self.csv_contratos]:
                csv["texto"] = csv["texto"].str.replace("[una laptop]", "")
                csv["texto"] = csv["texto"].str.replace("[un celular]", "un celular") 
                csv["texto"] = csv["texto"].str.replace("[y_2]", "")
            self.c_decreto.at[nro_fila, "texto"] = self.c_decreto.at[nro_fila, "texto"].replace("[#]", f"{self.count_parrafo}°")
            self.frases_lineas(documento=document, text=self.c_decreto["texto"][nro_fila], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))

            self.clausula_beneficio_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
            self.clausulas_contrato(nro_o_tipo_clausula="6") # escribe claustula beneficio de contrato
            
        elif "Laptop" in self.persona.nombres_beneficios_seleccionados and "Celular" in self.persona.nombres_beneficios_seleccionados: # Ambos
            self.count_parrafo += 1
            self.index_nro_subt_clausulas += 1
            nro_fila = 19
            self.index_csv_contrato = 18
            for csv in [self.c_decreto, self.csv_contratos]:
                csv["texto"] = csv["texto"].str.replace("[una laptop]", "una laptop") 
                csv["texto"] = csv["texto"].str.replace("[un celular]", "un celular") 
                csv["texto"] = csv["texto"].str.replace("[y_2]", " y ")
            self.c_decreto.at[nro_fila, "texto"] = self.c_decreto.at[nro_fila, "texto"].replace("[#]", f"{self.count_parrafo}°")
            self.frases_lineas(documento=document, text=self.c_decreto["texto"][nro_fila], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, palabra_bold=self.persona.nombre, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))

            self.clausula_beneficio_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
            self.clausulas_contrato(nro_o_tipo_clausula="6") # escribe claustula beneficio de contrato

        # Parrafo finales de contrato

        self.clausulas_contrato(nro_o_tipo_clausula="finales")

        #############################################

        for i_siempre in [21,22,23,24,25,26,27,28]:
            #print("Finalizó ultiomas clausulas decreto")
            
            if i_siempre == 21 and not self.persona.sala_cuna:
                continue  # Salta la iteración si i_siempre es 21 y sala_cuna no es True

            self.count_parrafo += 1
            self.c_decreto.at[i_siempre, "texto"] = self.c_decreto.at[i_siempre, "texto"].replace("[#]", f"{self.count_parrafo}°")
            # Solo agregar el párrafo si i_siempre no es 20 o si es 20 y sala_cuna es True

            if i_siempre in [25, 26, 27]:
                palabra_en_negrita = ""
            else:
                palabra_en_negrita = self.persona.nombre

            self.frases_lineas(documento=document, text=self.c_decreto["texto"][i_siempre], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, palabra_bold=palabra_en_negrita, sangria_derecha = Cm(0), sangria_izq = Cm(1), sangria_primera_linea = Cm(-0.5))
        
        self.frases_lineas(documento=document, text="\nANÓTESE, DISTRIBÚYASE Y REGÍSTRESE\n\n\n", posicion=WD_ALIGN_PARAGRAPH.LEFT, fuente=12, bold_bool=True, subrayado=False)
        self.frases_lineas(documento=document, text="ALCALDESA\n\n", posicion=WD_ALIGN_PARAGRAPH.LEFT, fuente=12, bold_bool=True, subrayado=False)
        self.frases_lineas(documento=document, text="SECRETARIO MUNICIPAL", posicion=WD_ALIGN_PARAGRAPH.RIGHT, fuente=12, bold_bool=True, subrayado=False)

        if self.cargo.departamento != "Municipal": # Educacion o salud
            subdi_o_depto = f"- Departamento de {self.cargo.departamento}"
        elif self.cargo.departamento == "Municipal":
            subdi_o_depto = ""
            pass 
        
        
        list_lineas = [f"{self.cargo.visadora_1}/{self.cargo.visadora_2}/{self.cargo.redactora}", 
                       "Distribución:", 
                       "- SIAPER Contraloría General de la República", 
                        "- Secretaría Municipal", 
                        "- Contraloría Municipal", 
                        "- Subdirección de Finanzas", # agregar en condicional: depto o subdireccion, condicionar que se escriba solo si existe
                        f"- {self.cargo.direccion}", 
                        "- Departamento de Personal - Sección Remuneraciones", 
                        "- Departamento D.O - Dpto. Bienestar y Calidad de Vida", 
                        "- Interesado", 
                        "- Oficina de Partes"] # hacer primera linea variable texto libre prellenada
        
        if subdi_o_depto != "":
            print("Agregando a la distribuación a: ", subdi_o_depto)
            list_lineas.insert(6, subdi_o_depto)

        for text_lineas in list_lineas: 
            #print("Escribiendo distribución")
            if text_lineas == "Distribución:": # subrayado
                self.frases_lineas(documento=document, text=text_lineas, posicion=WD_ALIGN_PARAGRAPH.LEFT, fuente=10, bold_bool=False, space_after=Pt(1.5), subrayado=True)
            elif text_lineas == "- SIAPER Contraloría General de la República" or text_lineas == f"- {self.cargo.direccion}" or text_lineas == "- Interesado": # con negrita
                self.frases_lineas(documento=document, text=text_lineas, posicion=WD_ALIGN_PARAGRAPH.LEFT, fuente=10, bold_bool=True, space_after=Pt(1.5), subrayado=False)
            else: # sin negrita
                    self.frases_lineas(documento=document, text=text_lineas, posicion=WD_ALIGN_PARAGRAPH.LEFT, fuente=10, bold_bool=False, space_after=Pt(1.5), subrayado=False)
        print("Finalizó distribución")

    def frases_lineas(self, documento, text, posicion, fuente, bold_bool, space_after=Pt(0), subrayado=False, palabra_bold = "", sangria_derecha = Cm(0), sangria_izq = Cm(0), sangria_primera_linea = Cm(0)):  
        # Agregar frases
        frase = documento.add_paragraph()
        
        if palabra_bold:
            # Dividir el texto en partes antes y después de la palabra en negrita
            parts = text.split(palabra_bold)
            
            # Agregar la parte antes de la palabra en negrita
            if parts[0]:
                run = frase.add_run(parts[0])
                run.bold = bold_bool
                run.font.size = Pt(fuente)
                run.underline = subrayado
            
            # Agregar la palabra en negrita
            run_bold = frase.add_run(palabra_bold)
            run_bold.bold = True
            run_bold.font.size = Pt(fuente)
            run_bold.underline = subrayado
            
            # Agregar la parte después de la palabra en negrita
            if len(parts) > 1 and parts[1]:
                run = frase.add_run(parts[1])
                run.bold = bold_bool
                run.font.size = Pt(fuente)
                run.underline = subrayado
        else:
            run = frase.add_run(text)  # Permite editar formato texto
            run.bold = bold_bool  # True o False
            run.font.size = Pt(fuente)  # Establecer el tamaño de la fuente
            run.underline = subrayado  # Subrayado
        
        frase.alignment = posicion # Posicion: WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if space_after:
            frase.paragraph_format.space_after = space_after  # Ajustar el espacio después del párrafo
        if sangria_derecha:
            frase.paragraph_format.right_indent = sangria_derecha  # Margen derecho, dentro del margen ya existente en la seccion izq 3 y derecha 2.54
        if sangria_primera_linea:
            frase.paragraph_format.first_line_indent = sangria_primera_linea
        if sangria_izq:
            frase.paragraph_format.left_indent = sangria_izq  # Margen izquierdo, dentro del margen ya existente en la seccion izq 3 y derecha 2.54
        
    def escribir_parrafos_con_vinetas(self, filas_parrafos_en_vineta, doc, margen_izquierdo=Cm(1.75), margen_derecho=Cm(0), posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, sangria_primera_linea=Cm(-0.5)):
        """
        Escribe párrafos con viñetas a partir de filas de una columna de un DataFrame en un documento de Word.
        
        :param doc: documento docx en el cual se está escribiendo
        :param list_p_cometidos: DataFrame de pandas que contiene los datos.
        :param columna: Nombre de la columna que contiene el texto.
        :param margen_izquierdo: Margen izquierdo en cm.
        :param margen_derecho: Margen derecho en cm.
        :param sangria_primera_linea: Margen primera linea en cm.
        """
        # Iterar sobre los índices especificados y agregar párrafos con viñetas
        if type(filas_parrafos_en_vineta) == str: # caso de cometido suma alzada, un parrafo
            p = doc.add_paragraph()
            p.add_run(filas_parrafos_en_vineta)
            p.style = 'List Bullet'
            p.alignment = posicion  # Posicion: WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.JUSTIFY
            # Margenes de los parrados
            #p.paragraph_format.space_after = space_after
            p.paragraph_format.first_line_indent = sangria_primera_linea
            p.paragraph_format.right_indent = margen_derecho
            p.paragraph_format.left_indent = margen_izquierdo
            return
        else: # Lista de strings, caso de varios parrafos
            for texto in filas_parrafos_en_vineta:
                print("Escribimos en viñeta: ", texto)
                p = doc.add_paragraph()
                p.add_run(texto)
                p.style = 'List Bullet' # podria hacerlo como parametro para poner opcion de numerar
                p.alignment = posicion  # Posicion: WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.JUSTIFY
                # Margenes de los parrados
                #p.paragraph_format.space_after = space_after
                p.paragraph_format.first_line_indent = sangria_primera_linea
                p.paragraph_format.right_indent = margen_derecho
                p.paragraph_format.left_indent = margen_izquierdo
            return
    
    def obtener_y_escribir_cometidos(self, docu):
        # Obtener los cometidos del cargo
        if self.cargo.departamento == "Salud":
            self.lista_lineas_cometidos = []
            # Leer el archivo CSV y filtrar las filas que coincidan con self.cargo.serv_salud
            with open(self.path_cometidos_salud, newline='', encoding='utf-8') as csvfile: 
                reader = csv.reader(csvfile, delimiter='|')
                for row in reader:
                    if row[0].strip().lower() == self.cargo.serv_salud.strip().lower() and row[2].strip() == self.cargo.especialidad_salud.strip():
                        self.lista_lineas_cometidos.append(row[3].strip())
            
            # Escribir las filas correspondientes en el documento
            self.escribir_parrafos_con_vinetas(filas_parrafos_en_vineta=self.lista_lineas_cometidos, doc=docu)
        elif self.cargo.departamento == "Municipal":
            print(f"Nro cometido SA: {self.cargo.numero_cometido_sa}")
            print(f"Cometido SA: {self.cargo.text_cometido_sa}")
            indice_nro_cometido = self.lista_numero_cometido_sa.index(int(self.cargo.numero_cometido_sa))
            self.cometido_sa =  self.lista_text_cometido_sa[indice_nro_cometido]
            # Escribimos el título del cometido subrallado y con su respectivo número
            self.escribir_parrafos_con_vinetas(filas_parrafos_en_vineta=self.cometido_sa, doc=docu)
        else: # educacion
            self.escribir_parrafos_con_vinetas(filas_parrafos_en_vineta=self.cargo.text_cometido_educacion, doc=docu)
            
    
    # Siendo self.cargo.fechas_dias_contratacion es una lista de strings con fechas en el formato "%d de %B de %Y"
    def obtener_fecha_menor_y_mayor(self, fechas):
        # Convertir las fechas a objetos datetime
        # fechas = fechas.sort() # ordena la lista de fechas de menor a mayor, se supone que ya fueron ordenadas en la funcion cel widget
        ### fechas_datetime = [datetime.strptime(fecha, "%d de %B de %Y") for fecha in fechas] # transforma string a datetime  02 dic

        # Encontrar la fecha más temprana
        ### self.fecha_menor = min(fechas_datetime) # datetime, tambien podria elegir el primer elemento de la lista 02 dic
        ### self.fecha_mayor = max(fechas_datetime) # datetime, tambien podria elegir el ultimo elemento de la lista 02 dic
        print("fechas (ordenadas): ", fechas)
        self.fecha_menor = fechas[0]
        self.fecha_mayor = fechas[-1]

        # Convertir la fecha más temprana de nuevo a string en el formato original
        ### return self.fecha_menor.strftime("%d de %B de %Y"), self.fecha_mayor.strftime("%d de %B de %Y") # entrega strings
        return self.fecha_menor, self.fecha_mayor
    
    def validar_correo(self, *args): # paso un número variable de argumentos posicionales a la función
        # Expresión regular para validar correos electrónicos
        self.persona.mail= self.entrada_p_mail.get()
        patron = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if re.match(patron, self.persona.mail): # patron de validación coincide con patron del mail digitado
            self.persona.validez_mail.set(True) # Modificar la variable para activar la traza
            self.bool_mail = self.persona.validez_mail.get()
            print("El correo es válido.")
        else:
            self.persona.validez_mail.set(False)
            self.bool_mail = self.persona.validez_mail.get()
            print("El correo no es válido.")
            self.mostrar_error("Mail incorrecto", "El mail ingresado no es valido, vuelvalo a intentar.")
        return self.bool_mail


    def verificar_rut(self, *args):
        print("Verificando RUT con función verificar_rut...")
        # Obtenemos el rut y el dígito verificador desde las entradas
        self.persona.rut = self.entrada_p_rut.get()
        self.persona.dig_ver = self.entrada_p_rut_dig_ver.get()

        # Ya tenemos separados el número del RUT y el dígito verificador
        rut_sin_dv = self.persona.rut.replace('.', '')  # Eliminar puntos si existen del string
        dv_proporcionado = self.persona.dig_ver # string

        # Convertir el número del RUT a una lista de enteros
        rut_numeros = list(map(int, rut_sin_dv))

        # Multiplicar cada dígito por los números 2, 3, 4, 5, 6, y 7 de derecha a izquierda
        multiplicadores = [2, 3, 4, 5, 6, 7]
        suma = 0
        for i, numero in enumerate(reversed(rut_numeros)):
            suma += numero * multiplicadores[i % len(multiplicadores)]

        # Calcular el dígito verificador
        resto = suma % 11
        dv_calculado = 11 - resto
        if dv_calculado == 11:
            dv_calculado = '0'
        elif dv_calculado == 10:
            dv_calculado = 'K'
        else:
            dv_calculado = str(dv_calculado)

        # Comparar el dígito verificador calculado con el proporcionado
        if dv_calculado != dv_proporcionado.upper():
            print("dv calculado: ", dv_calculado)
            print("dv proporcionado: ", dv_proporcionado)
            print("RUT incorrecto :()")
            self.mostrar_error("RUT incorrecto", "El RUT ingresado es incorrecto.")
            self.persona.validez_rut.set(False) # Modificar la variable para activar la traza y llama a verificar_rut
            self.valor = self.persona.validez_rut.get()
            return self.valor # False, si es falso se corta la función aqui, sino no entra al if y queda como true
        print("dv calculado: ", dv_calculado)
        print("dv proporcionado: ", dv_proporcionado)
        print("RUT correcto :)")

        self.persona.validez_rut.set(True)
        self.valor = self.persona.validez_rut.get()
        return self.valor # True
            #return dv_calculado == dv_proporcionado.upper() # Compara strings

    def join_art_nombre(self, genero):
        if genero == "Femenino":
            return f"la Sra.", f"la prestadora", f"la servidora"
        elif genero == "Masculino":
            return f"el Sr.", f"el prestador", f"el servidor"
        elif genero == "Otro":
            return f"", f"la persona prestadora", f"la persona servidora"

    def guardar_en_path(self):
        self.verificar_rut()
        self.validar_correo()
        if self.persona.validez_rut.get() == True and self.persona.validez_mail.get() == True:
            print("Validez del rut: ", self.persona.validez_rut.get())
            print("Validez del mail: ", self.persona.validez_mail.get())
            # Guarda variables
            self.guardar()
            # Crea documento word
            self.crear_word_decretos() # activa la creación del word de contrato por dentro
            # Guardar el documento en un directorio específico
            # Cuadro de diálogo para guardar el archivo

            # Iterar sobre todos los párrafos y `runs` del documento para modificar la fuente TIMES NEW ROMAN
            for parrafo in document.paragraphs:
                for run in parrafo.runs:
                    run.font.name = 'Times New Roman'  # Establecer la fuente
            for parrafo in document_contrato.paragraphs:
                for run in parrafo.runs:
                    run.font.name = 'Times New Roman'  # Establecer la fuente
            file_path_decreto = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], initialfile=f"Decreto {self.cargo.tipo_decretoycontrato} {self.persona.nombre}.docx")
            file_path_contrato = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], initialfile=f"Contrato {self.cargo.tipo_decretoycontrato} {self.persona.nombre}.docx")
            if file_path_decreto:
                document.save(file_path_decreto)
                print("Documento Word DECRETO creado con éxito en:", file_path_decreto)
            if file_path_contrato:
                document_contrato.save(file_path_contrato)
                print("Documento Word CONTRATO creado con éxito en:", file_path_contrato)
        elif self.persona.validez_rut.get() == False or self.persona.validez_mail.get()  == False:
            print("Validez del rut (else): ", self.persona.validez_rut.get())
            print("Validez del mail (else): ", self.persona.validez_mail.get())
            self.mostrar_error("Error", "No se puede crear los documentos porque hay errores en los datos ingresados.")
 
    def word_contratos(self):
        self.nombre_csv_alcalde = ["Alcaldesa_Camila_Merino.csv", "Alcalde_S_Rodrigo_Zalaquett.csv"]
        # Selección de alcaldesa o alcalde suplente
        if self.cargo.alcaldia == "Camila Merino Catalán":
            indice_alcalde = 0 
        elif self.cargo.alcaldia == "Rodrigo Zalaquett (S)":
            indice_alcalde = 1 
        
        file_path_contratos = os.path.join(path_a_utilizar, "clausulas_csv", "contratos.csv")
        self.csv_contratos = pd.read_csv(file_path_contratos, sep="|", encoding="utf-8", header=None, usecols=[1], names=["texto"])
        file_path_alcalde = os.path.join(path_a_utilizar, "clausulas_csv", self.nombre_csv_alcalde[indice_alcalde])
        csv_alcalde = pd.read_csv(file_path_alcalde, sep="|", encoding="utf-8", header=None, usecols=[0,1], names=["por_reemplazar","texto"])

        self.csv_contratos["texto"] = self.csv_contratos["texto"].astype("string") 
        csv_alcalde["por_reemplazar"] = csv_alcalde["por_reemplazar"].astype("string")
        csv_alcalde["texto"] = csv_alcalde["texto"].astype("string") # Hacemos columna string, podria haberlo hecho en el read_csv

        # Variables alcalde
        self.art_alcalde_s = csv_alcalde["texto"][0] # Artículo alcalde
        self.nombre_alcalde = csv_alcalde["texto"][1] # Nombre alcalde
        self.rut_alcalde = csv_alcalde["texto"][2] # Rut alcalde
        self.nacionalidad_alcalde = csv_alcalde["texto"][3] # Nacionalidad alcalde: chilena/chileno
        self.estado_civil_alcalde = csv_alcalde["texto"][4] # Estado civil alcalde
        self.profesion_alcalde = csv_alcalde["texto"][5] # Profesión alcalde

        #################################
        ###### Reemplazo variables ######
        #################################

        for index in range(len(csv_alcalde["por_reemplazar"])): # Reemplazamos información de alcalde (de csv_alcalde) en primer parrafo del contrato (self.csv_contratos)
            self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace(csv_alcalde["por_reemplazar"][index], csv_alcalde["texto"][index])

        # Reemplazo información de fecha del instrumento contrato
        self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[DÍA de MES de AÑO CONTRATO]", self.cargo.fecha_instrumento_contrato)
        # Reemplazo información de la persona en el contrato
        self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[NOMBRE PERSONA]", self.persona.nombre)
        self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[RUT PERSONA]", f"{self.rut_con_puntos}-{self.persona.dig_ver}")
        self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[NACIONALIDAD PERSONA]", f"de nacionalidad {self.persona.nacionalidad}")
        self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[ESTADO CIVIL PERSONA]", self.persona.estado_civil)
        self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[PROFESIÓN PERSONA]", self.persona.profesion)

        self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[ARTÍCULO + servidor/servidora]", self.art_servidor)
        if self.cargo.medio_validador == "Informe":
            self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[DIRECTIVO A QUIEN SE INFORMA]", self.directivo_a_quien_se_informa)
        self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[Total/Mensual]", self.cargo.tipo_renta)

        # Domicilio
        if self.persona.aclaracion_domicilio == "": # no existe aclaración de domicilio
            self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[DOMICILIO PERSONA]", f"{self.persona.domicilio} #{self.persona.nro_domicilio}, comuna de {self.persona.comuna}")
        else: # existe aclaración de domicilio
            self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[DOMICILIO PERSONA]", f"{self.persona.domicilio} #{self.persona.nro_domicilio}, {self.persona.aclaracion_domicilio}, comuna de {self.persona.comuna}") 

        ##################################
        #### Parrafo inicial contrato ####
        ##################################
        
        self.frases_lineas(documento=document_contrato, text=self.csv_contratos["texto"][0], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False, palabra_bold=self.persona.nombre) 
        
        # Clausulas contratos
        return None

    ## Insertar un salto de página
    ## document.add_page_break()

    def clausulas_contrato(self, nro_o_tipo_clausula):
        # Lista de subtitulos clausulas
        list_nro_subt_clausulas = ["PRIMERO", "SEGUNDO", "TERCERO", "CUARTO", "QUINTO", "SEXTO", "SEPTIMO", 
                                    "OCTAVO", "NOVENO", "DÉCIMO", "DÉCIMO PRIMERO", "DÉCIMO SEGUNDO", "DÉCIMO TERCERO", "DECIMO CUARTO",
                                    "DÉCIMO QUINTO", "DÉCIMO SEXTO", "DÉCIMO SEPTIMO", "DÉCIMO OCTAVO", "DÉCIMO NOVENO", "VIGÉSIMO"]
            
        indices_clausulas_finales_contrato_siempre = [21, 22, 23, 24, 25]

        ##################################
        #### Clausulas contrato ####
        ##################################        
                
        # Subtitulo de contrato: Agregar negrita y subrayado al parrafo, alineado a la izq.

        if nro_o_tipo_clausula == "1":
            ## Primera clausula
            #self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[COMETIDO TEXT]", f"\n\n{self.cargo.text_cometido_salud}") # Va en salud, ed, y municipal
            
            self.frases_lineas(documento=document_contrato, text=f"{list_nro_subt_clausulas[self.index_nro_subt_clausulas]}:", posicion=WD_ALIGN_PARAGRAPH.LEFT, fuente=12, bold_bool=True, subrayado=True)
            
            ### Salud
            if self.cargo.departamento == "Salud":
                self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[ESPECIALIDAD SALUD]", self.cargo.especialidad_salud)
                self.index_csv_contrato = 1
                if self.cargo.tipo_contrato == "Programa":
                    self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[NOMBRE PROGRAMA]", self.cargo.programa) 
                    self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[, en el NOMBRE PROGRAMA]", f", en el {self.cargo.programa}")
                elif self.cargo.tipo_contrato == "Suma alzada":
                    self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[, en el NOMBRE PROGRAMA]", "")
                self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[COSAM/CESFAM/Departamento de Salud Municipal/SAPU]", "Departamento de Salud Municipal de Vitacura") 
                # Titulo 
                self.frases_lineas(documento=document_contrato, text=self.csv_contratos["texto"][self.index_csv_contrato], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False)
                
            ### Educación
            elif self.cargo.departamento == "Educación":
                self.index_csv_contrato = 2
                self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[COSAM/CESFAM/Departamento de Salud Municipal/SAPU]", "Departamento de Educación Municipal de Vitacura") 
                self.frases_lineas(documento=document_contrato, text=self.csv_contratos["texto"][self.index_csv_contrato], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False)
            ### Municipal
            elif self.cargo.departamento == "Municipal":
                if self.cargo.tipo_contrato == "Programa": # Tambien lleva text de cometido
                    self.index_csv_contrato = 3
                    self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[NOMBRE PROGRAMA]", self.cargo.programa) 
                elif self.cargo.tipo_contrato == "Suma alzada":
                    self.index_csv_contrato = 4
                self.frases_lineas(documento=document_contrato, text=self.csv_contratos["texto"][self.index_csv_contrato], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False)
                self.frases_lineas(documento=document_contrato, text=f"Cometido N°{self.cargo.numero_cometido_sa}:", posicion=WD_ALIGN_PARAGRAPH.LEFT, fuente=12, bold_bool=False, subrayado=True, sangria_izq=Cm(0.75))

            # Escribe cometido con viñeta
            self.obtener_y_escribir_cometidos(docu=document_contrato)

        elif nro_o_tipo_clausula == "2":
            ## Segunda clausula
            self.index_nro_subt_clausulas += 1

            if self.cargo.renta_uf_clp == "CLP":
                self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[RB IMP RL]", f"de ${self.RB}.- con la deducción del {self.imp}% por concepto de Impuesto a la Renta, lo que da un líquido a pagar de ${self.cargo.renta_liquida}.-") 
            elif self.cargo.renta_uf_clp == "UF":
                self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[RB IMP RL]", f"de UF {self.RB}.- (impuesto incluido)") 

            self.frases_lineas(documento=document_contrato, text=f"{list_nro_subt_clausulas[self.index_nro_subt_clausulas]}:", posicion=WD_ALIGN_PARAGRAPH.LEFT, fuente=12, bold_bool=True, subrayado=True)

            if self.cargo.medio_validador == "Informe":
                self.index_csv_contrato = 5
                self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[Final/Mensual]", self.cargo.mensual_final) # "Mensual" o "Final"
                if self.cargo.mensual_final == "Final":
                    self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[un Informe Final/Informes Mensuales]", "un Informe Final")
                elif self.cargo.mensual_final == "Mensual":
                    self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[un Informe Final/Informes Mensuales]", "Informes Mensuales")
                clausula_segunda_contrato = self.csv_contratos["texto"][self.index_csv_contrato]

            elif self.cargo.medio_validador == "Certificado":
                self.index_csv_contrato = 6
                if self.cargo.mensual_final == "Final":
                    self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[Certificaciones de Cumplimientos Mensuales y previa visación de las correspondientes Boletas de Honorarios]", "Certificación de Cumplimiento y previa visación de la correspondiente Boleta de Honorarios")
                elif self.cargo.mensual_final == "Mensual":
                    self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[Certificaciones de Cumplimientos Mensuales y previa visación de las correspondientes Boletas de Honorarios]", "Certificaciones de Cumplimientos Mensuales y previa visación de las correspondientes Boletas de Honorarios")
                clausula_segunda_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
            # elif talleristas 7
            self.frases_lineas(documento=document_contrato, text=clausula_segunda_contrato, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False)

        elif nro_o_tipo_clausula == "3":
            ## Tercera clausula: ley 21.133 BBHH
            self.index_nro_subt_clausulas += 1
            self.frases_lineas(documento=document_contrato, text=f"{list_nro_subt_clausulas[self.index_nro_subt_clausulas]}:", posicion=WD_ALIGN_PARAGRAPH.LEFT, fuente=12, bold_bool=True, subrayado=True)

            self.index_csv_contrato = 8
            clausula_tercera_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
            self.frases_lineas(documento=document_contrato, text=clausula_tercera_contrato, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False)
            
        elif nro_o_tipo_clausula == "4":
            # Cuarta clausula: duración contrato
            if self.cargo.periodo_contratacion == "Por periodo":
                self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[DÍA de MES de AÑO INICIO]", self.cargo.fecha_inicio)
                self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[DÍA de MES de AÑO TÉRMINO]", self.cargo.fecha_termino)
            elif self.cargo.periodo_contratacion == "Por días":
                self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("desde el [DÍA de MES de AÑO INICIO] hasta el [DÍA de MES de AÑO TÉRMINO]", f"por los días {self.fechas_unidas}")
            elif self.cargo.periodo_contratacion == "Por un día":
                self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("desde el [DÍA de MES de AÑO INICIO] hasta el [DÍA de MES de AÑO TÉRMINO]", f"por el día {self.cargo.fechas_dias_contratacion[0]}")
            self.index_nro_subt_clausulas += 1
            self.frases_lineas(documento=document_contrato, text=f"{list_nro_subt_clausulas[self.index_nro_subt_clausulas]}:", posicion=WD_ALIGN_PARAGRAPH.LEFT, fuente=12, bold_bool=True, subrayado=True)
            self.index_csv_contrato = 9
                
            clausula_cuarta_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
            self.frases_lineas(documento=document_contrato, text=clausula_cuarta_contrato, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False)

        elif nro_o_tipo_clausula == "5":
            print("clausula beneficio contrato : elif nro_o_tipo_clausula == '5':")
            # Quinta clausula
            if self.cargo.medio_validador != "Certificado":
                self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[MES TERMINO de AÑO]", self.mes_termino_de_anio)  
                self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[MES INICIO]", self.mes_inicio)  
                self.csv_contratos["texto"] = self.csv_contratos["texto"].str.replace("[MES PENÚLTIMO]", self.mes_penultimo) 
                self.index_nro_subt_clausulas += 1
                self.frases_lineas(documento=document_contrato, text=f"{list_nro_subt_clausulas[self.index_nro_subt_clausulas]}:", posicion=WD_ALIGN_PARAGRAPH.LEFT, fuente=12, bold_bool=True, subrayado=True)
                self.index_csv_contrato = 10
                clausula_quinta_contrato = self.csv_contratos["texto"][self.index_csv_contrato]
                self.frases_lineas(documento=document_contrato, text=clausula_quinta_contrato, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False)
            else: # en caso de que medio validador sea certificado, no corre esta clausula
                pass

        # Sexta clausula: Beneficios contrato
        elif nro_o_tipo_clausula == "6":
            print("clausula beneficio contrato : elif nro_o_tipo_clausula == '6':")
            #self.index_nro_subt_clausulas += 1
            self.frases_lineas(documento=document_contrato, text=f"{list_nro_subt_clausulas[self.index_nro_subt_clausulas]}:", posicion=WD_ALIGN_PARAGRAPH.LEFT, fuente=12, bold_bool=True, subrayado=True)
            self.frases_lineas(documento=document_contrato, text=self.clausula_beneficio_contrato, posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False)

        # Clausualas finales que van siempre
        elif nro_o_tipo_clausula == "finales":
            for i_siempre in indices_clausulas_finales_contrato_siempre:
                self.index_nro_subt_clausulas += 1
                self.frases_lineas(documento=document_contrato, text=f"{list_nro_subt_clausulas[self.index_nro_subt_clausulas]}:", posicion=WD_ALIGN_PARAGRAPH.LEFT, fuente=12, bold_bool=True, subrayado=True)
                self.frases_lineas(documento=document_contrato, text=self.csv_contratos["texto"][i_siempre], posicion=WD_ALIGN_PARAGRAPH.JUSTIFY, fuente=12, bold_bool=False, subrayado=False)

            # Nombres para firmas
            ## Persona
            self.frases_lineas(documento=document_contrato, text=f"\n\n\n{self.persona.nombre.upper()}\nR.U.T. {self.rut_con_puntos}-{self.persona.dig_ver}", posicion=WD_ALIGN_PARAGRAPH.CENTER, fuente=12, bold_bool=True, subrayado=False, sangria_derecha=Cm(10))
            ## Alcalde(sa)
            self.frases_lineas(documento=document_contrato, text=f"{self.nombre_alcalde.upper()}\nR.U.T. {self.rut_alcalde}\n{self.art_alcalde_s.upper()}", posicion=WD_ALIGN_PARAGRAPH.CENTER, fuente=12, bold_bool=True, subrayado=False, sangria_izq=Cm(10)) 
            
            # Distribución visación y redacción
            ultima_linea_contrato = f"{self.cargo.visadora_1}/{self.cargo.visadora_2}/{self.cargo.redactora}"

            self.frases_lineas(documento=document_contrato, text=ultima_linea_contrato, posicion=WD_ALIGN_PARAGRAPH.LEFT, fuente=11, bold_bool=True, subrayado=False)
            print("Contrato finalizado (finalizó escritura de clausulas finales)")

        return None
    
    def crear_widget_con_label(self, frame, texto_label, nombre_variable, fila_grid, columna_grid, nombre_widget, estado="normal", lista_opciones=None, ancho=30, font=None, command=None, tooltip=None, set_valor_inicial=""):
        # Crear el Label
        label = ttk.Label(frame, text=texto_label, style="TLabel", width=30)
        label.grid(row=fila_grid, column=columna_grid, padx=5, pady=5, sticky="w")
        
        # Crear la variable correspondiente
        if nombre_widget == 'Entry':
            variable = tk.StringVar(name=nombre_variable, value=set_valor_inicial)
            widget = tk.Entry(frame, textvariable=variable, state=estado, width=ancho, font=font)

        elif nombre_widget == 'Combobox':
            variable = tk.StringVar(name=nombre_variable, value=set_valor_inicial)
            variable.set(set_valor_inicial)
            widget = ttkb.Combobox(frame, textvariable=variable, values=lista_opciones, state=estado, width=ancho, font=font, height=10) # postcommand=lambda: variable.set(variable.get()) para que se actualice el valor del combobox, ttkb permite hacer busqueda en combobox
            def update_nacionalidad(event):
                self.persona.nacionalidad = variable.get().title()
            # Buscador en lista de opciones
            widget.bind('<KeyRelease>', lambda event: self._filter_combobox(event, widget, lista_opciones))
            widget.bind("<<ComboboxSelected>>", lambda event: update_nacionalidad) # para que se actualice el valor del combobox
            
            

        elif nombre_widget == 'DateEntry':
            variable = tk.StringVar(name=nombre_variable, value=set_valor_inicial)
            widget = DateEntry(frame, textvariable=variable, width=ancho, dateformat='%d/%m/%Y', font=font, bootstyle="primary", firstweekday=0)

        elif nombre_widget == 'Checkbutton':
            variable = tk.BooleanVar(name=nombre_variable, value=False)
            widget = tk.Checkbutton(frame, variable=variable, state=estado, width=ancho, font=font, command=command)
        else:
            raise ValueError("Tipo de widget no soportado")
        
        # Colocar el widget en la posición especificada
        widget.grid(row=fila_grid, column=columna_grid+1, padx=5, pady=5, sticky="w")
        
        # Agregar tooltip si se proporciona
        if tooltip:
            def on_enter(event):
                tooltip_label = ttk.Label(frame, text=tooltip, background="yellow", font=font, sticky="w")
                tooltip_label.place(x=event.x_root, y=event.y_root)
                widget.tooltip_label = tooltip_label

            def on_leave(event):
                widget.tooltip_label.destroy()

            widget.bind("<Enter>", on_enter)
            widget.bind("<Leave>", on_leave)

    def actualizar_opciones_combobox(self, nueva_lista, combobox_name): # sirve como postcommand para combobox
        # Actualizar las opciones del Combobox
        combobox_name['values'] = nueva_lista
        print("Opciones actualizadas")
    
    def _filter_combobox(self, event, combobox_name, values):
        value = event.widget.get()
        if value == '':
            combobox_name['values'] = values
        else:
            data = []
            for item in values:
                if value.lower() in item.lower():
                    data.append(item)
            combobox_name['values'] = data

    def mostrar_error(self, titulo, mensaje): # Esta función crea una ventana de mensaje de error utilizando Tkinter
        root = tk.Tk()
        root.withdraw()  # Ocultar la ventana principal de Tkinter
        messagebox.showerror(titulo, mensaje)
        root.destroy()

    def destruir_widget(self, name_widget):
        # Destruir el widget si existe
        if hasattr(self, name_widget):
            self.name_widget.destroy()


if __name__ == "__main__": # Envuelve la inicialización de la aplicación y el bucle principal de Tkinter en un bloque try-except
    try:
        root = tk.Tk()
        app = Aplicacion(root)
        root.mainloop()

    except Exception as e:
        # Captura cualquier excepción y muestra el mensaje de error
        root = tk.Tk()
        root.withdraw()  # Ocultar la ventana principal de Tkinter
        messagebox.showerror("Error", f"Ocurrió un error: {str(e)}")
        root.destroy()


# PENDIENTE EN OPTIMIZACION: CREAR FUNCIONES PARA CREACION DE WIDGETS, VARIABLES Y PARRAFOS, SEPARAR EL ARCHIVO .PY EN VARIOS ARCHIVOS .PY. REVISAR TIEMPOS DE EJECUCION Y OPTIMIZAR CODIGO.    
# PENDIENTE: EN VEZ DE IMPRIMIR LA INFORMACIÓN DEL CARGO Y LA PERSONA, CREAR Y GUARDAR LA INFORMACION EN UN ARCHIVO CSV
