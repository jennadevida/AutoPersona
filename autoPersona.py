
#!/usr/bin/env python3
# -*- coding: utf-8 -*-

### seguir en este, actualizar luego subiendo a github

### ideas: ir numerando con +1 a medida que se van agregando parrafos, esto para evitar duplicacion de numero o nros faltantes de parrafos

import os
import ctypes

import pandas as pd

from tkinter import *
from tkinter import ttk

from PIL import Image, ImageTk  # Importar Pillow

from tkcalendar import DateEntry

from docx import Document 
from docx.shared import Inches

import datetime

#from customtkinter import *

def get_entradassv_var():
    global SSV_var
    SSV_var = entrada_SSV.get() # Variable creada

'''
def get_mod_var():
    global tipo_mod_var
    tipo_mod_var = entryModificacion.get() # Variable creada
'''

def get_av_var():
    global entrada_AV_var
    entrada_AV_var = entrada_AV.get() # Variable creada

def modulo_modificacion(): # PENDIENTE: HACER QUE SE ABRA PESTAÑA O PONER COLUMNA A LA DERECHA
    
    ############ casilla 2

    label_2 = ttk.Label(master = frm_b, text="Solicitud solicitante vía", style="TLabel")
    label_2.grid(row=2, column=0, sticky=N+W)

    global entrada_SSV
    entrada_SSV = StringVar() # podria ser int binario tambien com IntVarblbl() 
    
    memo_2 = Radiobutton(frm_b, text = "Memo", padx = 1, fg="black", variable=entrada_SSV, command=get_entradassv_var, value='Memo')
    memo_2.grid(row=2, column=1, sticky=N+W)
    mail_2 = Radiobutton(frm_b, text = "Correo electrónico", padx = 1, fg="black", variable=entrada_SSV, command=get_entradassv_var, value='Correo electrónico')
    mail_2.grid(row=2, column=2, sticky=N+W)

    '''
    global SSV_var
    entrada_SSV_var = str(entrada_SSV.get()) #Variable creada: entradaElegida
    '''
    ############ casilla 3

    # Dropdown menu options 
    list_3 = ["Beneficio", "Plazo", "Renta", "Cometido"]

    label_3 = Label(master = frm_b, text = "Tipo de Modificación", padx=5)
    label_3.grid(row=4, column=0, sticky=E)

    # Create Dropdown menu 
    global entryModificacion
    entryModificacion = StringVar()
    entryModificacion.set("Seleccionar") 

    menu_3 = ttk.Combobox(frm_b, textvariable = entryModificacion, values = list_3, style="TCombobox") 
    menu_3.grid(row=4, column=1, sticky=W)

    global tipo_mod_var
    tipo_mod_var = "{}".format(entryModificacion.get()) # Variable creada
    menu_3.set(tipo_mod_var) #Variable creada


    ### NO ESTA FUNCIONANDO ### AVANZAR EN ESTO LUEGO DE HACER MODULO DE SALUD/EDUCACION

    # menucombobox.delete("0", tk.END) # this will clear the field after button click

    # casilla 4

    label_4 = Label(master = frm_b, text="Aprobación vía", padx=5)
    label_4.grid(row=5, column=0, sticky=N+W)

    global entrada_AV
    entrada_AV = StringVar() # podria ser int binario tambien com IntVarblbl() 

    memo_4 = Radiobutton(frm_b, text = "Memo", padx = 1, fg="black", variable=entrada_AV, command=get_av_var, value='Memo')
    memo_4.grid(row=13, column=1, sticky=N+W)
    mail_4 = Radiobutton(frm_b, text = "Correo electrónico", padx = 1, fg="black", variable=entrada_AV, command=get_av_var, value='Correo electrónico')
    mail_4.grid(row=13, column=2, sticky=N+W)

    ### FIX: se marca ERROR SSV y MOD ###


    # casilla 5 :  Fecha aprobación modificación

    # Add Calendar
    
    # Add Button and Label

    sig2a_button = Button(frm_b, text='Siguiente 2a', command=det_direccion) 
    sig2a_button.grid(row=132, column=2)

def modulo_modificacion_2():
    '''
    ## Parece que esto va despues
    fecha = cal_5.get_date()
    fecha_label.config(text=f"Selected Date: {fecha_5a_var}")
    '''

    date_5 = Label(frm_a, text='Fecha aprobación modificación')
    date_5.grid(row=13, column=0, sticky=N+W)  

    cal_5 = DateEntry(frm_a, width=12, background='darkblue', foreground='white', borderwidth=2, date_pattern="dd-mm-yyyy")
    cal_5.grid(row=13, column=1, sticky=N+W) 

    get_date_button = Button(frm_a, text="Guardar fecha", command = get_fecha_5a_var)
    get_date_button.grid(row=13, column=2, sticky=N+W) 

    fecha = Label(frm_a, text="") #variabkle?
    fecha.grid(row=13, column=3, sticky=N+W) 

    # casilla 6
    '''
    ## Parece que esto va despues
    fecha_5a_var = cal_6.get_date()
    fecha_5a_var_label_6.config(text=f"Selected Date: {fecha_5a_var}")
    '''

    date_6 = Label(frm_a, text='Fecha instumento contrat.')
    date_6.grid(row=13, column=0, sticky=N+W)  

    cal_6 = DateEntry(frm_a, width=12, background='darkblue', foreground='white', borderwidth=2, date_pattern="dd-mm-yyyy")
    cal_6.grid(row=13, column=1, sticky=N+W) 

    get_date_button = Button(frm_a, text="Guardar fecha", command = get_fecha_5a_var)
    get_date_button.grid(row=13, column=2, sticky=N+W) 

    fecha_5a_var_label_6 = Label(frm_a, text="")
    fecha_5a_var_label_6.grid(row=13, column=3, sticky=N+W) 

    print(fecha_5a_var_label_6)

    # la casilla 7 está por fuera

    # casilla 8

    label_8 = Label(master = frm_a, text="Número de decreto")  
    label_8.grid(row=13, column=0, sticky=N+W)
    
    entry_8 = Entry(master = frm_a)  
    entry_8.grid(row=13, column=1, sticky=N+W)

    nro_decreto = entry_8.get()

    print(nro_decreto)   

    # falta para este modulo la fecha de decreto siaper aprovacion modificacion 
    '''
    sig_button = Button(frm_b, text='Siguiente 2', command=det_direccion) 
    sig_button.grid(row=132, column=2)
    '''

def det_direccion():

    # casilla 7: Dirección solicitante EDITING

    # PENDIENTE: hacer csv de direccionres y subdirecciones, a esto se le podria asociar los programas
    list_7 = ["Dirección de Personas",
            "Secretaría Comunal de Planificación",
            "Dirección de Salud y Educación, y Demás Incorporados en la Gestión Municipal",
            "Dirección de Asesoría Jurídica",
            "Dirección de Sustentabilidad e Innovación",
            "Dirección de Administración y Finanzas",
            "Dirección de Asesoría Urbana",
            "Dirección de Desarrollo Comunitario",
            "Dirección de Informática",
            "Dirección de Obras Municipales",
            "Dirección de Comunicaciones, Asuntos Corporativos y Prensa",
            "Dirección de Medio Ambiente, Aseo y Ornato",
            "Dirección de Infraestructura Comunal",
            "Dirección de Tránsito y Transporte Público",
            "Dirección de Seguridad Pública",
            "Dirección de Control",
            "Secretaría Municipal",
            "Administración Municipal",
            ]

    label_7 = ttk.Label(master = frm_a, text="Dirección solicitante", style="TLabel")
    label_7.grid(row=2, column=0, sticky=E)

    # Create Dropdown menu 
    global direccionSolicitante
    direccionSolicitante = StringVar()
    direccionSolicitante.set("Seleccionar")

    global menu_7
    menu_7 = ttk.Combobox(frm_a, width = 50, textvariable = direccionSolicitante, values = list_7, style="TCombobox")
    menu_7.grid(row=2, column=1, columnspan=3, sticky=W) # columnspan=3 significa que se extiende por 3 columnas

    menu_7.bind("<<ComboboxSelected>>", choice) # Liga una función al evento de selección para manejar lo que sucede cuando el usuario selecciona una opción.
    print("pasó bind")

    #sig2b_button = Button(frm_b, text='Siguiente 2b', command=salud()) 
    #sig2b_button.grid(row=2, column=2)

def choice(event):

    print("entró a choice")

    global direccion_sol_var
    direccion_sol_var = menu_7.get() # Variable creada

    if direccion_sol_var == "Dirección de Salud y Educación, y Demás Incorporados en la Gestión Municipal": 
        print("entro al id para salud/ed")
        label_ems = ttk.Label(master = frm_a, text="Departamento", style="TLabel")
        label_ems.grid(row=2, column=3, sticky=E)

        # Create Dropdown menu 
        global depto
        depto = StringVar()
        depto.set("Seleccionar")

        global menu_ems
        menu_ems = ttk.Combobox(frm_a, textvariable = depto, values = ["Salud", "Educación"], style="TCombobox") # si no escribimos ancho, este se ajusta por si solo
        menu_ems.grid(row=2, column=4, sticky=W)

        menu_ems.bind("<<ComboboxSelected>>", salud) # Liga una función al evento de selección para manejar lo que sucede cuando el usuario selecciona una opción.
        
        print("pasó servicio salud")        

    else:
        depto = "Municipal"

        #else: # ARREGLAR FIX ESTOOOO 

        # ALGO DEBE PASAR PARA QUE AL MOMENTO DE CAMBIAR LA DIRECCION, SE DEJE DE NOSTRAR EL WIDGET DE DEPERTAMENTO SALUD/EDUCACION

    ############ casilla parrafo 5

    ## falta: if para mostrar el tipo de contrato según seleccion de direccion/depto

    label_p5 = ttk.Label(master = frm_a, text="Seleccione (falta if): ", style="TLabel")
    label_p5.grid(row=3, column=0, sticky=E)

    global entrada_prog_esep_sumalz
    entrada_prog_esep_sumalz = StringVar() # podria ser int binario tambien com IntVarblbl() 
            
    p5_prog = Radiobutton(frm_a, text = "Programa", padx = 1, fg="black", variable=entrada_prog_esep_sumalz, command=get_tipo_contrato_var, value='Programa')
    p5_prog.grid(row=3, column=1, sticky=W)
    p5_ed_sep = Radiobutton(frm_a, text = "Educación SEP", padx = 1, fg="black", variable=entrada_prog_esep_sumalz, command=get_tipo_contrato_var, value='Educación SEP')
    p5_ed_sep.grid(row=3, column=2, sticky=W)
    p5_sum_alz = Radiobutton(frm_a, text = "Suma alzada", padx = 1, fg="black", variable=entrada_prog_esep_sumalz, command=get_tipo_contrato_var, value='Suma Alzada')
    p5_sum_alz.grid(row=3, column=3, sticky=W)

def salud(event):
    global depto
    depto = menu_ems.get() # Variable modifica
    print("seleccionaste :", depto)

    if depto == "Salud":
        list_salud = ["COSAM",
                "CESFAM",
                "Departamento de Salud"]

        label_salud = Label(master = frm_a, text="Servicio salud", padx=5)
        label_salud.grid(row=2, column=5, sticky=E)

        # Create Dropdown menu 
        global salud_seleccion
        salud_seleccion = StringVar()
        salud_seleccion.set("Seleccionar")

        global menu_salud
        menu_salud = ttk.Combobox(frm_a, width=15, textvariable = salud_seleccion, values = list_salud, style="TCombobox")
        menu_salud.grid(row=2, column=6, sticky=W)

        #global salud_seleccion_var
        #salud_seleccion_var = menu_salud.get() # Variable creada

        menu_salud.bind("<<ComboboxSelected>>", salud_exito) # Liga una función al evento de selección para manejar lo que sucede cuando el usuario selecciona una opción.
        
        print("pasó salud")

        #sig2c_button = Button(frm_b, text='Siguiente 2c', command=salud) 
        #sig2c_button.grid(row=13, column=3)
    else:
        None

def salud_exito(event):
    global salud_seleccion_var
    salud_seleccion_var = menu_salud.get()
    print("Servicio de salud seleccionado: ", salud_seleccion_var)
    
# imprime en terminal, luego tengo que pasarlo a un variable

def funcion_tipo_decreto(event): 

    global tipo_decreto_var
    tipo_decreto_var = entradaElegida.get()
    print("Tipo de decreto seleccionado: ", tipo_decreto_var)

    #sig1_button.grid_forget()

    if tipo_decreto_var == "Modificación":
        print("entró")
        modulo_modificacion() # LUEGO DEBO ENCADENWAR ESTA FUNCION 
    else:
        det_direccion()

    '''
    # REVISAR FIX ESTO ERROR EN QYE NO RECONOCE VARIABLE, PERO LA ESCRIBE IGUAL?
    print("Tipo de decreto: ", tipo_decreto_var)
    print("Solicitud solicitante vía: ", "SSV_var") 
    print("Tipo de modificación: ", tipo_mod_var)
    print("Aprobación vía: ", entrada_AV_var)
    #print("Fecha aprobación modificación: {}".format(fecha_5a_var_label_6)) 
    '''

def get_tipo_contrato_var(): # funcion ejecutada al seleccionar el tipo de contrato (5: a, b, c , d)
    global tipo_contrato_var 
    tipo_contrato_var = entrada_prog_esep_sumalz.get()

    print("tipo de contrato: ", tipo_contrato_var)

    # Separador: "|",  a modificar:  <<pendiente: para programas y ed sep>> 
    vyc_5_csv = pd.read_csv("vyc_5.csv", sep="|", encoding="utf-8", header=None, usecols=[1], names=["texto"]) 

    # escribir el docx, dependiendo si es municipal, salud o educacion
    vyc_5_csv["texto"] = vyc_5_csv["texto"].astype("string") # hacemos columna string, podria haberlo hecho en el read_csv

    print("pre 5")

    global content_5
    ### escribimos en el word cambiando los parametros
    if tipo_contrato_var == "Suma Alzada": # al parecer los suma alzada solo son los municipales
        if depto == "Municipal": # no deberia haber depto salud/educacion en suma alzada, pero nos aseguramos por si acaso
            content_5 = vyc_5_csv["texto"][3] # fila para tipo municipal, educacion no tiene este parrafo al parecer
    elif tipo_contrato_var == "Programa":
        if depto == "Salud":
            content_5 = vyc_5_csv["texto"][2] # fila para tipo salud
        elif depto == "Educación" or depto == "Municipal":

            clausula_5a() # se abre ventana para ingresar datos de la casilla 5

            # Convertir date (objeto) a string
            global date_str_5a1
            global date_str_5a2
            date_str_5a1 = fecha_5a1_var.strftime("%d-%m-%Y")
            date_str_5a2 = fecha_5a2_var.strftime("%d-%m-%Y")

            # Separador: "|",  a modificar:  <<nombre_programa>> <<nro_decreto_5a1>> <<fecha_decreto_5a_dma>>
            vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("<<nombre_programa>>", entrada_programa_var)
            vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("<<nro_decreto_5a1>>", nro_decreto_5a1_var)
            vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("<<fecha_decreto_5a1_dma>>", date_str_5a1)
            vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("<<nro_decreto_5a2>>", nro_decreto_5a2_var)
            vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("<<fecha_decreto_5a2_dma>>", date_str_5a1) # reemplazamos los parametros

            content_5 = vyc_5_csv["texto"][0] # fila para tipo programa no sep (educacion o municipal)

    elif tipo_contrato_var == "Educación SEP":

        clausula_5b() # se abre ventana para ingresar datos de la casilla 5

        global date_str_5b
        date_str_5b = fecha_5b_var.strftime("%d-%m-%Y")

        vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("<<nro_memo_dem>>", nro_memoDEM_5b_var)
        vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("<<fecha_memo_dma>>", fecha_5b_var)
        vyc_5_csv["texto"] = vyc_5_csv["texto"].str.replace("<<direccion_sol_var>>", direccion_sol_var) # reemplazamos los parametros

        content_5 = vyc_5_csv["texto"][1] # fila para tipo educacion SEP
    else:
        content_5 = "ERROR: Debe seleccionar un 'Tipo de contrato'" # en caso de error

    print("post 5")
    print(content_5)
    print(type(content_5))

def clausula_5a():

    ## Nombre del programa
    nombre_programa = ttk.Label(master = frm_a, text='Nombre del programa', style="TLabel") # cuando tengamos la BBDD con los programas, se podrá seleccionar
    nombre_programa.grid(row=4, column=0, sticky=E)

    entrada_programa = Entry(frm_a)
    entrada_programa.grid(row=4, column=1, sticky=W)

    global entrada_programa_var
    entrada_programa_var = entrada_programa.get()

    ## Decreto creación (5a.1)
    nro_decreto_5a1 = Label(frm_a, text='Número decreto creación programa') # cuando tengamos la BBDD con los programas, se podrá seleccionar
    nro_decreto_5a1.grid(row=5, column=0, sticky=E)

    entrada_nro_decreto_5a1 = Entry(frm_a)
    entrada_nro_decreto_5a1.grid(row=5, column=1, sticky=W)

    global nro_decreto_5a1_var
    nro_decreto_5a1_var = entrada_nro_decreto_5a1.get()

    print(nro_decreto_5a1_var)
    
    ## Fecha decreto creación (5a.1)

    date_5a1 = Label(frm_a, text='Fecha creación decreto programa')
    date_5a1.grid(row=5, column=2, sticky=E)  

    cal_5a1 = DateEntry(frm_a, width=12, background='darkblue', foreground='white', borderwidth=2, date_pattern="dd/mm/yyyy")
    cal_5a1.grid(row=5, column=3, sticky=W) 

    global fecha_5a1_var
    fecha_5a1_var = cal_5a1.get_date()

    print(fecha_5a1_var)

    ## Decreto moficación de programa (5a.2)

    nro_decreto_5a2 = Label(frm_a, text='Número modificación de decreto programa') # cuando tengamos la BBDD con los programas, se podrá seleccionar
    nro_decreto_5a2.grid(row=6, column=0, sticky=E)

    entrada_nro_decreto_5a2 = Entry(frm_a)
    entrada_nro_decreto_5a2.grid(row=6, column=1, sticky=W)

    global nro_decreto_5a2_var
    nro_decreto_5a2_var = entrada_nro_decreto_5a2.get()

    print(nro_decreto_5a2_var)

    ## Fecha decreto modificación de programa (5a.2)

    date_5a2 = Label(frm_a, text='Fecha modificación decreto programa')
    date_5a2.grid(row=6, column=2, sticky=E)  

    cal_5a2 = DateEntry(frm_a, width=12, background='darkblue', foreground='white', borderwidth=2, date_pattern="dd/mm/yyyy")
    cal_5a2.grid(row=6, column=3, sticky=W) 

    global fecha_5a2_var
    fecha_5a2_var = cal_5a2.get_date()

    print(fecha_5a2_var)

def clausula_5b():

    ## Número de memo DEM (5b)
    nro_memoDEM_5b = ttk.Label(master = frm_a, text='Número memo DEM', style = "TLabel") # cuando tengamos la BBDD con los programas, se podrá seleccionar
    nro_memoDEM_5b.grid(row=4, column=0, sticky=E)

    entrada_nro_memoDEM_5b = Entry(frm_a)
    entrada_nro_memoDEM_5b.grid(row=4, column=1, sticky=W)

    global nro_memoDEM_5b_var
    nro_memoDEM_5b_var = entrada_nro_memoDEM_5b.get()

    print(nro_memoDEM_5b_var)
    
    ## Fecha memo DEM (5b)

    date_5b = Label(frm_a, text="Fecha memeo DEM")
    date_5b.grid(row=4, column=2, sticky=E)  

    cal_5b = DateEntry(frm_a, width=12, background='darkblue', foreground='white', borderwidth=2, date_pattern="dd/mm/yyyy")
    cal_5b.grid(row=4, column=3, sticky=W) 

    global fecha_5b_var
    fecha_5b_var = cal_5b.get_date()

    print(fecha_5b_var)

# path
# path = os.getcwd()
# print(path)

# OJO: todo se está guardando en "c:\Users\jfz\" desde el equino, no en drive

# mejora calidad de imagen de la interfaz a crear
ctypes.windll.shcore.SetProcessDpiAwareness(True)

# creación ventana
# ventana = CTk() custom, para el final, modificar alfinal
ventana = Tk()
ventana.geometry("1700x800")
#ventana.configure(bg="#00c4b4") # Change the background color using configure
ventana.title("Automatizacón decretos y contratos a honorarios") # titulo de la ventana
#.set_appearance_mode("dark") # modo oscyro 

# Crear estilo para personalizar widgets ttk
'''
style = ttk.Style()
style.configure("TButton", background="#00c4b4", foreground="black") # highlightbackground
style.configure("TLabel", background="#00c4b4", foreground="black")
#style.configure("Outline.TButton", background="#00c4b4", foreground="black")
#style.configure("Link.TButton", background="#00c4b4", foreground="black")
style.configure("TCombobox", backgroud="#00c4b4" , fieldbackground="#00c4b4")
'''

# creación de 3 frames (grupos/secciones)

frm_a = ttk.Frame(ventana) # frame a
frm_b = ttk.Frame(ventana) # frame b
#frm_c = ttk.Frame() # frame c

#frm_a.grid_columnconfigure(0, weight=1)
#frm_a.grid_columnconfigure(0, weight=2)

frm_a.grid_rowconfigure(1, minsize=10)
frm_a.place(x = 50, y = 50)

frm_b.grid_rowconfigure(2, minsize=10)
frm_b.place(x = 50, y = 500)

### frm_a.grid_rowconfigure(1, minsize=10)

###################################
########## LOGO VITACURA ##########
###################################

# Cargar imagen PNG usando Pillow
imagen = Image.open("logos-vitacura_sineslogan_hor.png")

# Redimensionar la imagen
imagen = imagen.resize((121, 40))  # Nuevo tamaño (ancho, alto), resize usanfo cambiar tamaño imagen con 30% de la original

# Convertir la imagen redimensionada en un formato que tkinter pueda usar
imagen_tk = ImageTk.PhotoImage(imagen)

# Crear un Label y asignar la imagen
label_imagen = Label(frm_a, image=imagen_tk)
label_imagen.grid(row=0, column=0, columnspan=2, padx=1, sticky=N+W) # distancia para abajo

############ casilla 1

list_1 = ["En fecha", "Regularización", "Modificación"]

label_1 = ttk.Label(master = frm_a, text="Tipo de decreto", style="TLabel")
label_1.grid(row=1, column=0, sticky=E)

# Create Dropdown menu 
entradaElegida = StringVar()
entradaElegida.set("Seleccionar")

tipo_elegido = ttk.Combobox(frm_a, textvariable = entradaElegida, values = list_1, style="TCombobox") 
tipo_elegido.grid(row=1, column=1, sticky=W)

tipo_elegido.bind("<<ComboboxSelected>>", funcion_tipo_decreto)

# arreglar espacios grises quitando el sticky de los botones

#################################### en adelante intergrar respuestas de acasillas anteriores ########################

# REVISAR VALORES DE VARIABLES, NO ESTAN BIEN, ESTO PARA VIERNES 6

# Submit button 
# Whenever we click the submit button, our submitted 
# option is printed ---Testing purpose 

'''
sig1_button = Button(frm_b, text='Siguiente 1', command=funcion_tipo_decreto) 
sig1_button.grid(row=130, column=1)
'''

# se abre ventana
# debe posicoinarse despues de llamar a todos los widgets

ventana.mainloop()

'''
# casilla cerrar: quit

# escribinos frame c con boton "quit"

label_c = ttk.Label(master = frm_c, text = "Vistos y considerando")
boton_c = ttk.Button(master = frm_c, text = "Quit", command = ventana.destroy)
label_c.pack()
boton_c.pack()

'''

###################################
########## CREACIÓN WORD ##########
###################################

document = Document() 
document.add_heading('Vitacura', 0)

# leemos parafos y escribimos vistos y condiciones

######################################
# vistos y considerando, parrafo 1 y 2
######################################

vyc_1_2 = open("vyc_1+2_siempre.dat", "r", encoding='utf-8')
content_1_2 = vyc_1_2.readlines() # lista con lineas del texto, finaliza con \n
vyc_1_2.close()

# escribir el docx
for nro_parrafo in content_1_2:
    document.add_paragraph(nro_parrafo)

###################################
# vistos y considerando, parrafo 3
###################################
'''
vyc_3 = open("vyc_3.dat", "r", encoding='utf-8')
content_3 = vyc_3.readlines() # lista con lineas del texto, finaliza con \n
vyc_3.close()
'''

# Separador: "|",  a modificar:  <<serviciosalud>> 
vyc_3_csv = pd.read_csv("vyc_3.csv", sep="|", encoding="utf-8", header=None, usecols=[1], names=["texto"]) 

# escribir el docx, dependiendo si es municipal, salud o educacion

vyc_3_csv["texto"] = vyc_3_csv["texto"].astype("string") # hacemos columna string, podria haberlo hecho en el read_csv

if direccion_sol_var == "Dirección de Salud y Educación, y Demás Incorporados en la Gestión Municipal": 
    print("depto: ", depto)
    if depto == "Salud":
        vyc_3_csv["texto"] = vyc_3_csv["texto"].str.replace("<<serviciosalud>>", salud_seleccion_var) # str.replace reemplaza incluso dentro de la oracion a diferencia de solamente replace que requiere que toda la celda sea igual
        content_3 = vyc_3_csv["texto"][2] # fila para tipo salud
    elif depto == "Educación":
        content_3 = vyc_3_csv["texto"][1] # fila para tipo educacion
else: # (depto = "Municipal")
    print("depto: ", depto)
    content_3 = vyc_3_csv["texto"][0] # fila para tipo municipales

document.add_paragraph(content_3)

###################################
# vistos y considerando, parrafo 4
###################################

vyc_4 = open("vyc_4_siempre.dat", "r", encoding='utf-8')
content_4 = vyc_4.readlines() # lista con lineas del texto, finaliza con \n
vyc_4.close()

# escribir el docx
for nro_parrafo in content_4:
    document.add_paragraph(nro_parrafo)

###################################
# vistos y considerando, parrafo 5
###################################


## Trabajando aqui, en la funcion 

document.add_paragraph(content_5)

###################################
# vistos y considerando, parrafo 6
###################################

# vistos y considerando, parrafo 7
# vistos y considerando, parrafo 8
# vistos y considerando, parrafo 9
# vistos y considerando, parrafo 10


# guardar el doc

document.save("test_autoPersona.docx") 

# al final debo hace un packing .exe